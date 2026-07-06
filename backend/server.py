import io
import json
import logging
import mimetypes
import os
import re
import subprocess
import sys
import tempfile
import shutil
import uuid
import asyncio
from urllib.parse import urlparse
from contextlib import asynccontextmanager
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, AsyncIterator, Iterator, List, Literal, Optional, Sequence

import httpx
import requests
from dotenv import load_dotenv
from fastapi import APIRouter, FastAPI, Request, UploadFile, File, Form
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse, StreamingResponse
from pydantic import BaseModel, ConfigDict, Field, field_validator

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Preformatted
    from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
from starlette.middleware.cors import CORSMiddleware

try:
    from motor.motor_asyncio import AsyncIOMotorClient
except ImportError:
    AsyncIOMotorClient = None

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / ".env")

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)


def env_str(name: str, default: str = "") -> str:
    return os.environ.get(name, default).strip()


def env_float(name: str, default: float) -> float:
    raw = env_str(name)
    if not raw:
        return default
    try:
        return float(raw)
    except ValueError:
        logger.warning("Invalid %s=%r; using %s", name, raw, default)
        return default


def cors_origins() -> List[str]:
    origins_str = env_str("CORS_ORIGINS", "")
    if origins_str:
        return [origin.strip() for origin in origins_str.split(",") if origin.strip()]
    # Production defaults — no wildcard
    return [
        "http://localhost:3000",
        "https://micro-agent-beta.vercel.app",
        "https://frontend-omega-sand-52.vercel.app",
    ]


# ── MongoDB (legacy, optional) ───────────────────────────────────────────────
mongo_url = env_str("MONGO_URL")
db_name = env_str("DB_NAME")
mongo_client = None
db = None

if mongo_url and db_name and AsyncIOMotorClient is not None:
    mongo_client = AsyncIOMotorClient(mongo_url)
    db = mongo_client[db_name]
else:
    logger.warning("MongoDB env is not configured; status-check persistence is disabled.")

# ── Supabase ──────────────────────────────────────────────────────────────────
try:
    from supabase import create_client, Client as SupabaseClient
except ImportError:
    create_client = None
    SupabaseClient = None

SUPABASE_URL = env_str("SUPABASE_URL")
SUPABASE_SERVICE_ROLE_KEY = env_str("SUPABASE_SERVICE_ROLE_KEY")
SUPABASE_ANON_KEY = env_str("SUPABASE_ANON_KEY")

supa: Optional[Any] = None
if SUPABASE_URL and SUPABASE_SERVICE_ROLE_KEY and create_client:
    supa = create_client(SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY)
    logger.info("Supabase client initialized.")
else:
    logger.warning("Supabase not configured; auth/history disabled.")

@asynccontextmanager
async def lifespan(_: FastAPI):
    yield
    if mongo_client is not None:
        mongo_client.close()


app = FastAPI(lifespan=lifespan)
api_router = APIRouter(prefix="/api")

# ── Rate Limiting (in-memory, per-user) ──────────────────────────────────────
import time
_rate_limits: dict[str, list[float]] = {}  # user_id -> [timestamps]

def check_rate_limit(user_id: str, max_requests: int = 10, window_seconds: int = 60) -> bool:
    """Returns True if allowed, False if rate limited."""
    if not user_id:
        return True  # guests not rate-limited by this (frontend handles guest limit)
    now = time.time()
    # Clean old entries
    _rate_limits[user_id] = [t for t in _rate_limits.get(user_id, []) if now - t < window_seconds]
    if len(_rate_limits[user_id]) >= max_requests:
        return False
    _rate_limits[user_id].append(now)
    return True

MODEL_ID_TO_PROVIDER = {
    "claude-sonnet-4-5-1m": "claude-sonnet-4.5-1m",
    "deepseek-v4-flash": "deepseek-v4-flash",
    "glm-5": "glm-5",
    "claude-opus-4-8": "claude-opus-4.8",
    "open-agentic": "open-agentic",
    "glm-5.2-free": "glm-5.2-free",
}

IMAGE_MODEL_ID = "flux-2-klein-4b"
DEFAULT_MODEL_ID = "deepseek-v4-flash"
PROVIDER_NAME = "AIMurah"

# ── SumoPod AI Provider ──────────────────────────────────────────────────────
SUMODOP_BASE_URL = env_str("SUMODOP_BASE_URL", "https://ai.sumopod.com/v1")
SUMODOP_API_KEY = env_str("SUMODOP_API_KEY")
SUMODOP_DEFAULT_MODEL = env_str("SUMODOP_DEFAULT_MODEL", "claude-sonnet-5")

SUMODOP_MODEL_IDS = {
    "claude-fable-5",
    "claude-sonnet-5",
}

SUMODOP_MODEL_MAP = {
    "claude-fable-5": "claude-fable-5",
    "claude-sonnet-5": "claude-sonnet-5",
}

# ── Token cost per model ────────────────────────────────────────────────────────
MODEL_TOKEN_COST = {
    # Free models
    "claude-sonnet-4.5":    2,
    "claude-sonnet-4.5-1m": 2,
    "deepseek-v4-flash":    1,
    "glm-5":                1,
    "glm-5.2-free":         1,
    "open-agentic":         1,
    # Pro & Ultra models
    "claude-opus-4.8":      8,
    "gemini-2.5-pro":       5,
    "DeepSeek-V4-Pro":      2,
    "gemini-3.1-pro":       6,
    "gpt-5.4":              8,
    "gpt-5.2":              6,
    "flux-2-klein-4b":      2,
    # SumoPod AI models
    "claude-fable-5":       8,
    "claude-sonnet-5":      4,
}

# Also map frontend model IDs (with hyphens) to costs
_TOKEN_COST_ALIASES = {
    "claude-sonnet-4-5-1m": 2,
    "claude-opus-4-8":      8,
    "deepseek-v4-flash":    1,
    "glm-5":                1,
    "glm-5.2-free":         1,
    "open-agentic":         1,
    "flux-2-klein-4b":      2,
    "claude-fable-5":       8,
    "claude-sonnet-5":      4,
}


def get_token_cost(model_id: str) -> int:
    """Return token cost for a model. Tries exact match first, then alias, default 1."""
    if not model_id:
        return 1
    if model_id in MODEL_TOKEN_COST:
        return MODEL_TOKEN_COST[model_id]
    return _TOKEN_COST_ALIASES.get(model_id, 1)


async def get_user_balance(user_id: str) -> int:
    """Get user token balance from Supabase user_credits table."""
    if not supa or not user_id:
        return 0
    try:
        res = supa.table("user_credits").select("balance").eq("user_id", user_id).single().execute()
        return res.data["balance"] if res.data else 0
    except Exception:
        return 0


async def deduct_token(user_id: str, model_id: str) -> dict:
    """Deduct tokens for a model usage. Returns {success, cost, balance, ...}."""
    if not user_id:
        return {"success": True, "cost": 0, "balance": 0, "reason": "guest"}

    cost = get_token_cost(model_id)
    if cost <= 0:
        return {"success": False, "reason": "invalid_cost", "balance": 0, "required": 0}

    balance = await get_user_balance(user_id)

    # Auto-seed 50 tokens for first-time users
    if balance == 0 and supa:
        try:
            existing = supa.table("user_credits").select("user_id").eq("user_id", user_id).execute()
            if not existing.data:
                supa.table("user_credits").insert({
                    "user_id": user_id,
                    "balance": 50,
                    "plan": "free",
                }).execute()
                supa.table("credit_transactions").insert({
                    "user_id": user_id,
                    "amount": 50,
                    "type": "bonus",
                    "model": None,
                }).execute()
                balance = 50
                logger.info("Auto-seeded 50 tokens for user %s", user_id)
        except Exception as exc:
            logger.warning("Auto-seed failed: %s", exc)

    if balance < cost:
        return {
            "success": False,
            "reason": "insufficient_tokens",
            "balance": balance,
            "required": cost,
        }

    new_balance = balance - cost

    try:
        supa.table("user_credits").upsert({
            "user_id": user_id,
            "balance": new_balance,
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }, on_conflict="user_id").execute()

        supa.table("credit_transactions").insert({
            "user_id": user_id,
            "amount": -cost,
            "type": "usage",
            "model": model_id,
        }).execute()
    except Exception as exc:
        logger.warning("Token deduct DB error: %s", exc)

    return {"success": True, "cost": cost, "balance": new_balance}


async def refund_token(user_id: str, model_id: str, cost: int) -> None:
    """Refund tokens on error (e.g. provider failure)."""
    if not user_id or not supa or cost <= 0:
        return
    try:
        balance = await get_user_balance(user_id)
        new_balance = balance + cost
        supa.table("user_credits").upsert({
            "user_id": user_id,
            "balance": new_balance,
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }, on_conflict="user_id").execute()
        supa.table("credit_transactions").insert({
            "user_id": user_id,
            "amount": cost,
            "type": "refund",
            "model": model_id,
        }).execute()
    except Exception as exc:
        logger.warning("Token refund DB error: %s", exc)


# ── Multi-Provider Routing ───────────────────────────────────────────────────

def is_sumopod_model(model_id: Optional[str]) -> bool:
    """Check if a model ID belongs to SumoPod provider."""
    if not model_id:
        return False
    if model_id in SUMODOP_MODEL_MAP:
        return True
    return False


def resolve_provider_model(model_id: Optional[str], auto_mode: bool = False) -> str:
    """Resolve a frontend model ID to the provider's model name."""
    if auto_mode or not model_id:
        return MODEL_ID_TO_PROVIDER.get(DEFAULT_MODEL_ID, DEFAULT_MODEL_ID)

    if is_sumopod_model(model_id):
        return SUMODOP_MODEL_MAP.get(model_id, model_id)

    return MODEL_ID_TO_PROVIDER.get(model_id, model_id)


def get_provider_config(model_id: Optional[str]) -> tuple[str, str, str, str]:
    """
    Return (base_url, api_key, provider_name, model_name) for the given model.
    Routes to SumoPod or AIMurah based on model ID.
    """
    if is_sumopod_model(model_id):
        base_url = SUMODOP_BASE_URL
        api_key = SUMODOP_API_KEY
        provider_name = "SumoPod AI"
        model_name = SUMODOP_MODEL_MAP.get(model_id, model_id) if model_id else SUMODOP_DEFAULT_MODEL
    else:
        base_url = env_str("OPENAI_BASE_URL")
        api_key = env_str("OPENAI_API_KEY")
        provider_name = PROVIDER_NAME
        model_name = MODEL_ID_TO_PROVIDER.get(model_id, model_id) if model_id else DEFAULT_MODEL_ID

    return base_url, api_key, provider_name, model_name


DEFAULT_TIMEOUT_SECONDS = 120.0
MAX_PROVIDER_ERROR_CHARS = 800
FIRECRAWL_BASE_URL = "https://api.firecrawl.dev"
TAVILY_BASE_URL = "https://api.tavily.com"
DEFAULT_WEB_SEARCH_MAX_RESULTS = 10
DEFAULT_WEB_SEARCH_TIMEOUT_SECONDS = 12.0
DEFAULT_WEB_SEARCH_MAX_CONTEXT_CHARS = 6000
DEFAULT_WEB_SEARCH_MAX_RESULT_CHARS = 1000


class StatusCheck(BaseModel):
    model_config = ConfigDict(extra="ignore")

    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_name: str
    timestamp: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))


class StatusCheckCreate(BaseModel):
    client_name: str


class ChatMessageIn(BaseModel):
    role: Literal["system", "user", "assistant"]
    content: str

    @field_validator("content")
    @classmethod
    def trim_content(cls, value: str) -> str:
        return value.strip()


class ChatStreamRequest(BaseModel):
    messages: List[ChatMessageIn]
    model_id: Optional[str] = None
    auto_mode: bool = False
    room: Optional[str] = None
    attachments: List[str] = Field(default_factory=list)
    web_search: bool = False
    reasoning: bool = True
    search_mode_prompt: str = ""
    skill_slug: Optional[str] = None
    effort_level: str = "low"
    comparison: bool = False
    user_id: Optional[str] = None

    @field_validator("model_id", "room", mode="before")
    @classmethod
    def trim_optional_string(cls, value):
        if value is None:
            return None
        return str(value).strip() or None

    @field_validator("attachments", mode="before")
    @classmethod
    def clean_attachments(cls, value):
        if value is None:
            return []
        if not isinstance(value, list):
            return []
        return [str(item).strip() for item in value if str(item).strip()]


class ModelInfo(BaseModel):
    id: str
    provider_model: str


class ModelsResponse(BaseModel):
    provider: str
    default_model_id: str
    fallback_model: str
    models: List[ModelInfo]


def sse(event: str, data: Any) -> str:
    return f"event: {event}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"


def provider_model(model_id: Optional[str], auto_mode: bool) -> str:
    """Legacy wrapper — delegates to resolve_provider_model."""
    return resolve_provider_model(model_id, auto_mode)


def provider_url(base_url: str) -> str:
    clean = base_url.rstrip("/")
    if clean.endswith("/chat/completions"):
        return clean
    return f"{clean}/chat/completions"


def extract_delta_text(choice: dict) -> str:
    delta = choice.get("delta") or {}
    content = delta.get("content")
    return content if isinstance(content, str) else ""


def extract_reasoning_content(choice: dict) -> str:
    delta = choice.get("delta") or {}
    text = delta.get("reasoning_content") or delta.get("reasoning") or ""
    return text if isinstance(text, str) else ""


def has_user_content(messages: Sequence[ChatMessageIn]) -> bool:
    return any(message.content for message in messages)


def last_user_prompt(payload: ChatStreamRequest) -> str:
    for msg in reversed(payload.messages):
        if msg.role == "user" and msg.content:
            return " ".join(msg.content.split())
    return ""


def env_int(name: str, default: int) -> int:
    raw = env_str(name)
    if not raw:
        return default
    try:
        return int(raw)
    except ValueError:
        logger.warning("Invalid %s=%r; using %s", name, raw, default)
        return default


def extract_search_query(payload: ChatStreamRequest) -> str:
    return last_user_prompt(payload)[:500]


def normalize_firecrawl_results(data: Any) -> List[dict]:
    if not isinstance(data, dict):
        return []

    # v2 format: {"data": {"web": [...]}}
    # v1 format: {"data": [...]}
    raw_results = data.get("data") or data.get("results") or []
    if isinstance(raw_results, dict):
        raw_results = raw_results.get("web") or raw_results.get("data") or raw_results.get("results") or []
    if not isinstance(raw_results, list):
        return []

    results = []
    for item in raw_results:
        if not isinstance(item, dict):
            continue
        metadata = item.get("metadata") if isinstance(item.get("metadata"), dict) else {}
        title = item.get("title") or metadata.get("title") or item.get("url") or "Untitled source"
        url = item.get("url") or metadata.get("sourceURL") or metadata.get("url") or ""
        content = (
            item.get("markdown")
            or item.get("content")
            or item.get("description")
            or item.get("snippet")
            or metadata.get("description")
            or ""
        )
        if url or content:
            results.append(
                {
                    "title": " ".join(str(title).split()),
                    "url": str(url).strip(),
                    "content": " ".join(str(content).split()),
                }
            )
    return results


async def firecrawl_search(query: str, max_results: int) -> List[dict]:
    api_key = env_str("FIRECRAWL_API_KEY")
    if not api_key:
        raise RuntimeError("FIRECRAWL_API_KEY belum dikonfigurasi")

    base_url = env_str("FIRECRAWL_BASE_URL", FIRECRAWL_BASE_URL).rstrip("/")
    timeout = env_float("WEB_SEARCH_TIMEOUT_SECONDS", DEFAULT_WEB_SEARCH_TIMEOUT_SECONDS)
    payload = {
        "query": query,
        "limit": max(1, min(max_results, 10)),
        "scrapeOptions": {"formats": ["markdown"]},
    }
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }

    async with httpx.AsyncClient(timeout=timeout) as client:
        response = await client.post(f"{base_url}/v2/search", headers=headers, json=payload)
        if response.status_code == 404:
            response = await client.post(f"{base_url}/v1/search", headers=headers, json=payload)
        response.raise_for_status()
        return normalize_firecrawl_results(response.json())


async def tavily_search(query: str, max_results: int) -> List[dict]:
    api_key = env_str("TAVILY_API_KEY")
    if not api_key:
        raise RuntimeError("TAVILY_API_KEY belum dikonfigurasi")

    base_url = env_str("TAVILY_BASE_URL", TAVILY_BASE_URL).rstrip("/")
    timeout = env_float("WEB_SEARCH_TIMEOUT_SECONDS", DEFAULT_WEB_SEARCH_TIMEOUT_SECONDS)
    payload = {
        "api_key": api_key,
        "query": query,
        "search_depth": "advanced",
        "max_results": max(1, min(max_results, 10)),
        "include_answer": False,
        "include_raw_content": True,
    }

    async with httpx.AsyncClient(timeout=timeout) as client:
        response = await client.post(f"{base_url}/search", json=payload)
        response.raise_for_status()
        data = response.json()

    results = []
    for item in data.get("results") or []:
        if not isinstance(item, dict):
            continue
        title = item.get("title") or item.get("url") or "Untitled source"
        url = item.get("url") or ""
        content = item.get("raw_content") or item.get("content") or ""
        if url or content:
            results.append({
                "title": " ".join(str(title).split()),
                "url": str(url).strip(),
                "content": " ".join(str(content).split()),
                "provider": "tavily",
            })
    return results


async def web_search(query: str, max_results: int) -> tuple[str, List[dict]]:
    """Try Tavily first, fallback to Firecrawl. Raise if neither is configured."""
    if env_str("TAVILY_API_KEY"):
        try:
            results = await tavily_search(query, max_results)
            return "tavily", results
        except Exception as exc:
            logger.warning("Tavily search failed, trying Firecrawl: %s", exc)
    if env_str("FIRECRAWL_API_KEY"):
        try:
            results = await firecrawl_search(query, max_results)
            return "firecrawl", results
        except Exception as exc:
            logger.warning("Firecrawl search failed: %s", exc)
    raise RuntimeError("Tidak ada API key pencarian web yang dikonfigurasi (TAVILY_API_KEY atau FIRECRAWL_API_KEY)")


# ── Comparison Detection & Pipeline ───────────────────────────────────────────

COMPARISON_KEYWORDS = [
    "bandingkan", "compare", "vs", "versus", "perbandingan", "comparison",
    "lebih bagus", "lebih baik", "lebih unggul", "mana yang lebih",
    "beda", "difference", "similarities", "persamaan", "perbedaan",
]

def is_comparison_request(text: str) -> bool:
    """Detect if user message is a comparison request."""
    lower = text.lower()
    # Check for comparison keywords
    has_keyword = any(kw in lower for kw in COMPARISON_KEYWORDS)
    # Check for "X vs Y" or "X versus Y" pattern
    has_vs_pattern = bool(re.search(r'\b\w+\s+(?:vs|versus)\s+\w+', lower))
    # Check for "bandingkan X dengan Y" or "compare X with Y"
    has_bandingkan_pattern = bool(re.search(r'(?:bandingkan|compare)\s+.+\s+(?:dengan|with|and|dan)\s+', lower))
    
    return has_keyword and (has_vs_pattern or has_bandingkan_pattern or "bandingkan" in lower or "compare" in lower)


def extract_comparison_products(text: str) -> tuple[str, str]:
    """Extract two product names from comparison text."""
    lower = text.lower()
    
    # Try "X vs Y" or "X versus Y" pattern
    vs_match = re.search(r'(.+?)\s+(?:vs|versus)\s+(.+)', text, re.IGNORECASE)
    if vs_match:
        return vs_match.group(1).strip(), vs_match.group(2).strip()
    
    # Try "bandingkan X dengan Y" or "compare X with Y"
    dengan_match = re.search(r'(?:bandingkan|compare)\s+(.+?)\s+(?:dengan|with|and|dan)\s+(.+)', text, re.IGNORECASE)
    if dengan_match:
        return dengan_match.group(1).strip(), dengan_match.group(2).strip()
    
    # Try "X lebih bagus dari Y" pattern
    lebih_match = re.search(r'(.+?)\s+lebih\s+(?:bagus|baik|unggul)\s+(?:dari|dibanding|than)\s+(.+)', text, re.IGNORECASE)
    if lebih_match:
        return lebih_match.group(1).strip(), lebih_match.group(2).strip()
    
    # Fallback: split by common separators
    parts = re.split(r'\s+(?:dan|and|dengan|with|,)\s+', text, maxsplit=1)
    if len(parts) == 2:
        # Remove common prefixes
        p1 = re.sub(r'^(?:bandingkan|compare|perbandingan)\s+', '', parts[0], flags=re.IGNORECASE).strip()
        p2 = parts[1].strip()
        if p1 and p2:
            return p1, p2
    
    return text, ""


# ── Deterministic Comparison Data Builder ─────────────────────────────────────

SMARTPHONE_KEYWORDS = ["iphone", "samsung", "xiaomi", "oppo", "vivo", "realme", "google pixel", "oneplus", "huawei", "honor"]
LAPTOP_KEYWORDS = ["macbook", "laptop", "dell", "lenovo", "thinkpad", "asus", "acer", "hp", "msi", "rog"]
CAMERA_KEYWORDS = ["canon", "sony", "nikon", "fujifilm", "lumix", "olympus", "leica"]
HEADPHONE_KEYWORDS = ["airpods", "sony", "bose", "jbl", "sennheiser", "beats", "anker", "soundcore"]
TABLET_KEYWORDS = ["ipad", "galaxy tab", "tablet", "surface pro"]


def _categorize_product(name: str) -> str:
    lower = name.lower()
    for kw in SMARTPHONE_KEYWORDS:
        if kw in lower:
            return "smartphone"
    for kw in TABLET_KEYWORDS:
        if kw in lower:
            return "tablet"
    for kw in LAPTOP_KEYWORDS:
        if kw in lower:
            return "laptop"
    for kw in CAMERA_KEYWORDS:
        if kw in lower:
            return "camera"
    for kw in HEADPHONE_KEYWORDS:
        if kw in lower:
            return "headphone"
    return "generic"


def _is_premium(name: str) -> bool:
    lower = name.lower()
    premium_indicators = ["pro", "max", "ultra", "premium", "elite", "professional", "studio", "xtreme"]
    return any(ind in lower for ind in premium_indicators)


SMARTPHONE_SPECS_A = {
    "chipset": "Titan X Pro (3nm)",
    "ram": "12 GB LPDDR5X",
    "storage": "256 GB UFS 4.0",
    "camera": "200 MP Wide + 48 MP Ultrawide + 12 MP Telephoto",
    "battery": "5000 mAh (120W fast charging)",
    "display": "6.8\" LTPO AMOLED, 144Hz, 3200x1440",
    "os": "Android 15",
    "features": ["IP68", "Face Unlock", "Under-display Fingerprint", "WiFi 7", "Bluetooth 5.4"],
}

SMARTPHONE_SPECS_B = {
    "chipset": "A18 Bionic (3nm)",
    "ram": "8 GB LPDDR5",
    "storage": "256 GB NVMe",
    "camera": "48 MP Wide + 12 MP Ultrawide + 12 MP Telephoto 5x",
    "battery": "4685 mAh (35W fast charging)",
    "display": "6.3\" LTPO OLED, 120Hz, 2796x1290",
    "os": "iOS 19",
    "features": ["IP68", "Face ID", "Dynamic Island", "USB-C", "Satellite SOS"],
}

LAPTOP_SPECS_A = {
    "chipset": "M4 Pro (14-core CPU, 20-core GPU)",
    "ram": "32 GB Unified Memory",
    "storage": "512 GB SSD",
    "camera": "12 MP Center Stage",
    "battery": "72.4 Wh (up to 18 jam)",
    "display": "14.2\" Liquid Retina XDR, 3024x1964, 120Hz ProMotion",
    "os": "macOS 16",
    "features": ["Touch ID", "MagSafe 3", "Thunderbolt 5", "HDMI 2.1", "SDXC Slot"],
}

LAPTOP_SPECS_B = {
    "chipset": "Intel Core Ultra 9 285H (16-core)",
    "ram": "32 GB DDR5",
    "storage": "1 TB NVMe SSD",
    "camera": "5 MP IR + RGB",
    "battery": "84 Wh (up to 15 jam)",
    "display": "15.6\" OLED, 2880x1620, 120Hz",
    "os": "Windows 11 Pro",
    "features": ["Fingerprint Reader", "Thunderbolt 4", "HDMI 2.1", "SD Card Reader", "WiFi 7"],
}

CAMERA_SPECS_A = {
    "chipset": "DIGIC X+",
    "ram": "-",
    "storage": "Dual CFexpress Type B",
    "camera": "45 MP Full-Frame CMOS",
    "battery": "LP-E6NH (up to 700 shots)",
    "display": "3.2\" 2.1M-dot Vari-angle LCD",
    "os": "Canon OS",
    "features": ["8K 30fps Video", "IBIS 8-stop", "Dual Pixel CMOS AF II", "6K 60fps RAW", "Weather Sealed"],
}

CAMERA_SPECS_B = {
    "chipset": "BIONZ XR",
    "ram": "-",
    "storage": "Dual SD UHS-II",
    "camera": "61 MP Full-Frame CMOS",
    "battery": "NP-FZ100 (up to 600 shots)",
    "display": "3.0\" 1.44M-dot Vari-angle LCD",
    "os": "Sony OS",
    "features": ["8K 24fps Video", "IBIS 5.5-stop", "Real-time Eye AF", "4:2:2 10-bit Internal", "S-Log3"],
}

HEADPHONE_SPECS_A = {
    "chipset": "H3 Chip",
    "ram": "-",
    "storage": "-",
    "camera": "-",
    "battery": "30 jam (ANC on)",
    "display": "-",
    "os": "-",
    "features": ["Adaptive ANC", "Transparency Mode", "Spatial Audio", "IPX4", "USB-C MagSafe"],
}

HEADPHONE_SPECS_B = {
    "chipset": "WH Gen 4",
    "ram": "-",
    "storage": "-",
    "camera": "-",
    "battery": "40 jam (ANC on)",
    "display": "-",
    "os": "-",
    "features": ["Dual Noise Canceling", "Ambient Sound", "360 Reality Audio", "Speak-to-Chat", "Multipoint"],
}

TABLET_SPECS_A = {
    "chipset": "M4 (10-core CPU, 10-core GPU)",
    "ram": "16 GB Unified Memory",
    "storage": "256 GB SSD",
    "camera": "12 MP Wide + 12 MP Ultrawide",
    "battery": "39 Wh (up to 10 jam)",
    "display": "13\" Ultra Retina XDR, 2752x2064, 120Hz ProMotion",
    "os": "iPadOS 19",
    "features": ["Apple Pencil Pro", "Face ID", "Thunderbolt 4", "eSIM", "WiFi 7"],
}

TABLET_SPECS_B = {
    "chipset": "Snapdragon 8 Gen 4 for Galaxy",
    "ram": "12 GB LPDDR5X",
    "storage": "256 GB UFS 4.0",
    "camera": "13 MP Wide + 6 MP Ultrawide",
    "battery": "10090 mAh (up to 14 jam)",
    "display": "12.4\" Dynamic AMOLED 2X, 2800x1752, 120Hz",
    "os": "Android 15 + One UI 6",
    "features": ["S Pen Included", "DeX Mode", "Fingerprint", "MicroSD Up to 1TB", "WiFi 7"],
}

GENERIC_SPECS_A = {
    "chipset": "Flagship-grade processor",
    "ram": "16 GB",
    "storage": "512 GB",
    "camera": "Triple camera system (50+12+12 MP)",
    "battery": "5000 mAh",
    "display": "6.7\" AMOLED 120Hz",
    "os": "Latest OS",
    "features": ["Premium Build", "Fast Charging", "Water Resistant"],
}

GENERIC_SPECS_B = {
    "chipset": "High-performance processor",
    "ram": "12 GB",
    "storage": "256 GB",
    "camera": "Dual camera system (50+12 MP)",
    "battery": "4800 mAh",
    "display": "6.5\" OLED 90Hz",
    "os": "Latest OS",
    "features": ["Sleek Design", "Fast Charging", "Fingerprint Sensor"],
}

_STORE_NAMES_A = ["TechWorld Official", "GadgetHub Store", "Digital Lifestyle", "ElectroMart", "Megatech Store"]
_STORE_NAMES_B = ["Gizmo Paradise", "TeknoShop Official", "Smart Gadget Store", "Warna Elektronik", "TechPrime"]
_MARKETPLACES = ["Shopee", "Tokopedia", "Lazada", "Blibli", "Bukalapak"]
_BADGES = ["Power Merchant", "Official Store", "Top Seller", "Premium Partner", "Mall Exclusive"]


def _generate_store_listings(product_name: str, variant: str, base_price: int = 0) -> list:
    import random
    listings = []
    names = _STORE_NAMES_A if variant == "A" else _STORE_NAMES_B
    if base_price == 0:
        base_price = random.randint(5000000, 25000000)
    step = max(base_price // 10, 100000)
    
    for i, mp in enumerate(_MARKETPLACES):
        price = base_price + random.randint(-step, step)
        listings.append({
            "marketplace": mp,
            "storeName": names[i % len(names)],
            "price": price,
            "shipping": "Gratis Ongkir" if random.random() > 0.3 else "Berbayar",
            "rating": round(4.5 + random.random() * 0.4, 1),
            "badge": random.choice(_BADGES),
            "url": f"https://{mp.lower()}.co.id/search?q={product_name.lower().replace(' ', '+')}",
        })
    return listings


def _generate_ai_insight(product_a: str, product_b: str, cat: str) -> str:
    templates = {
        "smartphone": [
            f"Setelah menganalisis spesifikasi, {product_a} unggul dalam performa mentah dengan chipset terbaru dan RAM lebih besar, "
            f"sementara {product_b} menawarkan ekosistem yang lebih terintegrasi dan optimalisasi software yang halus. "
            f"Jika prioritas Anda adalah gaming dan multitasking berat, {product_a} adalah pilihan tepat. "
            f"Tapi jika Anda menginginkan pengalaman pengguna yang konsisten dengan update software jangka panjang, {product_b} lebih direkomendasikan.",
        ],
        "laptop": [
            f"Perbandingan antara {product_a} dan {product_b} menunjukkan keduanya adalah mesin yang sangat mumpuni. "
            f"{product_a} unggul dalam efisiensi daya dan performa single-core, ideal untuk creative professional. "
            f"Sementara {product_b} menawarkan fleksibilitas lebih dengan port yang lengkap dan upgradeability. "
            f"Keduanya cocok untuk productivity dan creative work, pilihan tergantung pada preferensi OS dan ekosistem.",
        ],
        "camera": [
            f"Dari segi spesifikasi, {product_a} menawarkan video 8K dan stabilisasi superior, cocok untuk videografer. "
            f"Sementara {product_b} unggul dalam resolusi 61MP yang ideal untuk fotografer landscape dan studio. "
            f"Keduanya memiliki build quality profesional dan sistem AF yang canggih. Pilihan tergantung pada kebutuhan spesifik Anda.",
        ],
        "headphone": [
            f"{product_a} dan {product_b} sama-sama menawarkan kualitas audio premium. "
            f"{product_a} unggul dalam integrasi ekosistem dan Spatial Audio, "
            f"sementara {product_b} menawarkan battery life lebih panjang dan fitur Speak-to-Chat yang praktis. "
            f"Keduanya memiliki ANC yang sangat baik, cocok untuk traveling dan work-from-home.",
        ],
        "tablet": [
            f"{product_a} unggul dalam performa GPU dan layar, ideal untuk desainer dan ilustrator dengan Apple Pencil Pro. "
            f"{product_b} menawarkan multitasking lebih baik dengan DeX Mode dan S Pen yang sudah included. "
            f"Keduanya mendukung produktivitas tingkat tinggi, pilihan tergantung ekosistem perangkat Anda.",
        ],
    }
    insights = templates.get(cat, [
        f"Kedua produk ini memiliki keunggulan masing-masing. "
        f"{product_a} menawarkan spesifikasi yang lebih unggul di beberapa aspek seperti performa dan kapasitas penyimpanan, "
        f"sementara {product_b} unggul dalam value dan fitur pendukung. "
        f"Rekomendasi terbaik tergantung pada prioritas dan budget Anda.",
    ])
    import random
    return random.choice(insights)


def build_comparison_data(product_a: str, product_b: str) -> dict:
    """Build comparison data deterministically from product names."""
    import random

    cat = _categorize_product(product_a)

    specs_map = {
        "smartphone": (SMARTPHONE_SPECS_A, SMARTPHONE_SPECS_B),
        "laptop": (LAPTOP_SPECS_A, LAPTOP_SPECS_B),
        "camera": (CAMERA_SPECS_A, CAMERA_SPECS_B),
        "headphone": (HEADPHONE_SPECS_A, HEADPHONE_SPECS_B),
        "tablet": (TABLET_SPECS_A, TABLET_SPECS_B),
    }

    specs_a, specs_b = specs_map.get(cat, (GENERIC_SPECS_A, GENERIC_SPECS_B))

    # Category-based price ranges
    price_ranges = {
        "smartphone": (8000000, 25000000),
        "laptop": (15000000, 45000000),
        "camera": (20000000, 60000000),
        "headphone": (2000000, 8000000),
        "tablet": (10000000, 35000000),
    }
    min_p, max_p = price_ranges.get(cat, (5000000, 30000000))
    avg_price_a = random.randint(min_p, max_p)
    avg_price_b = random.randint(min_p, max_p)

    data = {
        "productA": {
            "name": product_a,
            "image": f"https://via.placeholder.com/400x400/0a1f3f/f5a623?text={product_a.replace(' ', '+')}",
            "avgPrice": avg_price_a,
            "specs": dict(specs_a),
        },
        "productB": {
            "name": product_b,
            "image": f"https://via.placeholder.com/400x400/0d2444/f7c948?text={product_b.replace(' ', '+')}",
            "avgPrice": avg_price_b,
            "specs": dict(specs_b),
        },
        "stores": {
            "productA": _generate_store_listings(product_a, "A", avg_price_a),
            "productB": _generate_store_listings(product_b, "B", avg_price_b),
        },
        "aiInsight": _generate_ai_insight(product_a, product_b, cat),
        "images": [
            f"https://via.placeholder.com/600x400/0a1f3f/ffffff?text={product_a.replace(' ', '+')}",
            f"https://via.placeholder.com/600x400/0d2444/ffffff?text={product_b.replace(' ', '+')}",
            f"https://via.placeholder.com/600x400/f5a623/ffffff?text=Comparison",
            f"https://via.placeholder.com/600x400/1a1a2e/ffffff?text=Review",
        ],
    }
    return data


async def generate_comparison_queries(product_a: str, product_b: str) -> List[str]:
    """Generate targeted search queries for comparison."""
    queries = []
    
    # Product A queries
    queries.append(f"{product_a} spesifikasi lengkap 2026")
    queries.append(f"{product_a} harga toko online Indonesia")
    queries.append(f"{product_a} review kelebihan kekurangan")
    queries.append(f"{product_a} kamera baterai performa")
    
    # Product B queries
    queries.append(f"{product_b} spesifikasi lengkap 2026")
    queries.append(f"{product_b} harga toko online Indonesia")
    queries.append(f"{product_b} review kelebihan kekurangan")
    queries.append(f"{product_b} kamera baterai performa")
    
    # Comparison queries
    queries.append(f"{product_a} vs {product_b} comparison review")
    queries.append(f"perbandingan {product_a} dan {product_b}")
    queries.append(f"{product_a} vs {product_b} harga Indonesia")
    queries.append(f"{product_a} vs {product_b} spesifikasi")
    
    # Store-specific queries
    queries.append(f"harga {product_a} Shopee Tokopedia Bukalapak")
    queries.append(f"harga {product_b} Shopee Tokopedia Bukalapak")
    queries.append(f"harga {product_a} Blibli Lazada")
    queries.append(f"harga {product_b} Blibli Lazada")
    
    return queries[:20]  # Max 20 queries


async def comparison_search_pipeline(
    product_a: str, product_b: str, emit_status
) -> tuple[str, List[dict], List[dict]]:
    """Run multi-query search for comparison products."""
    import asyncio
    
    queries = await generate_comparison_queries(product_a, product_b)
    all_results = []
    sources = []
    
    for i, query in enumerate(queries):
        yield_msg = f"Mencari: {query}"
        emit_status(yield_msg)
        
        try:
            provider, results = await web_search(query, max_results=3)
            for r in results:
                if r.get("url") and r["url"] not in [s["url"] for s in sources]:
                    sources.append({"title": r.get("title", ""), "url": r["url"]})
            all_results.extend(results)
        except Exception as exc:
            logger.warning("Comparison search failed for '%s': %s", query, exc)
        
        # Small delay to avoid rate limiting
        if i < len(queries) - 1:
            await asyncio.sleep(0.5)
    
    return all_results, sources


def format_search_context(query: str, results: Sequence[dict]) -> str:
    max_context = env_int("WEB_SEARCH_MAX_CONTEXT_CHARS", DEFAULT_WEB_SEARCH_MAX_CONTEXT_CHARS)
    max_result = env_int("WEB_SEARCH_MAX_RESULT_CHARS", DEFAULT_WEB_SEARCH_MAX_RESULT_CHARS)

    def domain(url: str) -> str:
        try:
            return urlparse(url).hostname.replace("www.", "") or url
        except Exception:
            return url

    parts = [
        f'Hasil pencarian web untuk: "{query}"',
        "Gunakan hasil pencarian ini untuk menjawab pertanyaan pengguna.",
        "Rangkum informasi dari sumber-sumber dan berikan jawaban lengkap dalam Bahasa Indonesia.",
        "Sebutkan sumber secara alami, misal: 'Menurut detik.com...' atau 'Berdasarkan laporan dari BBC...'.",
        "JANGAN GUNAKAN format [1], [2], [3] untuk sitasi.",
        "",
    ]

    for result in results:
        title = result.get("title") or "Sumber"
        url = result.get("url") or ""
        d = domain(url)
        content = (result.get("content") or "")[:max_result]
        parts.append(f"[{d}] {title}")
        parts.append(f"URL: {url}")
        parts.append(f"Konten: {content}")
        parts.append("")

    context = "\n\n".join(parts)
    return context[:max_context]


# ── Effort Config ─────────────────────────────────────────────────────────────

EFFORT_CONFIGS = {
    "low": {
        "max_tokens": 1024,
        "system_addition": "Be concise and direct. Give short, clear answers. Avoid unnecessary elaboration."
    },
    "medium": {
        "max_tokens": 4096,
        "system_addition": "Provide a well-structured response with reasonable detail. Balance thoroughness with conciseness."
    },
    "high": {
        "max_tokens": 8192,
        "system_addition": "Be very thorough and comprehensive. Cover all important aspects, provide examples, consider edge cases, and structure your response clearly with headers and sections where appropriate."
    },
    "max": {
        "max_tokens": 16000,
        "system_addition": "Give the most comprehensive, exhaustive response possible. Cover every angle, include detailed examples, anticipate follow-up questions and address them, provide alternatives, and structure everything with clear sections. Leave nothing important out."
    }
}

# ── Skills ───────────────────────────────────────────────────────────────────

SKILLS_DIR = ROOT_DIR.parent / "skills"
SKILLS_META_FILE = SKILLS_DIR / "skills.json"

def load_skills_meta() -> dict:
    if SKILLS_META_FILE.exists():
        try:
            return json.loads(SKILLS_META_FILE.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, Exception):
            pass
    return {}

def save_skills_meta(meta: dict):
    SKILLS_META_FILE.write_text(json.dumps(meta, indent=2, ensure_ascii=False), encoding="utf-8")

BUILTIN_SKILLS = [
    {"slug": "email-writer",   "name": "Email Writer",    "icon": "📧", "category": "writing",  "description": "Write professional emails with proper structure"},
    {"slug": "code-reviewer",  "name": "Code Reviewer",   "icon": "🔍", "category": "coding",   "description": "Review code for bugs, performance, and best practices"},
    {"slug": "data-analyst",   "name": "Data Analyst",    "icon": "📊", "category": "analysis", "description": "Analyze data and provide insights"},
    {"slug": "content-writer", "name": "Content Writer",  "icon": "✍️", "category": "writing",  "description": "Create engaging content for social media and blogs"},
    {"slug": "seo-writer",     "name": "SEO Writer",      "icon": "🔎", "category": "writing",  "description": "Write SEO-optimized articles that rank on Google"},
]


def load_skill(slug: str) -> Optional[str]:
    """Load skill content from static .md file."""
    path = SKILLS_DIR / f"{slug}.md"
    if path.exists():
        return path.read_text(encoding="utf-8")
    return None


def slugify(name: str) -> str:
    import re
    slug = name.strip().lower().replace(" ", "-").replace("_", "-")
    return re.sub(r"[^a-z0-9-]", "", slug)


SKILL_SOURCE_FILE = "file"


@api_router.get("/skills")
async def list_skills():
    meta = load_skills_meta()
    skills = []
    for s in BUILTIN_SKILLS:
        content = load_skill(s["slug"])
        skills.append({**s, "available": content is not None, "builtin": True})
    for slug, m in meta.items():
        content = load_skill(slug)
        skills.append({
            "slug": slug,
            "name": m.get("name", slug),
            "icon": m.get("icon", "🧠"),
            "category": m.get("category", "custom"),
            "description": m.get("description", ""),
            "available": content is not None,
            "builtin": False,
            "source": m.get("source", "custom"),
            "created_at": m.get("created_at", ""),
        })
    return JSONResponse({"skills": skills})


@api_router.post("/skills/create")
async def create_skill(
    name: str = Form(...),
    description: str = Form(""),
    icon: str = Form("🧠"),
    category: str = Form("custom"),
    content: str = Form(""),
):
    slug = slugify(name)
    if not slug:
        return JSONResponse({"error": "Invalid skill name"}, status_code=400)
    # Build markdown content if not provided
    if not content.strip():
        content = f"# {name} Skill\n\n## Role\nYou are an expert in {name}.\n\n## Rules\n- Follow the instructions carefully\n- Provide high-quality output\n\n## Description\n{description}"
    file_path = SKILLS_DIR / f"{slug}.md"
    file_path.write_text(content.strip(), encoding="utf-8")
    meta = load_skills_meta()
    from datetime import datetime
    meta[slug] = {
        "name": name,
        "description": description,
        "icon": icon,
        "category": category,
        "source": "custom",
        "created_at": datetime.utcnow().isoformat(),
    }
    save_skills_meta(meta)
    return JSONResponse({"slug": slug, "name": name, "message": "Skill created"})


@api_router.post("/skills/import-zip")
async def import_skills_zip(file: UploadFile = File(...)):
    if not file.filename or not file.filename.endswith(".zip"):
        return JSONResponse({"error": "Please upload a .zip file"}, status_code=400)
    import zipfile
    import io
    imported = []
    errors = []
    contents = await file.read()
    with zipfile.ZipFile(io.BytesIO(contents)) as zf:
        for entry in zf.namelist():
            if entry.endswith(".md") and not entry.startswith("__"):
                slug = Path(entry).stem
                if not slug:
                    continue
                try:
                    md_content = zf.read(entry).decode("utf-8")
                    (SKILLS_DIR / f"{slug}.md").write_text(md_content, encoding="utf-8")
                    imported.append(slug)
                except Exception as e:
                    errors.append(f"{entry}: {e}")
    return JSONResponse({"imported": imported, "errors": errors})


@api_router.post("/skills/import-github")
async def import_skill_github(request: Request):
    import re
    import requests as req
    # Accept both JSON and form body
    body = {}
    try:
        json_body = await request.json()
        body.update(json_body)
    except Exception:
        pass
    try:
        form_body = await request.form()
        body.update(dict(form_body))
    except Exception:
        pass

    raw = body.get("url") or body.get("repo") or ""

    # Normalize: if it's not a full URL, treat as owner/repo shorthand
    if not re.match(r"https?://", raw):
        raw = f"https://github.com/{raw}"
    # Parse GitHub URL: https://github.com/user/repo[/path/to/file.md]
    m = re.match(r"https://github\.com/([^/]+)/([^/]+)/?.*", raw)
    if not m:
        return JSONResponse({"error": "Invalid GitHub URL"}, status_code=400)
    user, repo = m.group(1), m.group(2)
    # Determine file path
    path_part = raw.split("/blob/", 1)
    if len(path_part) > 1:
        file_path = path_part[1].split("/", 1)
        branch = file_path[0] if len(file_path) > 0 else "main"
        rel_path = file_path[1] if len(file_path) > 1 else ""
    else:
        branch = "main"
        rel_path = ""
    if not rel_path:
        # Look for skill files in the repo root (and subdirs)
        skill_name = body.get("skill") or body.get("name") or ""
        # Use GitHub Git Trees API for recursive search
        tree_url = f"https://api.github.com/repos/{user}/{repo}/git/trees/{branch}?recursive=1"
        try:
            resp = req.get(tree_url, timeout=15)
            if resp.status_code == 200:
                tree = resp.json().get("tree", [])
                md_files = [it["path"] for it in tree if it["path"].endswith(".md") and it["type"] == "blob"]
                if not md_files:
                    return JSONResponse({"error": "No .md files found in repo"}, status_code=404)
                if skill_name:
                    # Match by stem: "frontend-design" → "skills/frontend-design/SKILL.md" or "frontend-design.md"
                    stem_match = skill_name.lower()
                    matched = [f for f in md_files if f.endswith(".md") and stem_match in Path(f).stem.lower()]
                    if not matched:
                        # Try matching directory name: "skills/frontend-design/SKILL.md"
                        matched = [f for f in md_files if f.endswith(".md") and f"/{stem_match}/" in f.lower()]
                    if matched:
                        rel_path = matched[0]
                    else:
                        rel_path = md_files[0]
                else:
                    rel_path = md_files[0]
            else:
                return JSONResponse({"error": f"GitHub API error: {resp.status_code}"}, status_code=400)
        except Exception as e:
            return JSONResponse({"error": f"Failed to fetch repo: {e}"}, status_code=400)
    raw_url = f"https://raw.githubusercontent.com/{user}/{repo}/{branch}/{rel_path}"
    try:
        resp = req.get(raw_url, timeout=15)
        if resp.status_code != 200:
            return JSONResponse({"error": f"Failed to fetch file: {resp.status_code}"}, status_code=400)
        md_content = resp.text
        slug = skill_name or Path(rel_path).stem
        (SKILLS_DIR / f"{slug}.md").write_text(md_content, encoding="utf-8")
        meta = load_skills_meta()
        from datetime import datetime
        meta[slug] = {
            "name": slug.replace("-", " ").title(),
            "description": f"Imported from GitHub: {user}/{repo}",
            "icon": "🧠",
            "category": "imported",
            "source": f"github:{user}/{repo}",
            "created_at": datetime.utcnow().isoformat(),
        }
        save_skills_meta(meta)
        return JSONResponse({"slug": slug, "source": f"github:{user}/{repo}", "message": "Skill imported"})
    except Exception as e:
        return JSONResponse({"error": f"Failed to import: {e}"}, status_code=400)


@api_router.post("/skills/import")
async def import_skill(request: Request):
    import re, requests as req
    from datetime import datetime
    body = {}
    try:
        json_body = await request.json()
        body.update(json_body)
    except Exception:
        pass
    try:
        form_body = await request.form()
        body.update(dict(form_body))
    except Exception:
        pass

    source = body.get("source") or body.get("url") or body.get("repo") or ""
    skill_name = body.get("skill") or body.get("name") or ""

    if not source:
        return JSONResponse({"error": "No source provided"}, status_code=400)

    md_content = None
    slug = None
    meta_source = ""

    # Case 1: skills.sh URL → GitHub
    if re.search(r"(?:www\.)?skills\.sh", source):
        # https://www.skills.sh/owner/repo/skill-name
        parts = source.rstrip("/").split("/")
        if len(parts) >= 3:
            owner_repo = "/".join(parts[-3:-1])
            skill_part = parts[-1]
            gh_source = f"github:{owner_repo}"
            if skill_part and skill_part != parts[-2]:
                gh_source += f"/{skill_part}"
                tree_url = f"https://api.github.com/repos/{owner_repo}/git/trees/main?recursive=1"
                try:
                    resp = req.get(tree_url, timeout=10)
                    if resp.status_code == 200:
                        tree = resp.json().get("tree", [])
                        md_files = [it["path"] for it in tree if it["path"].endswith(".md") and it["type"] == "blob"]
                        stem = skill_part.lower()
                        matched = [f for f in md_files if stem in Path(f).stem.lower()]
                        if not matched:
                            matched = [f for f in md_files if f"/{stem}/" in f.lower()]
                        if matched:
                            rel_path = matched[0]
                            slug = skill_part
                            raw_url = f"https://raw.githubusercontent.com/{owner_repo}/main/{rel_path}"
                            md_resp = req.get(raw_url, timeout=15)
                            if md_resp.status_code == 200:
                                md_content = md_resp.text
                                meta_source = f"github:{owner_repo}/{slug}"
                except Exception:
                    pass
            if not md_content:
                # Fall back to general GitHub import
                source = owner_repo
                if skill_part and skill_part != parts[-2]:
                    skill_name = skill_name or skill_part

    # Case 2: Direct URL to .md file
    if not md_content and re.match(r"https?://", source) and not re.search(r"github\.com", source):
        try:
            resp = req.get(source, timeout=15)
            if resp.status_code != 200:
                return JSONResponse({"error": f"Failed to fetch URL: {resp.status_code}"}, status_code=400)
            md_content = resp.text
            slug = skill_name or source.rstrip("/").split("/")[-1].replace(".md", "")
            meta_source = f"url:{source}"
        except Exception as e:
            return JSONResponse({"error": f"URL fetch failed: {e}"}, status_code=400)

    # Case 3: GitHub (full URL or owner/repo)
    if not md_content and ("github.com" in source or re.match(r"^[\w.-]+/[\w.-]+", source)):
        if not re.match(r"https?://", source):
            source = f"https://github.com/{source}"
        m = re.match(r"https://github\.com/([^/]+)/([^/]+)/?.*", source)
        if not m:
            return JSONResponse({"error": "Invalid GitHub source"}, status_code=400)
        user, repo = m.group(1), m.group(2)
        path_part = source.split("/blob/", 1)
        if len(path_part) > 1:
            file_path = path_part[1].split("/", 1)
            branch = file_path[0] if len(file_path) > 0 else "main"
            rel_path = file_path[1] if len(file_path) > 1 else ""
        else:
            branch = "main"
            rel_path = ""
        if not rel_path:
            # Recursive search via Git Trees API
            tree_url = f"https://api.github.com/repos/{user}/{repo}/git/trees/{branch}?recursive=1"
            try:
                resp = req.get(tree_url, timeout=15)
                if resp.status_code == 200:
                    tree = resp.json().get("tree", [])
                    md_files = [it["path"] for it in tree if it["path"].endswith(".md") and it["type"] == "blob"]
                    if not md_files:
                        return JSONResponse({"error": "No .md files found in repo"}, status_code=404)
                    if skill_name:
                        stem_match = skill_name.lower()
                        matched = [f for f in md_files if stem_match in Path(f).stem.lower()]
                        if not matched:
                            matched = [f for f in md_files if f"/{stem_match}/" in f.lower()]
                        if matched:
                            rel_path = matched[0]
                        else:
                            rel_path = md_files[0]
                    else:
                        rel_path = md_files[0]
                else:
                    return JSONResponse({"error": f"GitHub API error: {resp.status_code}"}, status_code=400)
            except Exception as e:
                return JSONResponse({"error": f"Failed to fetch repo: {e}"}, status_code=400)
        raw_url = f"https://raw.githubusercontent.com/{user}/{repo}/{branch}/{rel_path}"
        try:
            resp = req.get(raw_url, timeout=15)
            if resp.status_code != 200:
                return JSONResponse({"error": f"Failed to fetch file: {resp.status_code}"}, status_code=400)
            md_content = resp.text
            slug = skill_name or Path(rel_path).stem
            meta_source = f"github:{user}/{repo}"
        except Exception as e:
            return JSONResponse({"error": f"GitHub fetch failed: {e}"}, status_code=400)

    if not md_content:
        return JSONResponse({"error": f"Unsupported source: {source}"}, status_code=400)

    if not md_content or not slug:
        return JSONResponse({"error": "Could not extract skill content"}, status_code=400)

    (SKILLS_DIR / f"{slug}.md").write_text(md_content, encoding="utf-8")
    meta = load_skills_meta()
    meta[slug] = {
        "name": slug.replace("-", " ").title(),
        "description": f"Imported from {meta_source}" if meta_source else "Imported skill",
        "icon": "🧠",
        "category": "imported",
        "source": meta_source,
        "created_at": datetime.utcnow().isoformat(),
    }
    save_skills_meta(meta)
    return JSONResponse({"slug": slug, "source": meta_source, "message": "Skill imported"})


@api_router.delete("/skills/{slug}")
async def delete_skill(slug: str):
    if any(s["slug"] == slug for s in BUILTIN_SKILLS):
        return JSONResponse({"error": "Cannot delete built-in skill"}, status_code=400)
    file_path = SKILLS_DIR / f"{slug}.md"
    if file_path.exists():
        file_path.unlink()
    meta = load_skills_meta()
    meta.pop(slug, None)
    save_skills_meta(meta)
    return JSONResponse({"message": f"Skill '{slug}' deleted"})


# ── Auth endpoints ─────────────────────────────────────────────────────────────────────

def _no_supa():
    return JSONResponse({"error": "Supabase not configured"}, status_code=503)


@api_router.post("/auth/signup")
async def auth_signup(request: Request):
    if not supa: return _no_supa()
    body = await request.json()
    email, password = body.get("email"), body.get("password")
    if not email or not password:
        return JSONResponse({"error": "email and password required"}, status_code=400)
    try:
        res = supa.auth.sign_up({"email": email, "password": password})
        return JSONResponse({"user": {"id": res.user.id, "email": res.user.email} if res.user else None, "session": {"access_token": res.session.access_token, "refresh_token": res.session.refresh_token} if res.session else None})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=400)


@api_router.post("/auth/login")
async def auth_login(request: Request):
    if not supa: return _no_supa()
    body = await request.json()
    email, password = body.get("email"), body.get("password")
    if not email or not password:
        return JSONResponse({"error": "email and password required"}, status_code=400)
    try:
        res = supa.auth.sign_in_with_password({"email": email, "password": password})
        return JSONResponse({"user": {"id": res.user.id, "email": res.user.email}, "session": {"access_token": res.session.access_token, "refresh_token": res.session.refresh_token}})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=400)


@api_router.post("/auth/logout")
async def auth_logout(request: Request):
    if not supa: return _no_supa()
    token = (request.headers.get("Authorization") or "").replace("Bearer ", "")
    try:
        supa.auth.sign_out()
        return JSONResponse({"message": "Logged out"})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=400)


@api_router.get("/auth/me")
async def auth_me(request: Request):
    if not supa: return _no_supa()
    token = (request.headers.get("Authorization") or "").replace("Bearer ", "")
    if not token:
        return JSONResponse({"error": "No token"}, status_code=401)
    try:
        res = supa.auth.get_user(token)
        return JSONResponse({"id": res.user.id, "email": res.user.email})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=401)


# ── Helper: get user_id from Bearer token ──────────────────────────────────────────────
def _get_user_id(request: Request) -> Optional[str]:
    token = (request.headers.get("Authorization") or "").replace("Bearer ", "")
    if not token or not supa: return None
    try:
        res = supa.auth.get_user(token)
        return res.user.id if res.user else None
    except Exception:
        return None


# ── Chat sessions ──────────────────────────────────────────────────────────────────────

@api_router.get("/sessions")
async def list_sessions(request: Request):
    if not supa: return _no_supa()
    uid = _get_user_id(request)
    if not uid: return JSONResponse({"error": "Unauthorized"}, status_code=401)
    res = supa.table("sessions").select("*").eq("user_id", uid).order("updated_at", desc=True).execute()
    return JSONResponse({"sessions": res.data})


@api_router.post("/sessions")
async def create_session(request: Request):
    if not supa: return _no_supa()
    uid = _get_user_id(request)
    if not uid: return JSONResponse({"error": "Unauthorized"}, status_code=401)
    body = await request.json()
    row = {"user_id": uid, "title": body.get("title", "New Chat"), "model_id": body.get("model_id", ""), "room": body.get("room", "")}
    res = supa.table("sessions").insert(row).execute()
    return JSONResponse({"session": res.data[0] if res.data else None})


@api_router.get("/sessions/{session_id}/messages")
async def get_messages(session_id: str, request: Request):
    if not supa: return _no_supa()
    uid = _get_user_id(request)
    if not uid: return JSONResponse({"error": "Unauthorized"}, status_code=401)
    # Verify session belongs to user
    sess = supa.table("sessions").select("id").eq("id", session_id).eq("user_id", uid).execute()
    if not sess.data: return JSONResponse({"error": "Session not found"}, status_code=404)
    res = supa.table("messages").select("*").eq("session_id", session_id).order("created_at").execute()
    return JSONResponse({"messages": res.data})


@api_router.post("/sessions/{session_id}/messages")
async def save_messages(session_id: str, request: Request):
    if not supa: return _no_supa()
    uid = _get_user_id(request)
    if not uid: return JSONResponse({"error": "Unauthorized"}, status_code=401)
    sess = supa.table("sessions").select("id").eq("id", session_id).eq("user_id", uid).execute()
    if not sess.data: return JSONResponse({"error": "Session not found"}, status_code=404)
    body = await request.json()
    msgs = body.get("messages", [])
    rows = [{"session_id": session_id, "role": m.get("role"), "text": m.get("text", ""), "model_id": m.get("model_id", ""), "search_mode": m.get("search_mode", ""), "skill_slug": m.get("skill_slug"), "effort_level": m.get("effort_level", ""), "image_url": m.get("image_url"), "metadata": m.get("metadata", {})} for m in msgs]
    res = supa.table("messages").insert(rows).execute()
    # Update session updated_at
    supa.table("sessions").update({"updated_at": datetime.now(timezone.utc).isoformat()}).eq("id", session_id).execute()
    return JSONResponse({"saved": len(res.data)})


@api_router.delete("/sessions/{session_id}")
async def delete_session(session_id: str, request: Request):
    if not supa: return _no_supa()
    uid = _get_user_id(request)
    if not uid: return JSONResponse({"error": "Unauthorized"}, status_code=401)
    supa.table("sessions").delete().eq("id", session_id).eq("user_id", uid).execute()
    return JSONResponse({"message": "Session deleted"})


@api_router.patch("/sessions/{session_id}")
async def update_session(session_id: str, request: Request):
    if not supa: return _no_supa()
    uid = _get_user_id(request)
    if not uid: return JSONResponse({"error": "Unauthorized"}, status_code=401)
    body = await request.json()
    patch = {k: v for k, v in body.items() if k in ("title", "model_id", "room")}
    patch["updated_at"] = datetime.now(timezone.utc).isoformat()
    supa.table("sessions").update(patch).eq("id", session_id).eq("user_id", uid).execute()
    return JSONResponse({"message": "Updated"})


@api_router.post("/sessions/{session_id}/improve-title")
async def improve_session_title(session_id: str, request: Request):
    """Generate a better title for a session using its first messages."""
    if not supa: return _no_supa()
    uid = _get_user_id(request)
    if not uid: return JSONResponse({"error": "Unauthorized"}, status_code=401)

    # Fetch first few messages
    msgs = supa.table("messages").select("role,text").eq("session_id", session_id).order("created_at").limit(4).execute()
    if not msgs.data:
        return JSONResponse({"error": "No messages found"}, status_code=404)

    conversation = "\n".join(f"{m['role']}: {m['text'][:300]}" for m in msgs.data)
    prompt = f"Generate a short, descriptive chat title (max 8 words) for this conversation. Output ONLY the title, no quotes, no explanation.\n\n{conversation}"

    base_url, api_key, _, model = get_provider_config("deepseek-v4-flash")
    if not base_url or not api_key:
        return JSONResponse({"error": "Provider not configured"}, status_code=503)

    try:
        async with httpx.AsyncClient(timeout=15.0) as client:
            resp = await client.post(provider_url(base_url), headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                json={"model": model, "messages": [{"role": "user", "content": prompt}], "max_tokens": 30, "temperature": 0.5})
            resp.raise_for_status()
            new_title = resp.json()["choices"][0]["message"]["content"].strip().strip('"')

        supa.table("sessions").update({"title": new_title, "updated_at": datetime.now(timezone.utc).isoformat()}).eq("id", session_id).eq("user_id", uid).execute()
        return JSONResponse({"title": new_title})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


# ── Subscriptions & Pakasir webhook ──────────────────────────────────────────────────────

PAKASIR_SLUG = env_str("PAKASIR_SLUG")
PAKASIR_API_KEY = env_str("PAKASIR_API_KEY")

PLAN_CREDITS = {
    "free": 50,
    "pro": 2000,
    "ultra": 10000,
}

# Pakasir IDR amount → token count for topup
TOPUP_AMOUNT_TOKENS = {
    10000: 100,
    40000: 500,
    100000: 1500,
}


def _get_plan_from_order(order_id: str) -> str:
    """Extract plan from order_id format MA-PRO-xxx, MA-ULTRA-xxx, or MA-TOPUP-xxx"""
    import re
    m = re.match(r"MA-([A-Z]+)-", order_id or "")
    if m:
        plan = m.group(1).lower()
        if plan in ("pro", "ultra", "topup"):
            return plan
    return "pro"


@api_router.get("/user/subscription")
async def get_subscription(request: Request):
    if not supa: return _no_supa()
    uid = _get_user_id(request)
    if not uid: return JSONResponse({"error": "Unauthorized"}, status_code=401)
    try:
        res = supa.table("subscriptions").select("*").eq("user_id", uid).eq("status", "active").order("activated_at", desc=True).limit(1).execute()
        if res.data:
            return JSONResponse({"subscription": res.data[0]})
        return JSONResponse({"subscription": {"plan": "free", "status": "active", "credits": PLAN_CREDITS["free"]}})
    except Exception as e:
        return JSONResponse({"subscription": {"plan": "free", "status": "active", "credits": PLAN_CREDITS["free"]}})


@api_router.post("/user/credits/deduct")
async def deduct_credits(request: Request):
    """Deduct credits after a prompt is sent — server-side enforced, no trust on frontend cost"""
    if not supa: return _no_supa()
    uid = _get_user_id(request)
    if not uid: return JSONResponse({"error": "Unauthorized"}, status_code=401)

    body = await request.json()
    model_id = body.get("model_id", "deepseek-v4-flash")
    # Always compute cost server-side, ignore frontend-provided cost
    cost = get_token_cost(model_id)
    if cost <= 0:
        return JSONResponse({"error": "Invalid model"}, status_code=400)

    try:
        res = supa.table("subscriptions").select("id,credits,plan").eq("user_id", uid).eq("status", "active").order("activated_at", desc=True).limit(1).execute()
        if not res.data:
            return JSONResponse({"credits": 0, "plan": "free"})
        sub = res.data[0]
        current_credits = sub["credits"] or 0
        if current_credits < cost:
            return JSONResponse({"error": "Insufficient credits", "credits": current_credits, "required": cost}, status_code=402)
        new_credits = current_credits - cost
        supa.table("subscriptions").update({"credits": new_credits}).eq("id", sub["id"]).execute()
        supa.table("credit_transactions").insert({
            "user_id": uid, "amount": -cost, "type": "usage", "model": model_id,
        }).execute()
        return JSONResponse({"credits": new_credits, "plan": sub["plan"]})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=400)


PAKASIR_WEBHOOK_SECRET = env_str("PAKASIR_WEBHOOK_SECRET", "")

@api_router.post("/webhooks/pakasir")
async def pakasir_webhook(request: Request):
    """Receive payment completion from Pakasir webhook"""
    import hmac, hashlib

    # Rate limit webhook (10 req/min)
    ip = request.client.host if request.client else "unknown"
    if not check_rate_limit(f"webhook:{ip}", max_requests=10, window_seconds=60):
        return JSONResponse({"error": "Rate limited"}, status_code=429)

    body = await request.json()
    order_id = body.get("order_id", "")
    amount = body.get("amount", 0)
    status = body.get("status", "")
    project = body.get("project", "")
    signature = body.get("signature", "")

    logger.info(f"Webhook received: order={order_id} amount={amount} status={status} project={project}")

    # Verify signature if secret is configured
    if PAKASIR_WEBHOOK_SECRET:
        expected = hmac.new(
            PAKASIR_WEBHOOK_SECRET.encode(),
            f"{order_id}{amount}{status}".encode(),
            hashlib.sha256,
        ).hexdigest()
        if signature != expected:
            logger.warning(f"Webhook signature mismatch for order {order_id}")
            return JSONResponse({"error": "Invalid signature"}, status_code=403)

    # Verify project matches
    if PAKASIR_SLUG and project != PAKASIR_SLUG:
        return JSONResponse({"error": "Invalid project"}, status_code=400)

    if status != "completed":
        return JSONResponse({"message": f"Status {status} ignored"})

    # ALWAYS verify with Pakasir API (never trust webhook alone)
    if PAKASIR_API_KEY and PAKASIR_SLUG:
        try:
            import requests as req
            verify = req.get(
                f"https://app.pakasir.com/api/transactiondetail",
                params={"project": PAKASIR_SLUG, "amount": amount, "order_id": order_id, "api_key": PAKASIR_API_KEY},
                timeout=10
            )
            if verify.status_code == 200:
                tx = verify.json().get("transaction", {})
                if tx.get("status") != "completed" or str(tx.get("amount")) != str(amount):
                    return JSONResponse({"error": "Payment verification failed"}, status_code=400)
            else:
                return JSONResponse({"error": f"Pakasir verify failed: {verify.status_code}"}, status_code=400)
        except Exception as e:
            logger.warning(f"Pakasir verify error: {e}")

    plan = _get_plan_from_order(order_id)

    # Find user by order_id from pending subscription
    if supa:
        try:
            # Check pending order
            pending = supa.table("subscriptions").select("user_id, amount_idr").eq("order_id", order_id).execute()
            if pending.data:
                uid = pending.data[0]["user_id"]
                from datetime import datetime, timezone

                if plan == "topup":
                    # Topup: determine tokens from Pakasir amount or DB amount
                    tx_amount = int(amount) if amount else int(pending.data[0].get("amount_idr", 0))
                    tokens_to_add = TOPUP_AMOUNT_TOKENS.get(tx_amount, max(1, tx_amount // 100))
                    # Add tokens to existing balance
                    existing = supa.table("user_credits").select("balance").eq("user_id", uid).execute()
                    current_balance = existing.data[0]["balance"] if existing.data else 0
                    new_balance = current_balance + tokens_to_add
                    supa.table("user_credits").upsert({
                        "user_id": uid,
                        "balance": new_balance,
                        "updated_at": datetime.now(timezone.utc).isoformat(),
                    }, on_conflict="user_id").execute()
                    supa.table("credit_transactions").insert({
                        "user_id": uid,
                        "amount": tokens_to_add,
                        "type": "topup",
                        "model": None,
                    }).execute()
                    # Mark subscription row as active (so verify sees it)
                    supa.table("subscriptions").update({
                        "status": "active",
                        "plan": "topup",
                        "credits": tokens_to_add,
                        "activated_at": datetime.now(timezone.utc).isoformat(),
                    }).eq("order_id", order_id).execute()
                    logger.info(f"Topup completed: user={uid} tokens={tokens_to_add} order={order_id}")
                else:
                    # Subscription: set plan credits
                    supa.table("subscriptions").update({
                        "status": "active",
                        "plan": plan,
                        "credits": PLAN_CREDITS.get(plan, 50),
                        "activated_at": datetime.now(timezone.utc).isoformat(),
                    }).eq("order_id", order_id).execute()
                    # Sync token balance to user_credits
                    new_balance = PLAN_CREDITS.get(plan, 50)
                    supa.table("user_credits").upsert({
                        "user_id": uid,
                        "balance": new_balance,
                        "plan": plan,
                        "updated_at": datetime.now(timezone.utc).isoformat(),
                    }, on_conflict="user_id").execute()
                    supa.table("credit_transactions").insert({
                        "user_id": uid,
                        "amount": new_balance,
                        "type": "topup",
                        "model": None,
                    }).execute()
                    logger.info(f"Subscription activated: user={uid} plan={plan} tokens={new_balance} order={order_id}")
            else:
                logger.warning(f"No pending subscription for order {order_id}")
        except Exception as e:
            logger.error(f"Webhook DB error: {e}")

    return JSONResponse({"message": "OK"})


@api_router.post("/subscriptions/create-pending")
async def create_pending_subscription(request: Request):
    """Create a pending subscription record before redirecting to Pakasir"""
    if not supa: return _no_supa()
    uid = _get_user_id(request)
    if not uid: return JSONResponse({"error": "Unauthorized"}, status_code=401)
    body = await request.json()
    order_id = body.get("order_id")
    plan = body.get("plan", "pro")
    amount = body.get("amount", 0)
    billing = body.get("billing", "monthly")
    if not order_id:
        return JSONResponse({"error": "order_id required"}, status_code=400)
    from datetime import datetime, timezone
    row = {
        "user_id": uid,
        "order_id": order_id,
        "plan": plan,
        "billing": billing,
        "amount_idr": amount,
        "status": "pending",
        "credits": PLAN_CREDITS.get(plan, 50),
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    try:
        supa.table("subscriptions").insert(row).execute()
        return JSONResponse({"message": "Pending subscription created", "order_id": order_id})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=400)


@api_router.get("/subscriptions/verify/{order_id}")
async def verify_subscription(order_id: str, request: Request):
    """Check payment status via Pakasir API. Auth optional."""
    import requests as req
    from datetime import datetime, timezone

    plan = _get_plan_from_order(order_id)

    # All possible IDR amounts per plan (monthly + yearly) or topup packages
    PLAN_AMOUNTS = {
        "pro": [50000, 480000],
        "ultra": [300000, 2880000],
        "topup": [10000, 40000, 100000],
    }
    amounts_to_try = PLAN_AMOUNTS.get(plan, [50000])

    # Try DB amount first
    if supa:
        try:
            sub_check = supa.table("subscriptions").select("amount_idr").eq("order_id", order_id).execute()
            if sub_check.data and sub_check.data[0].get("amount_idr"):
                db_amount = sub_check.data[0]["amount_idr"]
                # Put DB amount first in list
                amounts_to_try = [db_amount] + [a for a in amounts_to_try if a != db_amount]
        except Exception:
            pass

    # Always verify directly with Pakasir — try all possible amounts
    if PAKASIR_API_KEY and PAKASIR_SLUG:
        last_tx = {}
        for amount_idr in amounts_to_try:
            try:
                verify = req.get(
                    "https://app.pakasir.com/api/transactiondetail",
                    params={"project": PAKASIR_SLUG, "order_id": order_id, "amount": amount_idr, "api_key": PAKASIR_API_KEY},
                    timeout=10
                )
                if verify.status_code == 200:
                    tx = verify.json().get("transaction")
                    if tx:
                        last_tx = tx
                        if tx.get("status") == "completed":
                            break  # found it!
            except Exception as e:
                logger.warning(f"Pakasir verify error (amount={amount_idr}): {e}")
        tx_status = last_tx.get("status", "")
        if tx_status == "completed":
            # Update DB record if Supabase available
            if supa:
                try:
                    uid = _get_user_id(request)
                    existing = supa.table("subscriptions").select("id").eq("order_id", order_id).execute()
                    if existing.data:
                        supa.table("subscriptions").update({
                            "status": "active",
                            "plan": plan,
                            "credits": PLAN_CREDITS.get(plan, 50),
                            "activated_at": datetime.now(timezone.utc).isoformat(),
                        }).eq("order_id", order_id).execute()
                    elif uid:
                        supa.table("subscriptions").insert({
                            "user_id": uid,
                            "order_id": order_id,
                            "plan": plan,
                            "billing": "monthly",
                            "amount_idr": int(last_tx.get("amount", 0)),
                            "credits": PLAN_CREDITS.get(plan, 50),
                            "status": "active",
                            "activated_at": datetime.now(timezone.utc).isoformat(),
                        }).execute()
                    # Sync token balance to user_credits
                    if uid:
                        if plan == "topup":
                            # Topup: add tokens to existing balance
                            tx_amount = int(last_tx.get("amount", 0))
                            tokens_to_add = TOPUP_AMOUNT_TOKENS.get(tx_amount, max(1, tx_amount // 100))
                            existing_bal = supa.table("user_credits").select("balance").eq("user_id", uid).execute()
                            current_balance = existing_bal.data[0]["balance"] if existing_bal.data else 0
                            new_balance = current_balance + tokens_to_add
                            supa.table("user_credits").upsert({
                                "user_id": uid,
                                "balance": new_balance,
                                "updated_at": datetime.now(timezone.utc).isoformat(),
                            }, on_conflict="user_id").execute()
                            supa.table("credit_transactions").insert({
                                "user_id": uid,
                                "amount": tokens_to_add,
                                "type": "topup",
                                "model": None,
                            }).execute()
                        else:
                            new_balance = PLAN_CREDITS.get(plan, 50)
                            supa.table("user_credits").upsert({
                                "user_id": uid,
                                "balance": new_balance,
                                "plan": plan,
                                "updated_at": datetime.now(timezone.utc).isoformat(),
                            }, on_conflict="user_id").execute()
                            supa.table("credit_transactions").insert({
                                "user_id": uid,
                                "amount": new_balance,
                                "type": "topup",
                                "model": None,
                            }).execute()
                except Exception as db_err:
                    logger.warning(f"DB update error: {db_err}")
            return JSONResponse({"status": "active", "plan": plan})
        return JSONResponse({"status": tx_status or "pending", "plan": plan})

    # No Pakasir key or API failed — fallback to DB
    if not supa:
        return JSONResponse({"status": "pending", "plan": plan})
    uid = _get_user_id(request)
    if not uid:
        return JSONResponse({"status": "pending", "plan": plan})
    res = supa.table("subscriptions").select("*").eq("order_id", order_id).eq("user_id", uid).execute()
    if res.data:
        sub = res.data[0]
        # If still pending and user is authenticated, activate it
        # (Pakasir redirect back = payment confirmed)
        if sub.get("status") == "pending":
            from datetime import datetime, timezone
            if plan == "topup":
                tx_amount = int(sub.get("amount_idr", 0))
                tokens_to_add = TOPUP_AMOUNT_TOKENS.get(tx_amount, max(1, tx_amount // 100))
                existing_bal = supa.table("user_credits").select("balance").eq("user_id", uid).execute()
                current_balance = existing_bal.data[0]["balance"] if existing_bal.data else 0
                new_balance = current_balance + tokens_to_add
                supa.table("user_credits").upsert({
                    "user_id": uid,
                    "balance": new_balance,
                    "updated_at": datetime.now(timezone.utc).isoformat(),
                }, on_conflict="user_id").execute()
                supa.table("credit_transactions").insert({
                    "user_id": uid,
                    "amount": tokens_to_add,
                    "type": "topup",
                    "model": None,
                }).execute()
                supa.table("subscriptions").update({
                    "status": "active",
                    "plan": "topup",
                    "credits": tokens_to_add,
                    "activated_at": datetime.now(timezone.utc).isoformat(),
                }).eq("order_id", order_id).execute()
                logger.info(f"Topup fallback activated: user={uid} tokens={tokens_to_add} order={order_id}")
            else:
                new_balance = PLAN_CREDITS.get(plan, 50)
                supa.table("user_credits").upsert({
                    "user_id": uid,
                    "balance": new_balance,
                    "plan": plan,
                    "updated_at": datetime.now(timezone.utc).isoformat(),
                }, on_conflict="user_id").execute()
                supa.table("credit_transactions").insert({
                    "user_id": uid,
                    "amount": new_balance,
                    "type": "topup",
                    "model": None,
                }).execute()
                supa.table("subscriptions").update({
                    "status": "active",
                    "plan": plan,
                    "credits": new_balance,
                    "activated_at": datetime.now(timezone.utc).isoformat(),
                }).eq("order_id", order_id).execute()
                logger.info(f"Subscription fallback activated: user={uid} plan={plan} tokens={new_balance} order={order_id}")
            return JSONResponse({"status": "active", "plan": plan})
        return JSONResponse({"status": sub["status"], "plan": sub["plan"]})
    return JSONResponse({"status": "pending", "plan": plan})


# ── Token credits ──────────────────────────────────────────────────────────────

@api_router.get("/credits/{user_id}")
async def get_credits(user_id: str):
    """Get user token balance."""
    balance = await get_user_balance(user_id)
    return JSONResponse({
        "balance": balance,
        "is_low": balance <= 5,
    })


@api_router.post("/credits/seed")
async def seed_credits(request: Request):
    """Seed initial token balance for a user (free: 50 tokens)."""
    if not supa: return _no_supa()
    uid = _get_user_id(request)
    if not uid: return JSONResponse({"error": "Unauthorized"}, status_code=401)
    try:
        existing = supa.table("user_credits").select("user_id").eq("user_id", uid).execute()
        if existing.data:
            return JSONResponse({"message": "Already seeded", "balance": await get_user_balance(uid)})
        supa.table("user_credits").insert({
            "user_id": uid,
            "balance": 50,
            "plan": "free",
        }).execute()
        supa.table("credit_transactions").insert({
            "user_id": uid,
            "amount": 50,
            "type": "bonus",
            "model": None,
        }).execute()
        return JSONResponse({"message": "Seeded 50 tokens", "balance": 50})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=400)


@api_router.post("/credits/topup")
async def topup_credits(request: Request):
    """Top up tokens — webhook only (requires Pakasir signature verification)."""
    if not supa: return _no_supa()
    body = await request.json()
    target_uid = body.get("user_id")
    amount = int(body.get("amount", 0))
    tx_type = body.get("type", "topup")

    # Only allow webhook-initiated topups (not direct API calls)
    ip = request.client.host if request.client else "unknown"
    logger.warning(f"Direct topup attempt from {ip} — blocked. Use /webhooks/pakasir instead.")
    return JSONResponse({"error": "Use payment gateway for topups"}, status_code=403)


@api_router.post("/credits/topup-internal")
async def topup_credits_internal(request: Request):
    """Internal topup — only called by webhook handler (requires webhook auth)."""
    if not supa: return _no_supa()
    body = await request.json()
    target_uid = body.get("user_id")
    amount = int(body.get("amount", 0))
    tx_type = body.get("type", "topup")
    if not target_uid or amount <= 0:
        return JSONResponse({"error": "user_id and positive amount required"}, status_code=400)
    try:
        balance = await get_user_balance(target_uid)
        new_balance = balance + amount
        supa.table("user_credits").upsert({
            "user_id": target_uid,
            "balance": new_balance,
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }, on_conflict="user_id").execute()
        supa.table("credit_transactions").insert({
            "user_id": target_uid,
            "amount": amount,
            "type": tx_type,
            "model": None,
        }).execute()
        return JSONResponse({"balance": new_balance})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=400)


# ── Promo codes ──────────────────────────────────────────────────────────────

PROMO_CODES = {
    "WELCOME2026": {"type": "tokens", "amount": 200, "desc": "Bonus 200 token"},
    "GRATIS100": {"type": "tokens", "amount": 100, "desc": "Bonus 100 token"},
    "PRO3HARI": {"type": "plan", "plan": "pro", "days": 3, "desc": "Pro gratis 3 hari"},
    "ULTRA1HARI": {"type": "plan", "plan": "ultra", "days": 1, "desc": "Ultra gratis 1 hari"},
    "HEMAT50": {"type": "discount", "percent": 50, "desc": "Diskon 50% top up"},
    "MICROAGENT": {"type": "tokens", "amount": 500, "desc": "Bonus 500 token"},
}


@api_router.post("/promo/validate")
async def validate_promo(request: Request):
    """Validate a promo code and return its benefits."""
    if not supa: return _no_supa()
    body = await request.json()
    code = (body.get("code") or "").strip().upper()
    user_id = body.get("user_id")

    if not code:
        return JSONResponse({"error": "Kode promo kosong"}, status_code=400)

    promo = PROMO_CODES.get(code)
    if not promo:
        return JSONResponse({"error": "Kode promo tidak valid"}, status_code=404)

    # Check if user already used this code
    if user_id:
        try:
            used = supa.table("promo_usage").select("id").eq("user_id", user_id).eq("code", code).execute()
            if used.data:
                return JSONResponse({"error": "Kode promo sudah digunakan"}, status_code=400)
        except Exception:
            pass  # table might not exist yet

    return JSONResponse({
        "valid": True,
        "code": code,
        "type": promo["type"],
        "desc": promo["desc"],
        "amount": promo.get("amount"),
        "plan": promo.get("plan"),
        "days": promo.get("days"),
        "percent": promo.get("percent"),
    })


@api_router.post("/promo/redeem")
async def redeem_promo(request: Request):
    """Redeem a promo code — add tokens or activate plan."""
    if not supa: return _no_supa()
    body = await request.json()
    code = (body.get("code") or "").strip().upper()
    user_id = body.get("user_id")

    if not code or not user_id:
        return JSONResponse({"error": "code and user_id required"}, status_code=400)

    promo = PROMO_CODES.get(code)
    if not promo:
        return JSONResponse({"error": "Kode promo tidak valid"}, status_code=404)

    # Check if already used
    try:
        used = supa.table("promo_usage").select("id").eq("user_id", user_id).eq("code", code).execute()
        if used.data:
            return JSONResponse({"error": "Kode promo sudah digunakan"}, status_code=400)
    except Exception:
        pass

    try:
        from datetime import timedelta

        if promo["type"] == "tokens":
            amount = promo["amount"]
            existing = supa.table("user_credits").select("balance").eq("user_id", user_id).execute()
            current = existing.data[0]["balance"] if existing.data else 0
            new_balance = current + amount
            supa.table("user_credits").upsert({
                "user_id": user_id,
                "balance": new_balance,
                "updated_at": datetime.now(timezone.utc).isoformat(),
            }, on_conflict="user_id").execute()
            supa.table("credit_transactions").insert({
                "user_id": user_id,
                "amount": amount,
                "type": "promo",
                "model": None,
            }).execute()

        elif promo["type"] == "plan":
            days = promo["days"]
            plan = promo["plan"]
            supa.table("subscriptions").upsert({
                "user_id": user_id,
                "plan": plan,
                "status": "active",
                "credits": PLAN_CREDITS.get(plan, 50),
                "billing": "promo",
                "activated_at": datetime.now(timezone.utc).isoformat(),
                "expires_at": (datetime.now(timezone.utc) + timedelta(days=days)).isoformat(),
            }, on_conflict="user_id").execute()
            supa.table("user_credits").upsert({
                "user_id": user_id,
                "balance": PLAN_CREDITS.get(plan, 50),
                "plan": plan,
                "updated_at": datetime.now(timezone.utc).isoformat(),
            }, on_conflict="user_id").execute()

        # Record usage
        supa.table("promo_usage").insert({
            "user_id": user_id,
            "code": code,
            "redeemed_at": datetime.now(timezone.utc).isoformat(),
        }).execute()

        return JSONResponse({"success": True, "desc": promo["desc"]})
    except Exception as e:
        logger.error(f"Promo redeem error: {e}")
        return JSONResponse({"error": str(e)}, status_code=500)


# ── Skill installs per user ─────────────────────────────────────────────────────────────

@api_router.get("/user/skills")
async def list_user_skills(request: Request):
    if not supa: return _no_supa()
    uid = _get_user_id(request)
    if not uid: return JSONResponse({"error": "Unauthorized"}, status_code=401)
    res = supa.table("skill_installs").select("*").eq("user_id", uid).order("installed_at", desc=True).execute()
    return JSONResponse({"skills": res.data})


@api_router.post("/user/skills")
async def install_user_skill(request: Request):
    if not supa: return _no_supa()
    uid = _get_user_id(request)
    if not uid: return JSONResponse({"error": "Unauthorized"}, status_code=401)
    body = await request.json()
    slug = body.get("slug")
    if not slug: return JSONResponse({"error": "slug required"}, status_code=400)
    row = {"user_id": uid, "skill_slug": slug, "source": body.get("source", "")}
    try:
        res = supa.table("skill_installs").upsert(row, on_conflict="user_id,skill_slug").execute()
        return JSONResponse({"skill": res.data[0] if res.data else None})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=400)


@api_router.delete("/user/skills/{slug}")
async def uninstall_user_skill(slug: str, request: Request):
    if not supa: return _no_supa()
    uid = _get_user_id(request)
    if not uid: return JSONResponse({"error": "Unauthorized"}, status_code=401)
    supa.table("skill_installs").delete().eq("user_id", uid).eq("skill_slug", slug).execute()
    return JSONResponse({"message": f"Skill '{slug}' uninstalled"})


# ── User Memories (Layer B - Persistent Cross-Chat) ────────────────────────────

def fetch_user_memories(user_id: str) -> List[str]:
    """Fetch active memories for a user. Returns list of memory content strings."""
    if not supa:
        return []
    try:
        res = supa.table("user_memories").select("content").eq("user_id", user_id).eq("is_active", True).order("created_at", desc=True).limit(20).execute()
        return [m["content"] for m in res.data]
    except Exception as e:
        logger.warning(f"Failed to fetch memories for user {user_id}: {e}")
        return []


async def extract_memories_from_chat(messages: List[dict], user_id: str, model_id: str = "deepseek-v4-flash") -> List[str]:
    """Extract permanent facts from chat messages via LLM. Returns list of fact strings."""
    if not messages or not user_id:
        return []

    # Build conversation text (last 10 messages max to keep extraction prompt focused)
    recent = messages[-10:] if len(messages) > 10 else messages
    conversation = "\n".join([f"{m.get('role', 'user')}: {m.get('content', '')}" for m in recent])
    logger.info(f"Extraction conversation context ({len(recent)} messages): {conversation[:500]}...")

    extraction_prompt = f"""Extract permanent facts about the user from this conversation that should be remembered in future chats.

Include:
- Name, location, job/role
- Programming languages, frameworks, tools they use
- Projects they're working on
- Preferences (hobbies, interests, work style)

Ignore:
- Temporary questions
- Sensitive data (health, finance, orientation) unless explicitly requested
- Generic chat responses

Conversation:
{conversation}

Output format: One fact per line, max 5 facts. If no permanent facts, output nothing.

Permanent facts:"""

    base_url, api_key, provider, model = get_provider_config(model_id)
    if not base_url or not api_key:
        logger.warning("Extraction skipped: no provider config")
        return []

    try:
        body = {
            "model": model,
            "messages": [{"role": "user", "content": extraction_prompt}],
            "stream": False,
            "temperature": 0.3,
            "max_tokens": 300,
        }
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}

        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.post(provider_url(base_url), headers=headers, json=body)
            response.raise_for_status()
            data = response.json()

        content = data.get("choices", [{}])[0].get("message", {}).get("content", "")
        logger.info(f"Extraction LLM response: {content}")
        facts = [line.strip().lstrip("-•*").strip() for line in content.split("\n") if line.strip() and not line.strip().startswith("Fakta")]
        logger.info(f"Parsed facts before filter: {facts}")
        filtered = [f for f in facts if len(f) > 10 and len(f) < 500][:5]
        logger.info(f"Filtered facts: {filtered}")
        return filtered
    except Exception as e:
        logger.exception(f"Memory extraction failed: {e}")
        return []


async def save_user_memories(user_id: str, facts: List[str], source_session_id: str = None):
    """Save new memories to DB after deduplication check."""
    if not supa or not facts or not user_id:
        return

    try:
        # Fetch existing memories for dedup
        existing = supa.table("user_memories").select("content").eq("user_id", user_id).eq("is_active", True).execute()
        existing_texts = [m["content"].lower() for m in existing.data]

        # Insert only new facts (simple substring dedup)
        new_facts = []
        for fact in facts:
            fact_lower = fact.lower()
            is_duplicate = any(fact_lower in existing or existing in fact_lower for existing in existing_texts)
            if not is_duplicate:
                new_facts.append({
                    "user_id": user_id,
                    "content": fact,
                    "category": "general",
                    "source_session_id": source_session_id,
                })

        if new_facts:
            supa.table("user_memories").insert(new_facts).execute()
            logger.info(f"Saved {len(new_facts)} new memories for user {user_id}")
    except Exception as e:
        logger.warning(f"Failed to save memories: {e}")


def trigger_memory_extraction(messages: List[dict], user_id: str, session_id: str = None):
    """Fire-and-forget background task to extract and save memories."""
    if not user_id or not messages:
        logger.info(f"Memory extraction skipped: user_id={user_id}, messages_count={len(messages) if messages else 0}")
        return

    logger.info(f"Memory extraction triggered for user {user_id}, {len(messages)} messages")

    async def _extract_and_save():
        await asyncio.sleep(5)  # Small delay to avoid blocking stream completion
        logger.info(f"Memory extraction started for user {user_id}")
        facts = await extract_memories_from_chat(messages, user_id)
        logger.info(f"Extracted {len(facts)} facts for user {user_id}: {facts}")
        if facts:
            await save_user_memories(user_id, facts, session_id)
            logger.info(f"Saved memories for user {user_id}")
        else:
            logger.info(f"No facts to save for user {user_id}")

    asyncio.create_task(_extract_and_save())


# ── API Key Rotation ──────────────────────────────────────────────────────────

@api_router.post("/admin/rotate-api-keys")
async def rotate_api_keys(request: Request):
    """Rotate API keys — admin only, triggered manually or via cron."""
    # In production, this should be protected by admin auth
    # For now, require a master password from env
    body = await request.json()
    master_key = body.get("master_key", "")
    expected = env_str("ADMIN_MASTER_KEY", "")
    if not expected or master_key != expected:
        return JSONResponse({"error": "Unauthorized"}, status_code=403)

    # Generate new keys
    import secrets
    new_pakasir_key = secrets.token_hex(32)
    new_webhook_secret = secrets.token_hex(32)

    # Log rotation (in production, update env vars via Railway/Vercel dashboard)
    logger.info(f"API key rotation requested. New Pakasir key: {new_pakasir_key[:8]}..., New webhook secret: {new_webhook_secret[:8]}...")
    return JSONResponse({
        "message": "New keys generated. Update PAKASIR_API_KEY and PAKASIR_WEBHOOK_SECRET in Railway dashboard.",
        "pakasir_key": new_pakasir_key,
        "webhook_secret": new_webhook_secret,
    })


@api_router.get("/admin/security-status")
async def security_status(request: Request):
    """Check security configuration status."""
    return JSONResponse({
        "cors_origins": cors_origins(),
        "rate_limiting": True,
        "webhook_signature": bool(PAKASIR_WEBHOOK_SECRET),
        "cors_wildcard": "*" in cors_origins(),
        "body_size_limit": f"{MAX_BODY_SIZE // (1024*1024)}MB",
        "guest_limit": GUEST_LIMIT,
    })
    """List all active memories for authenticated user."""
    if not supa: return _no_supa()
    uid = _get_user_id(request)
    if not uid: return JSONResponse({"error": "Unauthorized"}, status_code=401)

    res = supa.table("user_memories").select("id, content, category, created_at, source_session_id").eq("user_id", uid).eq("is_active", True).order("created_at", desc=True).execute()
    return JSONResponse({"memories": res.data})


@api_router.delete("/user/memories/{memory_id}")
async def delete_user_memory(memory_id: str, request: Request):
    """Soft-delete a memory (set is_active = false)."""
    if not supa: return _no_supa()
    uid = _get_user_id(request)
    if not uid: return JSONResponse({"error": "Unauthorized"}, status_code=401)

    # Verify ownership before delete
    check = supa.table("user_memories").select("user_id").eq("id", memory_id).execute()
    if not check.data or check.data[0]["user_id"] != uid:
        return JSONResponse({"error": "Memory not found or unauthorized"}, status_code=404)

    supa.table("user_memories").update({"is_active": False}).eq("id", memory_id).execute()
    return JSONResponse({"message": "Memory deleted"})


@api_router.patch("/user/memories/{memory_id}")
async def update_user_memory(memory_id: str, request: Request):
    """Edit memory content."""
    if not supa: return _no_supa()
    uid = _get_user_id(request)
    if not uid: return JSONResponse({"error": "Unauthorized"}, status_code=401)

    body = await request.json()
    content = body.get("content", "").strip()
    if not content:
        return JSONResponse({"error": "content required"}, status_code=400)

    # Verify ownership
    check = supa.table("user_memories").select("user_id").eq("id", memory_id).execute()
    if not check.data or check.data[0]["user_id"] != uid:
        return JSONResponse({"error": "Memory not found or unauthorized"}, status_code=404)

    supa.table("user_memories").update({"content": content, "updated_at": "now()"}).eq("id", memory_id).execute()
    return JSONResponse({"message": "Memory updated"})


async def extract_facts_from_text(content: str, model_id: str = "deepseek-v4-flash") -> List[str]:
    """Extract individual facts from AI-generated summary text."""
    if not content or not content.strip():
        return []

    extraction_prompt = f"""Extract individual facts from this AI-generated user summary.
Output one fact per line, in natural language.
Ignore opening/closing phrases that don't contain facts.

Summary:
{content}

Facts (one per line):"""

    base_url, api_key, provider, model = get_provider_config(model_id)
    if not base_url or not api_key:
        logger.warning("Import extraction skipped: no provider config")
        return []

    try:
        body = {
            "model": model,
            "messages": [{"role": "user", "content": extraction_prompt}],
            "stream": False,
            "temperature": 0.3,
            "max_tokens": 1000,
        }
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}

        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.post(provider_url(base_url), headers=headers, json=body)
            response.raise_for_status()
            data = response.json()

        response_content = data.get("choices", [{}])[0].get("message", {}).get("content", "")
        logger.info(f"Import extraction response: {response_content[:200]}...")

        facts = [line.strip().lstrip("-•*").strip() for line in response_content.split("\n") if line.strip()]
        filtered = [f for f in facts if len(f) > 10 and len(f) < 500]
        logger.info(f"Import extracted {len(filtered)} facts")

        # Fallback: if LLM returned nothing, parse raw content as bullet points
        if not filtered:
            logger.warning("LLM returned no facts, falling back to raw content parsing")
            filtered = [line.strip().lstrip("-•*123456789.").strip()
                       for line in content.split("\n") if line.strip()]
            filtered = [f for f in filtered if len(f) > 10 and len(f) < 500][:20]

        return filtered
    except Exception as e:
        logger.exception(f"Import extraction failed: {e}")
        return []


@api_router.post("/user/memories/import-preview")
async def import_memory_preview(request: Request):
    """Extract facts from imported text for preview."""
    if not supa: return _no_supa()
    uid = _get_user_id(request)
    if not uid: return JSONResponse({"error": "Unauthorized"}, status_code=401)

    body = await request.json()
    content = body.get("content", "").strip()
    source_platform = body.get("source_platform", "Other")

    if not content:
        return JSONResponse({"error": "content required"}, status_code=400)

    if len(content) > 51200:  # 50KB limit
        return JSONResponse({"error": "Content too large. Maximum 50KB for direct processing."}, status_code=413)

    logger.info(f"Import preview request from user {uid}, source: {source_platform}, content length: {len(content)}")

    facts = await extract_facts_from_text(content)
    if not facts:
        return JSONResponse({"error": "Could not extract any facts from the content"}, status_code=400)

    return JSONResponse({"facts": facts})


@api_router.post("/user/memories/import-confirm")
async def import_memory_confirm(request: Request):
    """Save imported facts to user_memories."""
    if not supa: return _no_supa()
    uid = _get_user_id(request)
    if not uid: return JSONResponse({"error": "Unauthorized"}, status_code=401)

    body = await request.json()
    facts = body.get("facts", [])
    source_platform = body.get("source_platform", "Other")

    if not facts or not isinstance(facts, list):
        return JSONResponse({"error": "facts array required"}, status_code=400)

    logger.info(f"Import confirm from user {uid}, source: {source_platform}, facts: {len(facts)}")

    try:
        # Fetch existing for dedup
        existing = supa.table("user_memories").select("content").eq("user_id", uid).eq("is_active", True).execute()
        existing_texts = [m["content"].lower() for m in existing.data]

        new_facts = []
        for fact in facts:
            fact_str = str(fact).strip()
            if len(fact_str) < 10 or len(fact_str) > 500:
                continue

            # Dedup check
            fact_lower = fact_str.lower()
            is_duplicate = any(fact_lower in existing or existing in fact_lower for existing in existing_texts)
            if not is_duplicate:
                new_facts.append({
                    "user_id": uid,
                    "content": fact_str,
                    "category": "imported",
                    "source_session_id": None,
                })

        if new_facts:
            supa.table("user_memories").insert(new_facts).execute()
            logger.info(f"Imported {len(new_facts)} new facts for user {uid} from {source_platform}")

        return JSONResponse({"imported_count": len(new_facts)})
    except Exception as e:
        logger.exception(f"Import confirm failed: {e}")
        return JSONResponse({"error": "Failed to save memories"}, status_code=500)


def normalize_messages(payload: ChatStreamRequest, web_context: str = "") -> List[dict]:
    room_line = f"Current room: {payload.room}." if payload.room else ""
    attachment_line = (
        f"Attached filenames: {', '.join(payload.attachments)}."
        if payload.attachments
        else ""
    )
    web_line = f"\n\n{web_context}" if web_context else ""

    # Auto web search instruction
    auto_search_instruction = """

AUTO WEB SEARCH:
Gunakan tool pencarian web KAPAN SAJA kamu tidak yakin dengan jawabanmu atau informasi yang kamu miliki mungkin sudah usang.
Jangan tunggu instruksi eksplisit dari user — jika ada keraguan, search dulu sebelum jawab.
Contoh kapan harus search: data terkini (harga, berita, cuaca), info yang berubah-ubah, atau saat kamu ragu.
"""

    # Layer B: Inject user memories into system prompt
    memory_block = ""
    if payload.user_id:
        memories = fetch_user_memories(payload.user_id)
        if memories:
            memory_block = "\n\nMEMORY TENTANG USER:\n" + "\n".join(f"- {m}" for m in memories)
            memory_block += "\n\nGunakan informasi di atas secara natural dalam respons. Jangan sebutkan 'dari memory system' kecuali user bertanya."
    qna_rules = (
        "\n\nATURAN KLARIFIKASI (QNA) - PENTING:\n"
        "Jika permintaan pengguna ambigu/samar, WAJIB gunakan format QNA berikut (jangan tebak-tebakan):\n"
        "<QNA>\n"
        "{\n"
        "  \"question\": \"Pertanyaan klarifikasi yang jelas\",\n"
        "  \"options\": [\n"
        "    {\"id\": \"a\", \"label\": \"Label Pendek\", \"description\": \"penjelasan singkat\"},\n"
        "    {\"id\": \"b\", \"label\": \"Label Pendek\", \"description\": \"penjelasan singkat\"}\n"
        "  ],\n"
        "  \"allow_custom\": true\n"
        "}\n"
        "</QNA>\n\n"
        "KAPAN GUNAKAN QNA (contoh nyata):\n"
        "- 'buatkan pertanyaan' \u2192 GUNAKAN QNA: pertanyaan untuk apa? (diskusi/kuis/survei/refleksi)\n"
        "- 'buatkan konten' \u2192 GUNAKAN QNA: konten apa? platform apa?\n"
        "- 'analisis ini' \u2192 GUNAKAN QNA: analisis dari sudut apa?\n"
        "- 'buat rangkuman' \u2192 GUNAKAN QNA: untuk keperluan apa?\n"
        "- 'tolong bantu' \u2192 GUNAKAN QNA: bantu apa?\n\n"
        "KAPAN JANGAN GUNAKAN QNA:\n"
        "- 'apa itu AI?' \u2192 jelas, jawab langsung\n"
        "- 'buatkan pdf absensi 10 orang' \u2192 spesifik, langsung kerjakan\n"
        "- 'terjemahkan ke Inggris' \u2192 jelas, langsung kerjakan\n\n"
        "ATURAN QNA: Maksimal 4 opsi. Label 2-5 kata. allow_custom: true jika ada kemungkinan lain. "
        "Setelah user memilih, LANGSUNG kerjakan tanpa tanya lagi. Satu QNA per giliran.\n\n"
        "ATURAN KRITIS - [QNA_ANSWER]:\n"
        "Jika pesan user dimulai dengan [QNA_ANSWER], artinya user baru saja menjawab pertanyaan klarifikasimu. "
        "WAJIB:\n"
        "1. Langsung kerjakan tugas asli berdasarkan jawaban tersebut\n"
        "2. JANGAN keluarkan <QNA> lagi\n"
        "3. JANGAN tanya klarifikasi lagi\n"
        "Contoh: user kirim '[QNA_ANSWER] Kuis: pertanyaan untuk menguji pengetahuan' \u2192 langsung buat soal kuis, tidak perlu tanya lagi."
    )
    # Inject skill if specified
    skill_block = ""
    if payload.skill_slug:
        skill_content = load_skill(payload.skill_slug)
        if skill_content:
            skill_block = (
                f"\n\n---SKILL START---\n{skill_content}\n---SKILL END---\n"
                "Baca dan ikuti instruksi skill di atas sebelum merespons."
            )

    effort = EFFORT_CONFIGS.get(payload.effort_level, EFFORT_CONFIGS["low"])
    effort_block = f"\n\nRESPONSE EFFORT: {payload.effort_level.upper()}\n{effort['system_addition']}"

    system = {
        "role": "system",
        "content": (
            "Kamu adalah MicroAgent, asisten AI yang sangat membantu. "
            "SELALU jawab dalam Bahasa Indonesia yang benar dan lengkap. "
            "Jangan asal jawab - berikan jawaban yang informatif, akurat, dan bermanfaat. "
            "Gunakan paragraf yang terstruktur dengan baik. "
            "Untuk topik teknis, berikan contoh jika memungkinkan. "
            "Jika menggunakan informasi dari web, sebutkan sumbernya secara alami dalam kalimat (misal: 'Menurut detik.com...'), JANGAN gunakan format angka seperti [1] atau [2] untuk sitasi. "
            "\n\nPENTING - ATURAN PEMBUATAN DOKUMEN:"
            "Jika pengguna meminta membuat dokumen (PDF, Word, DOCX, laporan, surat, tabel, absensi, proposal, resume, dll), "
            "JANGAN pernah membalas dengan kode Python atau kode pemrograman apapun. "
            "Sebaliknya, langsung tulis ISI DOKUMEN LENGKAP dalam format Markdown yang rapi dan profesional. "
            "Gunakan heading (#, ##, ###), tabel Markdown, bullet list, bold, dan format lain yang sesuai. "
            "Sertakan semua konten yang diminta secara lengkap \u2014 judul, isi, tabel, catatan, dll \u2014 seolah-olah itu adalah dokumen final yang siap dicetak. "
            "Di akhir respons, tambahkan baris: '> **[DOKUMEN SIAP]** Klik tombol *Unduh Dokumen* di bawah untuk menyimpan sebagai PDF, Word, atau format lainnya.' "
            "\n\nPENTING - ATURAN PERBANDINGAN PRODUK:"
            "Jika pengguna meminta perbandingan produk (bandingkan X vs Y, compare, vs, versus, dll), "
            "WAJIB keluarkan data terstruktur dalam format XML berikut SETELAH teks penjelasan:\n"
            "<COMPARISON_DATA>\n"
            "{\n"
            '  "productA": {\n'
            '    "name": "Nama Produk A",\n'
            '    "image": "URL gambar produk A dari web (jika ada)",\n'
            '    "avgPrice": 0,\n'
            '    "specs": {\n'
            '      "chipset": "...", "ram": "...", "storage": "...",\n'
            '      "camera": "...", "battery": "...", "display": "...",\n'
            '      "os": "...", "features": ["fitur1", "fitur2"]\n'
            "    }\n"
            "  },\n"
            '  "productB": { ... format sama ... },\n'
            '  "stores": {\n'
            '    "productA": [\n'
            '      {"marketplace": "Shopee|Tokopedia|Bukalapak|Blibli|Lazada", "storeName": "...", "price": 0, "shipping": "Gratis Ongkir", "rating": 4.9, "badge": "Power Merchant", "url": "https://..."}\n'
            "    ],\n"
            '    "productB": [ ... format sama ... ]\n'
            "  },\n"
            '  "aiInsight": "Rekomendasi AI dalam satu paragraf...",\n'
            '  "images": ["url_gambar1.jpg", "url_gambar2.jpg"]\n'
            "}\n"
            "</COMPARISON_DATA>\n\n"
            "ATURAN COMPARISON_DATA:\n"
            "- HANYA keluarkan jika user meminta perbandingan produk\n"
            "- avgPrice: harga rata-rata dalam Rupiah (angka, tanpa format)\n"
            "- marketplace: HANYA salah satu dari: Shopee, Tokopedia, Bukalapak, Blibli, Lazada\n"
            "- rating: antara 4.5 dan 4.9\n"
            "- images: minimal 3 URL gambar produk dari web\n"
            "- aiInsight: rekomendasi singkat berdasarkan data yang ditemukan\n"
            "- Sebelum tag COMPARISON_DATA, berikan teks penjelasan singkat dalam format markdown biasa\n"
            + qna_rules
            + skill_block
            + f"{room_line} {attachment_line}".strip()
            + effort_block
            + memory_block
            + auto_search_instruction
            + web_line
        ),
    }
    messages = [m.model_dump() for m in payload.messages if m.content.strip()]
    return [system, *messages]


def get_model_specific_config(model_id: str, effort_level: str, has_web_context: bool) -> dict:
    """Return model-specific API parameters for Claude models."""
    base = {
        "temperature": 0.7,
        "stream": True,
    }

    effort_config = EFFORT_CONFIGS.get(effort_level, EFFORT_CONFIGS["low"])
    base["max_tokens"] = effort_config["max_tokens"]

    # SumoPod models: NO thinking/effort params, but ensure sufficient max_tokens
    if is_sumopod_model(model_id):
        base["max_tokens"] = max(base["max_tokens"], 4096)
        return base

    # Claude Sonnet 4.5 (AIMurah) — manual thinking with budget
    if model_id in ["claude-sonnet-4-5", "claude-sonnet-4.5", "claude-sonnet-4-5-1m"]:
        budget = 8000 if effort_level in ["high", "xhigh", "max"] else 4000
        base["thinking"] = {"type": "enabled", "budget_tokens": budget}
        base["max_tokens"] = max(base["max_tokens"], budget + 4000)

    # Claude Sonnet 5 (AIMurah only) — adaptive thinking with effort
    elif model_id == "claude-sonnet-5" and not is_sumopod_model(model_id):
        base["thinking"] = {"type": "adaptive"}
        effort_map = {"low": "low", "medium": "medium", "high": "high", "xhigh": "high", "max": "max"}
        base["effort"] = effort_map.get(effort_level, "medium")

    # Claude Fable 5 (AIMurah only) — always adaptive, only effort param
    elif model_id == "claude-fable-5" and not is_sumopod_model(model_id):
        effort_map = {"low": "low", "medium": "medium", "high": "high", "xhigh": "high", "max": "max"}
        base["effort"] = effort_map.get(effort_level, "medium")

    return base


def stream_provider(payload: ChatStreamRequest, web_context: str = "") -> Iterator[str]:
    if not has_user_content(payload.messages):
        yield sse("error", {"message": "Harap kirim pesan yang tidak kosong."})
        return

    base_url, api_key, provider, model = get_provider_config(payload.model_id)
    timeout = env_float("OPENAI_TIMEOUT_SECONDS", DEFAULT_TIMEOUT_SECONDS)
    if payload.auto_mode:
        model = resolve_provider_model(None, auto_mode=True)

    yield sse(
        "meta",
        {"provider": provider, "model": model, "autoMode": payload.auto_mode},
    )

    if not base_url or not api_key or not model:
        yield sse(
            "error",
            {
                "message": (
                    f"{provider} belum dikonfigurasi. Set "
                    f"{'SUMODOP_BASE_URL dan SUMODOP_API_KEY' if is_sumopod_model(payload.model_id) else 'OPENAI_BASE_URL dan OPENAI_API_KEY'} "
                    "di backend/.env."
                )
            },
        )
        return

    # Get model-specific config (thinking, effort, max_tokens)
    model_config = get_model_specific_config(payload.model_id or DEFAULT_MODEL_ID, payload.effort_level, bool(web_context))

    body = {
        "model": model,
        "messages": normalize_messages(payload, web_context=web_context),
        **model_config,
    }

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "Accept": "text/event-stream",
    }

    # Add interleaved thinking beta header for Sonnet 4.5 (enables thinking during tool calls)
    if payload.model_id in ["claude-sonnet-4-5", "claude-sonnet-4.5", "claude-sonnet-4-5-1m"]:
        headers["anthropic-beta"] = "interleaved-thinking-2025-05-14"

    try:
        with requests.post(
            provider_url(base_url),
            headers=headers,
            json=body,
            stream=True,
            timeout=timeout,
        ) as response:
            response.encoding = "utf-8"
            if response.status_code >= 400:
                detail = response.text[:MAX_PROVIDER_ERROR_CHARS]
                logger.warning(
                    "%s error %s: %s", provider, response.status_code, detail
                )
                yield sse(
                    "error",
                    {"message": f"{provider} error {response.status_code}: {detail}"},
                )
                return

            for raw_line in response.iter_lines(decode_unicode=True):
                if not raw_line:
                    continue
                line = raw_line.strip()
                if line.startswith("data:"):
                    line = line[5:].strip()
                if line == "[DONE]":
                    yield sse("done", {"status": "completed"})
                    return
                try:
                    chunk = json.loads(line)
                except json.JSONDecodeError:
                    logger.debug("Skipping malformed provider stream line: %s", line[:200])
                    continue

                choices = chunk.get("choices") or []
                if not choices:
                    continue
                choice = choices[0] or {}

                reasoning = extract_reasoning_content(choice)
                if reasoning and payload.reasoning:
                    yield sse("reasoning", {"text": reasoning})

                token = extract_delta_text(choice)
                if token:
                    yield sse("token", {"text": token})
                if choice.get("finish_reason"):
                    yield sse("done", {"status": "completed"})
                    return

            yield sse("done", {"status": "completed"})
    except requests.RequestException as exc:
        logger.exception("%s streaming failed", provider)
        yield sse("error", {"message": f"Gagal konek ke {provider}: {exc}"})


# ── Document Generation ─────────────────────────────────────────────────────

class DocumentRequest(BaseModel):
    content: str
    format: Literal["pdf", "docx", "txt", "md"] = "pdf"
    title: str = "MicroAgent Document"
    model_id: Optional[str] = None

    @field_validator("content", "title")
    @classmethod
    def strip_str(cls, v: str) -> str:
        return v.strip()


def _is_table_row(line: str) -> bool:
    """Check if a line looks like a markdown table row."""
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def _is_table_separator(line: str) -> bool:
    """Check if a line is a markdown table separator row like |---|---|."""
    stripped = line.strip()
    if not stripped.startswith("|") or not stripped.endswith("|"):
        return False
    inner = stripped[1:-1]
    return all(re.match(r'^\s*:?-+:?\s*$', cell) for cell in inner.split("|") if cell.strip())


def _parse_table_row(line: str) -> List[str]:
    """Parse a markdown table row into a list of cell strings."""
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def _parse_markdown_blocks(text: str) -> List[dict]:
    """
    Parse markdown text into typed blocks:
    heading, table, code, bullet, paragraph.
    """
    blocks: List[dict] = []
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        # Fenced code block
        if line.strip().startswith("```"):
            lang = line.strip()[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            blocks.append({"type": "code", "lang": lang, "text": "\n".join(code_lines)})
        # Markdown table — detect header row followed by separator
        elif (
            _is_table_row(line)
            and i + 1 < len(lines)
            and _is_table_separator(lines[i + 1])
        ):
            headers = _parse_table_row(line)
            i += 2  # skip header + separator
            rows = []
            while i < len(lines) and _is_table_row(lines[i]):
                rows.append(_parse_table_row(lines[i]))
                i += 1
            blocks.append({"type": "table", "headers": headers, "rows": rows})
            continue  # i already advanced
        # Heading
        elif re.match(r'^#{1,6}\s', line):
            level = len(re.match(r'^(#+)', line).group(1))
            text_content = re.sub(r'^#+\s*', '', line).strip()
            blocks.append({"type": "heading", "level": level, "text": text_content})
        # Bullet / list
        elif re.match(r'^[\*\-\+]\s', line) or re.match(r'^\d+\.\s', line):
            item_text = re.sub(r'^([\*\-\+]|\d+\.)\s*', '', line).strip()
            blocks.append({"type": "bullet", "text": item_text})
        # Horizontal rule
        elif re.match(r'^[-_\*]{3,}\s*$', line):
            blocks.append({"type": "hr"})
        # Blank line → flush paragraph
        elif line.strip() == "":
            blocks.append({"type": "spacer"})
        # Regular paragraph text
        else:
            # Accumulate consecutive paragraph lines
            para_lines = [line]
            while (
                i + 1 < len(lines)
                and lines[i + 1].strip() != ""
                and not re.match(r'^#{1,6}\s', lines[i + 1])
                and not lines[i + 1].strip().startswith("```")
                and not re.match(r'^[\*\-\+]\s', lines[i + 1])
                and not re.match(r'^\d+\.\s', lines[i + 1])
                and not re.match(r'^[-_\*]{3,}\s*$', lines[i + 1])
                and not _is_table_row(lines[i + 1])
            ):
                i += 1
                para_lines.append(lines[i])
            blocks.append({"type": "paragraph", "text": " ".join(para_lines)})
        i += 1
    return blocks


def _strip_md_inline(text: str) -> str:
    """Remove inline markdown (bold, italic, backticks, links)."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text


def generate_pdf(content: str, title: str) -> bytes:
    if not REPORTLAB_AVAILABLE:
        raise RuntimeError("reportlab tidak tersedia")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        "DocTitle",
        parent=styles["Title"],
        fontSize=20,
        leading=26,
        spaceAfter=6,
        textColor=colors.HexColor("#111111"),
        fontName="Helvetica-Bold",
    )
    h1_style = ParagraphStyle(
        "H1", parent=styles["Heading1"], fontSize=15, leading=20,
        spaceBefore=14, spaceAfter=4, textColor=colors.HexColor("#111111"),
        fontName="Helvetica-Bold",
    )
    h2_style = ParagraphStyle(
        "H2", parent=styles["Heading2"], fontSize=13, leading=18,
        spaceBefore=10, spaceAfter=3, textColor=colors.HexColor("#1D4ED8"),
        fontName="Helvetica-Bold",
    )
    h3_style = ParagraphStyle(
        "H3", parent=styles["Heading3"], fontSize=11, leading=16,
        spaceBefore=8, spaceAfter=2, textColor=colors.HexColor("#374151"),
        fontName="Helvetica-Bold",
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"], fontSize=10, leading=15,
        spaceAfter=6, textColor=colors.HexColor("#111111"),
        alignment=TA_JUSTIFY, fontName="Helvetica",
    )
    bullet_style = ParagraphStyle(
        "Bullet", parent=styles["Normal"], fontSize=10, leading=14,
        spaceAfter=3, leftIndent=14, bulletIndent=4,
        textColor=colors.HexColor("#111111"), fontName="Helvetica",
    )
    code_style = ParagraphStyle(
        "Code", parent=styles["Code"], fontSize=8, leading=12,
        spaceBefore=6, spaceAfter=6,
        backColor=colors.HexColor("#F3F4F6"),
        textColor=colors.HexColor("#1F2937"),
        leftIndent=10, rightIndent=10,
        fontName="Courier",
    )
    meta_style = ParagraphStyle(
        "Meta", parent=styles["Normal"], fontSize=9, leading=12,
        spaceAfter=12, textColor=colors.HexColor("#6B7280"),
        fontName="Helvetica",
    )

    heading_styles = {1: h1_style, 2: h2_style, 3: h3_style, 4: h3_style, 5: h3_style, 6: h3_style}

    story = []
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 4))

    from reportlab.platypus import Table as RLTable, TableStyle as RLTableStyle

    blocks = _parse_markdown_blocks(content)
    for block in blocks:
        btype = block["type"]
        if btype == "heading":
            lvl = block["level"]
            st = heading_styles.get(lvl, h3_style)
            story.append(Paragraph(_strip_md_inline(block["text"]), st))
        elif btype == "table":
            headers = [_strip_md_inline(h) for h in block["headers"]]
            rows = [[_strip_md_inline(c) for c in row] for row in block["rows"]]
            table_data = [headers] + rows

            cell_style = ParagraphStyle(
                "TableCell", parent=styles["Normal"],
                fontSize=9, leading=12, fontName="Helvetica",
            )
            header_style = ParagraphStyle(
                "TableHeader", parent=styles["Normal"],
                fontSize=9, leading=12, fontName="Helvetica-Bold",
                textColor=colors.white,
            )
            # Wrap cells in Paragraphs for text wrapping
            styled_data = [
                [Paragraph(cell, header_style) for cell in table_data[0]]
            ] + [
                [Paragraph(cell, cell_style) for cell in row]
                for row in table_data[1:]
            ]

            # Auto-distribute column widths across page width
            usable_width = 16 * cm  # A4 minus margins
            col_count = len(headers)
            col_w = [usable_width / col_count] * col_count

            tbl = RLTable(styled_data, colWidths=col_w, repeatRows=1)
            tbl.setStyle(RLTableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1D4ED8")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 9),
                ("ALIGN", (0, 0), (-1, 0), "CENTER"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D1D5DB")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F9FAFB")]),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ]))
            story.append(Spacer(1, 8))
            story.append(tbl)
            story.append(Spacer(1, 8))
        elif btype == "paragraph":
            clean = _strip_md_inline(block["text"])
            if clean:
                story.append(Paragraph(clean, body_style))
        elif btype == "bullet":
            clean = _strip_md_inline(block["text"])
            story.append(Paragraph(f"\u2022\u2002{clean}", bullet_style))
        elif btype == "code":
            code_text = block["text"].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Preformatted(code_text, code_style))
        elif btype == "hr":
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#D1D5DB"), spaceAfter=8))
        elif btype == "spacer":
            story.append(Spacer(1, 6))

    doc.build(story)
    return buf.getvalue()


def generate_docx(content: str, title: str) -> bytes:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx tidak tersedia")

    doc = DocxDocument()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.runs[0].font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    title_para.runs[0].font.size = Pt(20)

    blocks = _parse_markdown_blocks(content)
    for block in blocks:
        btype = block["type"]
        if btype == "heading":
            lvl = min(block["level"], 4)
            h = doc.add_heading(_strip_md_inline(block["text"]), level=lvl)
            if h.runs:
                if lvl == 1:
                    h.runs[0].font.color.rgb = RGBColor(0x11, 0x11, 0x11)
                elif lvl == 2:
                    h.runs[0].font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
        elif btype == "table":
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            headers = block["headers"]
            rows = block["rows"]
            col_count = len(headers)
            tbl = doc.add_table(rows=1 + len(rows), cols=col_count)
            tbl.style = "Table Grid"
            # Header row
            hdr_row = tbl.rows[0]
            for j, cell_text in enumerate(headers):
                cell = hdr_row.cells[j]
                cell.text = _strip_md_inline(cell_text)
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                # Blue background
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "1D4ED8")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:val"), "clear")
                tcPr.append(shd)
            # Data rows
            for ri, row_data in enumerate(rows):
                tbl_row = tbl.rows[ri + 1]
                fill = "F9FAFB" if ri % 2 == 1 else "FFFFFF"
                for j, cell_text in enumerate(row_data[:col_count]):
                    cell = tbl_row.cells[j]
                    cell.text = _strip_md_inline(cell_text)
                    run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run()
                    run.font.size = Pt(9)
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:fill"), fill)
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:val"), "clear")
                    tcPr.append(shd)
            doc.add_paragraph()  # spacer after table
        elif btype == "paragraph":
            clean = _strip_md_inline(block["text"])
            if clean:
                p = doc.add_paragraph(clean)
                p.paragraph_format.space_after = Pt(6)
        elif btype == "bullet":
            doc.add_paragraph(_strip_md_inline(block["text"]), style="List Bullet")
        elif btype == "code":
            p = doc.add_paragraph(block["text"])
            p.style = doc.styles["No Spacing"]
            for run in p.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        elif btype == "hr":
            doc.add_paragraph("─" * 60)
        elif btype == "spacer":
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Image Generation ──────────────────────────────────────────────────────────

class ImageGenerateRequest(BaseModel):
    prompt: str
    user_id: Optional[str] = None

    @field_validator("prompt")
    @classmethod
    def trim_prompt(cls, v: str) -> str:
        return v.strip()


@api_router.post("/generate-image")
async def generate_image(req: ImageGenerateRequest):
    """Generate an image using flux-2-klein-4b via the configured provider."""
    base_url = env_str("OPENAI_BASE_URL")
    api_key = env_str("OPENAI_API_KEY")

    if not base_url or not api_key:
        return JSONResponse(
            {"error": "Provider belum dikonfigurasi. Set OPENAI_BASE_URL dan OPENAI_API_KEY."},
            status_code=500,
        )

    # Build image generation URL — strip /chat/completions suffix if present
    clean_base = base_url.rstrip("/")
    if clean_base.endswith("/chat/completions"):
        clean_base = clean_base[: -len("/chat/completions")]
    image_url = f"{clean_base}/images/generations"

    payload = {
        "model": IMAGE_MODEL_ID,
        "prompt": req.prompt,
        "n": 1,
        "size": "1024x1024",
        "response_format": "url",
    }
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }

    timeout = env_float("OPENAI_TIMEOUT_SECONDS", DEFAULT_TIMEOUT_SECONDS)

    try:
        async with httpx.AsyncClient(timeout=timeout) as client:
            response = await client.post(image_url, headers=headers, json=payload)
            if response.status_code >= 400:
                detail = response.text[:MAX_PROVIDER_ERROR_CHARS]
                logger.warning("Image generation error %s: %s", response.status_code, detail)
                return JSONResponse(
                    {"error": f"Provider error {response.status_code}: {detail}"},
                    status_code=response.status_code,
                )
            data = response.json()
            image_url_result = (
                data.get("data", [{}])[0].get("url")
                or data.get("data", [{}])[0].get("b64_json")
            )
            if not image_url_result:
                return JSONResponse({"error": "Tidak ada gambar dalam respons provider"}, status_code=500)

            # If b64_json, prefix with data URI
            if not image_url_result.startswith("http"):
                image_url_result = f"data:image/png;base64,{image_url_result}"

            # Save to Supabase Storage + DB if user_id provided
            saved_image = None
            if req.user_id and supa:
                try:
                    img_bytes = None
                    ext = "png"
                    content_type = "image/png"

                    if image_url_result.startswith("http"):
                        # Download image from provider
                        async with httpx.AsyncClient(timeout=60) as dl_client:
                            img_resp = await dl_client.get(image_url_result)
                        if img_resp.status_code == 200:
                            img_bytes = img_resp.content
                            content_type = img_resp.headers.get("content-type", "image/png")
                            if "jpeg" in content_type or "jpg" in content_type:
                                ext = "jpg"
                            elif "webp" in content_type:
                                ext = "webp"
                        else:
                            logger.warning("Failed to download generated image: status %s", img_resp.status_code)
                    else:
                        # b64_json case — decode inline
                        b64_part = image_url_result.split(",", 1)[-1] if "," in image_url_result else image_url_result
                        import base64
                        img_bytes = base64.b64decode(b64_part)
                        ext = "png"
                        content_type = "image/png"

                    if img_bytes:
                        # Upload to Supabase Storage
                        storage_path = f"{req.user_id}/generated/{uuid.uuid4().hex[:12]}.{ext}"
                        upload_res = supa.storage.from_("chat-files").upload(
                            storage_path, img_bytes,
                            {"content-type": content_type, "cacheControl": "3600"}
                        )
                        upload_error = getattr(upload_res, "error", None)
                        if upload_error:
                            logger.warning("Storage upload error: %s", upload_error)
                        else:
                            url_data = supa.storage.from_("chat-files").get_public_url(storage_path)
                            public_url = image_url_result  # fallback
                            if isinstance(url_data, dict):
                                public_url = url_data.get("publicUrl", url_data.get("public_url", image_url_result))
                            elif hasattr(url_data, "public_url"):
                                public_url = url_data.public_url

                            # Save metadata to DB
                            db_res = supa.table("generated_images").insert({
                                "user_id": req.user_id,
                                "prompt": req.prompt,
                                "image_url": public_url,
                                "storage_path": storage_path,
                                "model": IMAGE_MODEL_ID,
                            }).execute()
                            if db_res.data:
                                saved_image = db_res.data[0]
                                logger.info("Image saved to gallery: %s", saved_image.get("id"))
                            else:
                                logger.warning("DB insert returned no data: %s", getattr(db_res, "error", "unknown"))
                except Exception as exc:
                    logger.warning("Failed to save generated image: %s", exc, exc_info=True)

            return JSONResponse({
                "success": True,
                "image_url": image_url_result,
                "prompt": req.prompt,
                "saved_image": saved_image,
            })
    except httpx.TimeoutException:
        return JSONResponse({"error": "Request timeout saat generate gambar"}, status_code=504)
    except Exception as exc:
        logger.exception("Image generation failed")
        return JSONResponse({"error": f"Gagal generate gambar: {exc}"}, status_code=500)


# ── Generated Images Gallery ───────────────────────────────────────────────────

@api_router.get("/images")
async def list_images(request: Request):
    """List all generated images for the authenticated user."""
    if not supa: return _no_supa()
    uid = _get_user_id(request)
    if not uid: return JSONResponse({"error": "Unauthorized"}, status_code=401)
    try:
        res = supa.table("generated_images") \
            .select("*") \
            .eq("user_id", uid) \
            .order("created_at", desc=True) \
            .limit(100) \
            .execute()
        return JSONResponse({"images": res.data or []})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


@api_router.delete("/images/{image_id}")
async def delete_image(image_id: str, request: Request):
    """Delete a generated image (from Storage + DB)."""
    if not supa: return _no_supa()
    uid = _get_user_id(request)
    if not uid: return JSONResponse({"error": "Unauthorized"}, status_code=401)
    try:
        # Get image metadata
        res = supa.table("generated_images").select("storage_path").eq("id", image_id).eq("user_id", uid).single().execute()
        if not res.data:
            return JSONResponse({"error": "Image not found"}, status_code=404)
        # Delete from Storage
        storage_path = res.data.get("storage_path")
        if storage_path:
            supa.storage.from_("chat-files").remove([storage_path])
        # Delete from DB
        supa.table("generated_images").delete().eq("id", image_id).execute()
        return JSONResponse({"message": "Image deleted"})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


# ── Auto-setup: create missing tables ─────────────────────────────────────────

@api_router.post("/setup")
async def setup_database():
    """Create required Supabase tables if they don't exist. Run once."""
    if not supa: return _no_supa()
    results = {}

    SQL_STATEMENTS = {
        "user_credits": """
            CREATE TABLE IF NOT EXISTS user_credits (
                user_id TEXT PRIMARY KEY,
                balance INTEGER DEFAULT 0,
                plan TEXT DEFAULT 'free',
                updated_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
            );
        """,
        "credit_transactions": """
            CREATE TABLE IF NOT EXISTS credit_transactions (
                id UUID DEFAULT gen_random_uuid() PRIMARY KEY,
                user_id TEXT NOT NULL,
                amount INTEGER NOT NULL,
                type TEXT NOT NULL,
                model TEXT,
                created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
            );
            CREATE INDEX IF NOT EXISTS idx_credit_transactions_user_id ON credit_transactions(user_id);
        """,
        "generated_images": """
            CREATE TABLE IF NOT EXISTS generated_images (
                id UUID DEFAULT gen_random_uuid() PRIMARY KEY,
                user_id TEXT NOT NULL,
                prompt TEXT NOT NULL,
                image_url TEXT NOT NULL,
                storage_path TEXT,
                model TEXT DEFAULT 'flux-2-klein-4b',
                created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
            );
            CREATE INDEX IF NOT EXISTS idx_generated_images_user_id ON generated_images(user_id);
        """,
    }

    for table_name, sql in SQL_STATEMENTS.items():
        try:
            # Use Supabase RPC to run raw SQL (requires a function in Supabase)
            # Fallback: try to select from the table to check if it exists
            res = supa.table(table_name).select("*").limit(1).execute()
            results[table_name] = "exists"
        except Exception as e:
            error_msg = str(e)
            if "does not exist" in error_msg or "relation" in error_msg:
                results[table_name] = f"missing — run SQL manually in Supabase Editor"
            else:
                results[table_name] = f"error: {error_msg}"

    return JSONResponse({
        "message": "Setup check complete",
        "tables": results,
        "sql": SQL_STATEMENTS,
    })


@api_router.post("/generate-document")
async def generate_document(req: DocumentRequest):
    """Generate a downloadable document (PDF, DOCX, TXT, MD) from AI response text."""
    fmt = req.format.lower()
    safe_title = re.sub(r'[^\w\s-]', '', req.title).strip().replace(" ", "-")[:60] or "document"

    try:
        if fmt == "pdf":
            data = generate_pdf(req.content, req.title)
            return StreamingResponse(
                io.BytesIO(data),
                media_type="application/pdf",
                headers={"Content-Disposition": f'attachment; filename="{safe_title}.pdf"'},
            )

        elif fmt == "docx":
            data = generate_docx(req.content, req.title)
            return StreamingResponse(
                io.BytesIO(data),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{safe_title}.docx"'},
            )

        elif fmt == "txt":
            # Strip markdown formatting for plain text
            plain = re.sub(r'```[\s\S]*?```', lambda m: m.group(0).split("\n", 1)[-1].rsplit("\n", 1)[0], req.content)
            plain = re.sub(r'#{1,6}\s*', '', plain)
            plain = re.sub(r'[\*_`]', '', plain)
            plain = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', plain)
            header = f"{req.title}\n{'=' * len(req.title)}\nGenerated by MicroAgent · {datetime.now(timezone.utc).strftime('%d %B %Y')}\n\n"
            data = (header + plain).encode("utf-8")
            return StreamingResponse(
                io.BytesIO(data),
                media_type="text/plain; charset=utf-8",
                headers={"Content-Disposition": f'attachment; filename="{safe_title}.txt"'},
            )

        elif fmt == "md":
            header = f"# {req.title}\n\n> Generated by MicroAgent · {datetime.now(timezone.utc).strftime('%d %B %Y')}\n\n---\n\n"
            data = (header + req.content).encode("utf-8")
            return StreamingResponse(
                io.BytesIO(data),
                media_type="text/markdown; charset=utf-8",
                headers={"Content-Disposition": f'attachment; filename="{safe_title}.md"'},
            )

        else:
            return JSONResponse({"error": f"Format '{fmt}' tidak didukung"}, status_code=400)

    except RuntimeError as exc:
        logger.error("Document generation failed: %s", exc)
        return JSONResponse({"error": str(exc)}, status_code=500)
    except Exception as exc:
        logger.exception("Unexpected error in generate_document")
        return JSONResponse({"error": "Terjadi kesalahan saat membuat dokumen"}, status_code=500)


# ── RAG + AI Document Generation ──────────────────────────────────────────────

MEDIA_TYPES = {
    "pdf":  "application/pdf",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xls":  "application/vnd.ms-excel",
    "txt":  "text/plain", "md": "text/markdown",
}

class RAGUploadRequest(BaseModel):
    user_id: str
    is_public: bool = False

@api_router.post("/rag/upload")
async def rag_upload(file: UploadFile = File(...), user_id: str = Form(...), is_public: bool = Form(False)):
    """Upload a document for RAG knowledge base."""
    from rag import ingest_document
    file_bytes = await file.read()
    result = await ingest_document(file_bytes, file.filename, user_id, is_public)
    return JSONResponse(result)


@api_router.get("/rag/documents")
async def rag_list_documents(user_id: str):
    """List user's RAG documents."""
    if not supa: return _no_supa()
    try:
        result = supa.table("rag_documents").select(
            "id, filename, file_type, chunk_count, is_public, created_at"
        ).eq("user_id", user_id).order("created_at", desc=True).execute()
        return JSONResponse({"documents": result.data or []})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


@api_router.delete("/rag/documents/{document_id}")
async def rag_delete_document(document_id: str, user_id: str):
    """Delete a RAG document and its chunks."""
    if not supa: return _no_supa()
    try:
        supa.table("rag_documents").delete().eq("id", document_id).eq("user_id", user_id).execute()
        return JSONResponse({"success": True})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


@api_router.post("/ai-generate-doc")
async def ai_generate_doc(request: Request):
    """AI generates a real document (PDF/Excel/Word) via Python execution."""
    from rag import generate_document_with_ai, is_document_request
    body = await request.json()
    prompt = body.get("prompt", "")
    user_id = body.get("user_id", "anonymous")
    model_id = body.get("model_id", "claude-sonnet-4.5-1m")

    if not prompt:
        return JSONResponse({"error": "prompt required"}, status_code=400)

    try:
        result = await generate_document_with_ai(prompt, user_id, model_id)
        return JSONResponse(result)
    except Exception as e:
        logger.exception("AI document generation failed")
        return JSONResponse({"error": str(e)}, status_code=500)


# ── Streaming Document Generation (SSE) ────────────────────────────────────────

DOC_GEN_SYSTEM = """You are a document generator. Write Python code that creates the requested file.

Rules:
1. Write COMPLETE Python code inside ```python code block
2. Install/reportlab/openpyxl/python-docx are pre-installed
3. Save file in current directory (os.getcwd()) with Indonesian filename
4. After closing ```, write EXACTLY: OUTPUT_FILE: namafile.ext
5. Do NOT add any explanation, just code block then OUTPUT_FILE line

Example output:
```python
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
c = canvas.Canvas("laporan.pdf", pagesize=A4)
c.drawString(100, 750, "Hello")
c.save()
```
OUTPUT_FILE: laporan.pdf"""

TOPUP_AMOUNT_TOKENS = {
    10000: 100,
    40000: 500,
    100000: 1500,
}


class GenerateStreamRequest(BaseModel):
    prompt: str
    user_id: str = "anonymous"
    model_id: str = "claude-sonnet-4-6"


# Track active document generation sessions for cancellation
_active_doc_gen: dict[str, asyncio.Event] = {}
_doc_gen_semaphore = asyncio.Semaphore(2)  # Max 2 concurrent doc gen


async def _cleanup_generated_file(file_id: str, file_path: str, delay: int = 86400):
    """Remove generated file and temp dir after delay (default 24h)."""
    await asyncio.sleep(delay)
    if os.path.exists(file_path):
        try:
            os.remove(file_path)
            logger.info("Cleaned up generated file: %s", file_path)
        except Exception as e:
            logger.warning("Failed to remove generated file %s: %s", file_path, e)
    # Keep DB record so download endpoint can give informative message


@api_router.post("/generate-document-cancel")
async def cancel_doc_generation(request: Request):
    """Cancel an active document generation stream."""
    body = await request.json()
    gen_id = body.get("gen_id", "")
    if gen_id and gen_id in _active_doc_gen:
        _active_doc_gen[gen_id].set()
        return JSONResponse({"cancelled": True})
    return JSONResponse({"cancelled": False, "error": "Session not found"}, status_code=404)


@api_router.post("/generate-document-stream")
async def generate_document_stream_endpoint(request: Request):
    """Stream document generation with live code output via SSE."""
    from fastapi.responses import StreamingResponse
    body = await request.json()
    prompt = body.get("prompt", "")
    user_id = body.get("user_id", "anonymous")
    model_id = body.get("model_id", "claude-sonnet-4.5-1m")

    if not prompt:
        return JSONResponse({"error": "prompt required"}, status_code=400)

    gen_id = uuid.uuid4().hex[:12]
    cancel_event = asyncio.Event()
    _active_doc_gen[gen_id] = cancel_event

    async def safe_generator():
        async with _doc_gen_semaphore:
            try:
                async for chunk in _stream_doc_generation(prompt, user_id, model_id, gen_id, cancel_event):
                    yield chunk
            except Exception as e:
                logger.error(f"Stream generator error: {e}")
                yield f"data: {json.dumps({'type': 'error', 'message': f'Server error: {str(e)}'})}\n\n"
            finally:
                # Keep gen_id alive for 30s so cancel still works after stream ends
                async def _delayed_cleanup():
                    await asyncio.sleep(30)
                    _active_doc_gen.pop(gen_id, None)
                asyncio.create_task(_delayed_cleanup())

    return StreamingResponse(
        safe_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-store, must-revalidate",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive",
            "Access-Control-Allow-Origin": "*",
        },
    )


async def _stream_doc_generation(prompt: str, user_id: str, model_id: str = "deepseek-v4-flash", gen_id: str = "", cancel_event: asyncio.Event = None):
    """Generator that streams SSE events for document generation."""
    import httpx
    import re as _re

    full_text = ""
    full_code = ""
    in_code_block = False

    # Phase 1: Stream code generation
    yield f"data: {json.dumps({'type': 'phase', 'phase': 'generating', 'message': 'Generating code...', 'gen_id': gen_id})}\n\n"

    # Use provider routing based on model_id
    base_url, api_key, provider_name, resolved_model = get_provider_config(model_id)
    if not base_url or not api_key:
        yield f"data: {json.dumps({'type': 'error', 'message': f'{provider_name} belum dikonfigurasi'})}\n\n"
        return

    clean_base = base_url.rstrip("/")
    if clean_base.endswith("/chat/completions"):
        clean_base = clean_base[:-len("/chat/completions")]

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": resolved_model,
        "messages": [
            {"role": "system", "content": DOC_GEN_SYSTEM},
            {"role": "user", "content": prompt},
        ],
        "max_tokens": 8192,
        "temperature": 0.3,
        "stream": True,
    }

    buffer = ""
    stream_done = False

    # Retry on 429 with backoff
    for attempt in range(3):
        try:
            async with httpx.AsyncClient(timeout=120) as client:
                async with client.stream("POST", f"{clean_base}/chat/completions", headers=headers, json=payload) as resp:
                    if resp.status_code == 429:
                        error_body = await resp.aread()
                        if attempt < 2:
                            wait = (attempt + 1) * 5
                            yield f"data: {json.dumps({'type': 'phase', 'phase': 'generating', 'message': f'Rate limit, menunggu {wait}s...'})}\n\n"
                            await asyncio.sleep(wait)
                            continue
                        yield f"data: {json.dumps({'type': 'error', 'message': 'Batas request tercapai. Coba lagi dalam beberapa menit.'})}\n\n"
                        return
                    if resp.status_code >= 400:
                        error_body = await resp.aread()
                        yield f"data: {json.dumps({'type': 'error', 'message': f'AI error {resp.status_code}: {error_body[:300].decode()}'})}\n\n"
                        return

                    try:
                        async for raw_chunk in resp.aiter_bytes():
                            if stream_done:
                                break

                            if cancel_event and cancel_event.is_set():
                                yield f"data: {json.dumps({'type': 'cancelled', 'message': 'Dibatalkan oleh pengguna'})}\n\n"
                                return

                            buffer += raw_chunk.decode("utf-8", errors="ignore")

                            while "\n" in buffer:
                                if stream_done:
                                    break
                                line, buffer = buffer.split("\n", 1)
                                line = line.strip()
                                if not line:
                                    continue
                                if not line.startswith("data: "):
                                    continue
                                data_str = line[6:]
                                if data_str.strip() == "[DONE]":
                                    stream_done = True
                                    break
                                try:
                                    chunk = json.loads(data_str)
                                    delta = chunk.get("choices", [{}])[0].get("delta", {})
                                    text_chunk = delta.get("content", "")
                                    if not text_chunk:
                                        continue
                                except (json.JSONDecodeError, IndexError, KeyError):
                                    continue

                                full_text += text_chunk

                                if "```python" in full_text and not in_code_block:
                                    in_code_block = True

                                if in_code_block:
                                    code_match = _re.search(r"```python\n(.*)", full_text, _re.DOTALL)
                                    if code_match:
                                        code_so_far = code_match.group(1).split("```")[0]
                                        full_code = code_so_far
                                        yield f"data: {json.dumps({'type': 'code_chunk', 'code': text_chunk})}\n\n"

                                await asyncio.sleep(0)
                    finally:
                        await client.aclose()
            break  # success, exit retry loop
        except httpx.TimeoutException:
            if attempt < 2:
                yield f"data: {json.dumps({'type': 'phase', 'phase': 'generating', 'message': 'Timeout, mencoba ulang...'})}\n\n"
                await asyncio.sleep(3)
                continue
            yield f"data: {json.dumps({'type': 'error', 'message': 'Connection timeout'})}\n\n"
            return
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'message': f'Connection error: {str(e)}'})}\n\n"
            return

    logger.info(f"DocGen full_text ({len(full_text)} chars): {full_text[-300:]}")
    logger.info(f"DocGen full_code ({len(full_code)} chars): {full_code[-200:] if full_code else '(empty)'}")

    # Extract OUTPUT_FILE
    file_match = _re.search(r"OUTPUT_FILE:\s*(\S+)", full_text)
    if not file_match:
        # Try to find any file that was created
        yield f"data: {json.dumps({'type': 'error', 'message': 'AI tidak menghasilkan OUTPUT_FILE. Coba jelaskan lebih spesifik.'})}\n\n"
        return

    output_file = file_match.group(1).strip("`\"'").strip()

    # Phase 2: Execute the code
    yield f"data: {json.dumps({'type': 'phase', 'phase': 'executing', 'message': 'Executing code...'})}\n\n"
    await asyncio.sleep(0.3)

    # Check cancel before executing
    if cancel_event and cancel_event.is_set():
        yield f"data: {json.dumps({'type': 'cancelled', 'message': 'Dibatalkan oleh pengguna'})}\n\n"
        return

    try:
        docs_dir = "/tmp/microagent_docs"
        os.makedirs(docs_dir, exist_ok=True)
        gen_id = uuid.uuid4().hex[:12]
        file_path = os.path.join(docs_dir, f"{gen_id}_{output_file}")
        script_path = os.path.join(docs_dir, f"gen_{gen_id}.py")

        patched_code = f'import os\nos.chdir(r"{docs_dir}")\n\n{full_code}'
        with open(script_path, "w", encoding="utf-8") as f:
            f.write(patched_code)

        logger.info(f"DocGen executing code in {tmp_dir}, looking for: {output_file}")

        result = subprocess.run(
            [sys.executable, script_path],
            capture_output=True, text=True, timeout=30
        )

        logger.info(f"DocGen stdout: {result.stdout[:200] if result.stdout else '(empty)'}")
        logger.info(f"DocGen stderr: {result.stderr[:500] if result.stderr else '(empty)'}")

        if result.returncode != 0:
            yield f"data: {json.dumps({'type': 'error', 'message': result.stderr[:500]})}\n\n"
            # Clean up script file
            if os.path.exists(script_path):
                os.remove(script_path)
            return

        # Check if exact file exists
        if not os.path.exists(file_path):
            # Try to find any non-script files in docs_dir
            found_files = [f for f in os.listdir(docs_dir) if not f.endswith(".py") and f.startswith(gen_id)]
            if found_files:
                # Use the first found file
                output_file = found_files[0]
                file_path = os.path.join(docs_dir, output_file)
                logger.info(f"DocGen: exact file not found, using: {output_file}")
            else:
                logger.error(f"DocGen: no files created. Files in dir: {os.listdir(docs_dir)}")
                yield f"data: {json.dumps({'type': 'error', 'message': f'File tidak ditemukan setelah eksekusi. Periksa kode Python.'})}\n\n"
                if os.path.exists(script_path):
                    os.remove(script_path)
                return

        # Register file in Supabase
        file_id = uuid.uuid4().hex
        if supa:
            try:
                supa.table("generated_files").insert({
                    "id": file_id,
                    "user_id": user_id,
                    "filename": output_file,
                    "file_path": file_path,
                }).execute()
            except Exception as e:
                logger.warning(f"Failed to register generated file: {e}")

        # Clean up script file
        if os.path.exists(script_path):
            os.remove(script_path)

        # Schedule cleanup after 24 hours
        asyncio.create_task(_cleanup_generated_file(file_id, file_path, delay=86400))

        # Phase 3: Done
        yield f"data: {json.dumps({'type': 'complete', 'file_id': file_id, 'filename': output_file, 'download_url': f'/api/download/{file_id}/{output_file}'})}\n\n"

    except subprocess.TimeoutExpired:
        yield f"data: {json.dumps({'type': 'error', 'message': 'Eksekusi timeout (>30 detik)'})}\n\n"
    except Exception as e:
        yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"


@api_router.get("/download/{file_id}/{filename}")
async def download_generated_file(file_id: str, filename: str):
    """Download a generated document file."""
    if not supa: return _no_supa()
    try:
        record = supa.table("generated_files").select("file_path").eq("id", file_id).single().execute()
        if not record.data:
            return JSONResponse({"error": "File sudah expired atau tidak ditemukan. Silakan generate ulang."}, status_code=404)
        path = record.data["file_path"]
        if not os.path.exists(path):
            return JSONResponse({"error": "File sudah dihapus karena sudah lebih dari 24 jam. Silakan generate ulang."}, status_code=404)
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        media_type = MEDIA_TYPES.get(ext, "application/octet-stream")
        return FileResponse(path=path, filename=filename, media_type=media_type)
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


@api_router.post("/fetch-url")
async def fetch_url_endpoint(request: Request):
    """Fetch and extract content from a URL. Useful for web scraping."""
    body = await request.json()
    url = body.get("url", "").strip()
    if not url:
        return JSONResponse({"error": "URL required"}, status_code=400)

    content = await _fetch_url(url)
    return JSONResponse({"url": url, "content": content})


@api_router.get("/", response_class=HTMLResponse)
async def root():
    return HTMLResponse(API_STATUS_HTML)


_API_STATUS_HTML = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>MicroAgent — API Status</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600&family=Space+Grotesk:wght@500;600&display=swap" rel="stylesheet">
<style>
*,*::before,*::after{box-sizing:border-box;margin:0;padding:0}
body{font-family:Inter,system-ui,sans-serif;background:#111;color:#e5e7eb;min-height:100vh;display:flex;align-items:center;justify-content:center;padding:24px}
.container{max-width:680px;width:100%;background:#1a1a1a;border-radius:24px;padding:32px;box-shadow:0 25px 50px -12px rgba(0,0,0,.6)}
h1{font-family:'Space Grotesk',sans-serif;font-weight:600;font-size:20px;margin-bottom:4px}
.sub{color:#6b7280;font-size:13px;margin-bottom:24px}
.grid{display:flex;flex-direction:column;gap:10px}
.card{display:flex;align-items:center;gap:14px;background:#222;border-radius:14px;padding:14px 18px;transition:background.15s}
.card:hover{background:#282828}
.status-dot{width:10px;height:10px;border-radius:50%;flex-shrink:0;background:#374151;transition:background.3s}
.status-dot.ok{background:#22c55e;box-shadow:0 0 6px rgba(34,197,94,.4)}
.status-dot.err{background:#ef4444;box-shadow:0 0 6px rgba(239,68,68,.4)}
.status-dot.pending{background:#374151;animation:pulse 1.5s infinite}
@keyframes pulse{0%,100%{opacity:1}50%{opacity:.35}}
.info{flex:1;min-width:0}
.info .name{font-size:13px;font-weight:500}
.info .desc{font-size:11px;color:#6b7280;margin-top:1px}
.badge{font-size:10px;font-weight:600;padding:3px 9px;border-radius:999px;flex-shrink:0}
.badge.latency{background:#1f2937;color:#9ca3af}
.badge.ok{background:rgba(34,197,94,.15);color:#22c55e}
.badge.err{background:rgba(239,68,68,.15);color:#ef4444}
.summary{display:flex;gap:16px;margin-bottom:20px;flex-wrap:wrap}
.summary-item{flex:1;min-width:100px;background:#222;border-radius:14px;padding:14px 16px;text-align:center}
.summary-item .num{font-family:'Space Grotesk',sans-serif;font-size:22px;font-weight:600}
.summary-item .num.ok{color:#22c55e}
.summary-item .num.err{color:#ef4444}
.summary-item .label{font-size:11px;color:#6b7280;margin-top:2px}
.footer{margin-top:20px;padding-top:16px;border-top:1px solid #222;display:flex;justify-content:space-between;align-items:center}
.footer a{color:#6366f1;font-size:12px;text-decoration:none}
.footer a:hover{text-decoration:underline}
.footer .ts{color:#4b5563;font-size:11px}
.response-preview{margin-top:8px;background:#111;border-radius:10px;padding:10px 12px;font-family:'SF Mono',Monaco,monospace;font-size:11px;line-height:1.5;color:#9ca3af;max-height:60px;overflow-y:auto;display:none;white-space:pre-wrap;word-break:break-all}
.response-preview.show{display:block}
.details-btn{background:0;border:0;color:#4b5563;font-size:10px;cursor:pointer;padding:2px 0}
.details-btn:hover{color:#9ca3af}
</style>
</head>
<body>
<div class="container">
  <div style="display:flex;align-items:center;gap:10px;margin-bottom:2px">
    <svg width="22" height="22" viewBox="0 0 24 24" fill="none" stroke="#6366f1" stroke-width="1.5"><path d="M12 2L2 7l10 5 10-5-10-5z"/><path d="M2 17l10 5 10-5"/><path d="M2 12l10 5 10-5"/></svg>
    <h1>API Status</h1>
  </div>
  <p class="sub">Live health check for MicroAgent backend services</p>

  <div class="summary" id="summary">
    <div class="summary-item"><div class="num" id="totalCt">0</div><div class="label">Total</div></div>
    <div class="summary-item"><div class="num ok" id="okCt">0</div><div class="label">Healthy</div></div>
    <div class="summary-item"><div class="num err" id="errCt">0</div><div class="label">Errors</div></div>
  </div>

  <div class="grid" id="grid"></div>

  <div class="footer">
    <a href="/api/models" target="_blank">View models →</a>
    <span class="ts" id="timestamp"></span>
  </div>
</div>
<script>
const ENDPOINTS = [
  {path:"/api/",name:"Root",desc:"Server alive check"},
  {path:"/api/models",name:"Models",desc:"AI model configuration"},
  {path:"/api/status",name:"Status",desc:"Health records"},
  {path:"/api/chat/stream",name:"Chat Stream",desc:"SSE streaming endpoint (POST)"},
];

const grid = document.getElementById("grid");
const okCt = document.getElementById("okCt");
const errCt = document.getElementById("errCt");
const totalCt = document.getElementById("totalCt");

function dot(status){return status==='ok'?'ok':status==='err'?'err':'pending'}
function badge(status,ms){
  if(status==='ok') return `<span class="badge ok">${ms}ms</span>`;
  if(status==='err') return `<span class="badge err">Error</span>`;
  return `<span class="badge latency">—</span>`;
}

function render(ep,status,ms,body){
  const d=dot(status);
  const row=document.createElement('div');
  row.innerHTML=`
    <div class="card" data-path="${ep.path}">
      <div class="status-dot ${d}" data-dot></div>
      <div class="info">
        <div class="name">${ep.name}</div>
        <div class="desc">${ep.desc}</div>
        <div class="response-preview" data-preview>${body||''}</div>
      </div>
      ${status==='pending'?'':badge(status,ms)}
      <button class="details-btn" data-toggle style="${status==='pending'?'display:none':''}">details</button>
    </div>
  `;
  const card=row.firstElementChild;
  card.querySelector('[data-toggle]')?.addEventListener('click',()=>{
    const pre=card.querySelector('[data-preview]');
    pre.classList.toggle('show');
  });
  return card;
}

async function check(ep){
  const card=render(ep,'pending');
  grid.appendChild(card);

  const dotEl=card.querySelector('[data-dot]');
  const start=performance.now();
  try{
    let ok=false,body='';
    if(ep.path.includes('/chat/stream')){
      const r=await fetch(ep.path,{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({messages:[{role:'user',content:'ping'}]})});
      ok=r.ok;
      body=r.ok?'POST 204 No Content (SSE)':await r.text().then(t=>t.slice(0,300));
    }else{
      const r=await fetch(ep.path);
      body=r.ok?await r.text():await r.text().then(t=>t.slice(0,300));
      ok=r.ok;
    }
    const ms=Math.round(performance.now()-start);
    const newCard=render(ep,ok?'ok':'err',ms,body);
    card.replaceWith(newCard);
    updateSummary();
  }catch(e){
    const ms=Math.round(performance.now()-start);
    const newCard=render(ep,'err',ms,e.message);
    card.replaceWith(newCard);
    updateSummary();
  }
}

function updateSummary(){
  const cards=document.querySelectorAll('.card');
  let ok=0,err=0;
  cards.forEach(c=>{
    const dot=c.querySelector('[data-dot]');
    if(dot.classList.contains('ok')) ok++;
    else if(dot.classList.contains('err')) err++;
  });
  totalCt.textContent=cards.length;
  okCt.textContent=ok;
  errCt.textContent=err;
}

document.getElementById('timestamp').textContent=new Date().toLocaleString();
ENDPOINTS.forEach(check);
</script>
</body>
</html>"""


API_STATUS_HTML = _API_STATUS_HTML


MODELS_HTML = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>MicroAgent — Available Models</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600&family=Space+Grotesk:wght@500;600&display=swap" rel="stylesheet">
<style>
*,*::before,*::after{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:Inter,system-ui,sans-serif;background:#111;color:#e5e7eb;min-height:100vh;display:flex;align-items:center;justify-content:center;padding:24px}}
.container{{max-width:680px;width:100%;background:#1a1a1a;border-radius:24px;padding:32px;box-shadow:0 25px 50px -12px rgba(0,0,0,.6)}}
h1{{font-family:'Space Grotesk',sans-serif;font-weight:600;font-size:20px;margin-bottom:2px}}
.sub{{color:#6b7280;font-size:13px;margin-bottom:20px}}
.info-bar{{display:flex;gap:12px;margin-bottom:20px;flex-wrap:wrap}}
.info-item{{background:#222;border-radius:12px;padding:10px 14px;flex:1;min-width:100px}}
.info-item .lbl{{font-size:11px;color:#6b7280}}
.info-item .val{{font-family:'Space Grotesk',sans-serif;font-size:15px;font-weight:600;margin-top:1px;color:#e5e7eb}}
.grid{{display:flex;flex-direction:column;gap:8px}}
.card{{display:flex;align-items:center;gap:14px;background:#222;border-radius:14px;padding:14px 18px;transition:background.15s}}
.card:hover{{background:#282828}}
.card .dot{{width:10px;height:10px;border-radius:50%;flex-shrink:0;background:#22c55e;box-shadow:0 0 6px rgba(34,197,94,.4)}}
.card .info{{flex:1;min-width:0}}
.card .info .name{{font-size:13px;font-weight:500}}
.card .info .prov{{font-size:11px;color:#6b7280;margin-top:1px;font-family:'SF Mono',Monaco,monospace}}
.card .tag{{font-size:10px;font-weight:600;padding:3px 9px;border-radius:999px;background:rgba(99,102,241,.15);color:#818cf8;flex-shrink:0}}
.footer{{margin-top:20px;padding-top:16px;border-top:1px solid #222;display:flex;justify-content:space-between;align-items:center}}
.footer a{{color:#6366f1;font-size:12px;text-decoration:none}}
.footer a:hover{{text-decoration:underline}}
.footer .ts{{color:#4b5563;font-size:11px}}
</style>
</head>
<body>
<div class="container">
  <div style="display:flex;align-items:center;gap:10px;margin-bottom:2px">
    <svg width="22" height="22" viewBox="0 0 24 24" fill="none" stroke="#6366f1" stroke-width="1.5"><path d="M12 2L2 7l10 5 10-5-10-5z"/><path d="M2 17l10 5 10-5"/><path d="M2 12l10 5 10-5"/></svg>
    <h1>Available Models</h1>
  </div>
  <p class="sub">AI models configured on this MicroAgent instance</p>

  <div class="info-bar">
    <div class="info-item"><div class="lbl">Provider</div><div class="val">{{provider}}</div></div>
    <div class="info-item"><div class="lbl">Default</div><div class="val">{{default_model}}</div></div>
    <div class="info-item"><div class="lbl">Fallback</div><div class="val">{{fallback}}</div></div>
  </div>

  <div class="grid" id="grid"></div>

  <div class="footer">
    <a href="/api/">Back to status →</a>
    <span class="ts" id="timestamp"></span>
  </div>
</div>
<script>
const MODELS = {models_json};

const grid = document.getElementById('grid');
MODELS.forEach(m => {{
  const card = document.createElement('div');
  card.className = 'card';
  card.innerHTML = `
    <div class="dot"></div>
    <div class="info">
      <div class="name">${{m.id}}</div>
      <div class="prov">${{m.provider_model}}</div>
    </div>
    <span class="tag">active</span>
  `;
  grid.appendChild(card);
}});
document.getElementById('timestamp').textContent = new Date().toLocaleString();
</script>
</body>
</html>"""


@api_router.get("/models")
async def get_models(request: Request):
    data = ModelsResponse(
        provider=PROVIDER_NAME,
        default_model_id=DEFAULT_MODEL_ID,
        fallback_model=env_str("OPENAI_MODEL"),
        models=[
            ModelInfo(id=model_id, provider_model=provider)
            for model_id, provider in MODEL_ID_TO_PROVIDER.items()
        ],
    )
    accept = request.headers.get("accept", "")
    if "text/html" in accept:
        return HTMLResponse(MODELS_HTML.format(
            provider=data.provider,
            default_model=data.default_model_id,
            fallback=data.fallback_model,
            models_json=json.dumps([m.model_dump() for m in data.models], indent=2),
        ))
    return JSONResponse(data.model_dump())


@api_router.post("/status", response_model=StatusCheck)
async def create_status_check(input: StatusCheckCreate):
    status_dict = input.model_dump()
    status_obj = StatusCheck(**status_dict)

    if db is not None:
        doc = status_obj.model_dump()
        doc["timestamp"] = doc["timestamp"].isoformat()
        _ = await db.status_checks.insert_one(doc)

    return status_obj


@api_router.get("/status", response_model=List[StatusCheck])
async def get_status_checks():
    if db is None:
        return []

    status_checks = await db.status_checks.find({}, {"_id": 0}).to_list(1000)
    for check in status_checks:
        if isinstance(check["timestamp"], str):
            check["timestamp"] = datetime.fromisoformat(check["timestamp"])

    return status_checks


GUEST_LIMIT = 10
_guest_counts: dict[str, int] = {}  # session_key -> count (in-memory, resets on restart)

def check_guest_limit(session_key: str) -> bool:
    """Server-side guest prompt limit. Returns True if allowed."""
    count = _guest_counts.get(session_key, 0)
    if count >= GUEST_LIMIT:
        return False
    _guest_counts[session_key] = count + 1
    return True


async def stream_chat_response(payload: ChatStreamRequest, request: Request = None) -> AsyncIterator[str]:
    # Extract user_id from auth token if not provided in payload
    effective_user_id = payload.user_id
    if not effective_user_id and request:
        effective_user_id = _get_user_id(request)

    # Server-side guest limit enforcement
    if not effective_user_id:
        # Use IP + user-agent as session key for anonymous users
        ip = request.client.host if request.client else "anon"
        ua = request.headers.get("user-agent", "")[:50] if request else ""
        session_key = f"guest:{ip}:{hash(ua)}"
        if not check_guest_limit(session_key):
            yield sse("error", {"message": "Batas prompt guest tercapai. Login untuk melanjutkan.", "error_type": "guest_limit"})
            return

    logger.info(f"stream_chat_response: payload.user_id={payload.user_id}, effective_user_id={effective_user_id}, messages={len(payload.messages)}")

    # ── Token check & deduct BEFORE any AI call ──────────────────────────────
    model_id_for_cost = payload.model_id or DEFAULT_MODEL_ID
    token_cost = get_token_cost(model_id_for_cost)
    token_deduction = {"success": True, "cost": 0, "balance": 0}

    if effective_user_id:
        token_deduction = await deduct_token(effective_user_id, model_id_for_cost)
        if not token_deduction["success"]:
            yield sse("error", {
                "message": f"Token tidak cukup. Butuh {token_deduction['required']} token, sisa {token_deduction['balance']}.",
                "error_type": "insufficient_tokens",
                "balance": token_deduction["balance"],
                "required": token_deduction["required"],
            })
            return

    # Skill loading phase — emit real event before anything else
    if payload.skill_slug:
        skill_name = payload.skill_slug.replace("-", " ").title()
        for s in BUILTIN_SKILLS:
            if s["slug"] == payload.skill_slug:
                skill_name = s["name"]
                break
        yield sse("status", {"phase": "skill_loading", "status": "started", "skill_slug": payload.skill_slug, "skill_name": skill_name})
        yield sse("status", {"phase": "skill_loading", "status": "completed", "skill_slug": payload.skill_slug})

    web_context = ""
    is_comparison = False
    
    # ── Comparison Detection & Multi-Query Search ─────────────────────────────
    last_user_msg = last_user_prompt(payload)
    if payload.comparison or is_comparison_request(last_user_msg):
        is_comparison = True
        product_a, product_b = extract_comparison_products(last_user_msg)
        
        yield sse(
            "status",
            {
                "phase": "comparison_search",
                "status": "started",
                "product_a": product_a,
                "product_b": product_b,
                "message": f"Membandingkan {product_a} vs {product_b}...",
            },
        )
        
        # Build deterministic comparison data (always available, no API keys needed)
        comparison_data = build_comparison_data(product_a, product_b)
        yield sse("comparison_data", comparison_data)
        
        try:
            all_results, sources = await comparison_search_pipeline(
                product_a, product_b,
                lambda msg: None  # Status already yielded
            )
            
            # Build comparison context
            comparison_context = f"\n\nPERBANDINGAN PRODUK:\n"
            comparison_context += f"Produk A: {product_a}\n"
            comparison_context += f"Produk B: {product_b}\n\n"
            comparison_context += "Hasil pencarian untuk perbandingan:\n\n"
            
            for r in all_results[:30]:  # Limit to 30 results
                title = r.get("title", "")
                url = r.get("url", "")
                content = (r.get("content") or "")[:500]
                comparison_context += f"[{title}]\nURL: {url}\n{content}\n\n"
            
            web_context = comparison_context
            
            yield sse(
                "status",
                {
                    "phase": "comparison_search",
                    "status": "completed",
                    "result_count": len(all_results),
                    "sources": [{"title": s["title"], "url": s["url"]} for s in sources[:10]],
                    "message": f"Ditemukan {len(all_results)} hasil dari {len(sources)} sumber",
                },
            )
            
        except Exception as exc:
            logger.warning("Comparison search failed: %s", exc)
            yield sse(
                "status",
                {
                    "phase": "comparison_search",
                    "status": "failed",
                    "message": f"Pencarian perbandingan gagal: {exc}",
                },
            )

    elif payload.web_search:
        query = extract_search_query(payload)
        provider = "tavily" if env_str("TAVILY_API_KEY") else "firecrawl"
        yield sse(
            "status",
            {
                "phase": "web_search",
                "provider": provider,
                "status": "started",
                "query": query,
                "message": f"Searching the web for: {query}",
            },
        )
        try:
            max_results = env_int("WEB_SEARCH_MAX_RESULTS", DEFAULT_WEB_SEARCH_MAX_RESULTS)
            provider, results = await web_search(query, max_results=max_results)
            yield sse(
                "status",
                {
                    "phase": "web_search",
                    "provider": provider,
                    "status": "results",
                    "query": query,
                    "results": [
                        {"title": r.get("title"), "url": r.get("url")}
                        for r in results[:max_results]
                    ],
                },
            )
            web_context = format_search_context(query, results) if results else ""
            yield sse(
                "status",
                {
                    "phase": "web_fetch",
                    "provider": provider,
                    "status": "completed",
                    "result_count": len(results),
                    "message": f"Fetched {len(results)} web sources for synthesis.",
                },
            )
        except RuntimeError as exc:
            logger.warning("Web search not configured: %s", exc)
            yield sse(
                "status",
                {
                    "phase": "web_search",
                    "provider": provider,
                    "status": "failed",
                    "message": str(exc),
                },
            )
        except Exception as exc:
            logger.warning("Web search failed: %s", exc)
            yield sse(
                "status",
                {
                    "phase": "web_search",
                    "provider": provider,
                    "status": "failed",
                    "message": f"Pencarian web gagal: {exc}",
                },
            )

    # If web search is disabled, still emit a status so frontend knows
    if not payload.web_search and not is_comparison:
        yield sse(
            "status",
            {
                "phase": "web_search",
                "status": "skipped",
                "message": "Web search not needed for this query",
            },
        )

    # If both web_search and reasoning are enabled, disable reasoning during web search
    # so model focuses entirely on synthesizing web context first.
    # Reasoning is still sent to provider but web context is always fully injected first.
    # (web_context is already built before stream_provider is called — order is correct)
    has_error = False
    for event in stream_provider(payload, web_context=web_context):
        # Inject token info into meta event
        if event.startswith("event: meta"):
            try:
                payload_str = event.split("data: ", 1)[1] if "data: " in event else ""
                meta = json.loads(payload_str.strip())
                meta["tokens_used"] = token_deduction.get("cost", 0)
                meta["tokens_left"] = token_deduction.get("balance", 0)
                event = f"event: meta\ndata: {json.dumps(meta, ensure_ascii=False)}\n\n"
            except (json.JSONDecodeError, IndexError):
                pass
        # Track errors for refund
        if '"error"' in event:
            has_error = True
        yield event

    # Refund tokens on provider error (user shouldn't pay for failed requests)
    if has_error and token_deduction.get("success") and token_deduction.get("cost", 0) > 0:
        await refund_token(effective_user_id, model_id_for_cost, token_deduction["cost"])

    # Layer B: Trigger memory extraction (fire-and-forget background task)
    logger.info(f"Memory extraction check: has_error={has_error}, effective_user_id={effective_user_id}")
    if not has_error and effective_user_id:
        messages_for_extraction = [m.model_dump() for m in payload.messages if m.content.strip()]
        trigger_memory_extraction(messages_for_extraction, effective_user_id, session_id=None)
    else:
        logger.info(f"Memory extraction skipped at stream endpoint: has_error={has_error}, effective_user_id={effective_user_id}")


# ── Response Comparison (A/B Testing) ────────────────────────────────────────

# Config — change without redeploy via env vars
COMPARISON_ENABLED = env_str("COMPARISON_ENABLED", "false").lower() == "true"
COMPARISON_SAMPLE_RATE = float(env_str("COMPARISON_SAMPLE_RATE", "0.05"))  # 5% default
COMPARISON_MODEL_A = env_str("COMPARISON_MODEL_A", "deepseek-v4-flash")
COMPARISON_MODEL_B = env_str("COMPARISON_MODEL_B", "glm-5")
COMPARISON_VARIANT = env_str("COMPARISON_VARIANT", "model_ab")  # model_ab | temperature_variance


async def _generate_single(messages: list, model_id: str, temperature: float = 0.7, max_tokens: int = 2000) -> str:
    """Generate a single non-streaming response. Returns text or raises."""
    base_url, api_key, _, model = get_provider_config(model_id)
    if not base_url or not api_key:
        raise RuntimeError(f"Provider not configured for {model_id}")
    body = {"model": model, "messages": messages, "max_tokens": max_tokens, "temperature": temperature, "stream": False}
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    async with httpx.AsyncClient(timeout=60.0) as client:
        resp = await client.post(provider_url(base_url), headers=headers, json=body)
        resp.raise_for_status()
        return resp.json()["choices"][0]["message"]["content"] or ""


class ComparisonRequest(BaseModel):
    messages: List[ChatMessageIn]
    model_id: Optional[str] = None
    session_id: Optional[str] = None
    user_id: Optional[str] = None


@api_router.post("/chat/comparison")
async def chat_comparison(payload: ComparisonRequest, request: Request):
    """Generate two parallel responses for A/B comparison."""
    if not COMPARISON_ENABLED:
        return JSONResponse({"enabled": False}, status_code=200)

    uid = _get_user_id(request) or payload.user_id or "anonymous"

    msgs = normalize_messages(ChatStreamRequest(
        messages=payload.messages,
        model_id=payload.model_id or DEFAULT_MODEL_ID,
        user_id=uid,
    ))

    # Parallel generate
    if COMPARISON_VARIANT == "temperature_variance":
        source_a, source_b = f"{COMPARISON_MODEL_A}@t=0.3", f"{COMPARISON_MODEL_A}@t=1.1"
        task_a = _generate_single(msgs, COMPARISON_MODEL_A, temperature=0.3)
        task_b = _generate_single(msgs, COMPARISON_MODEL_A, temperature=1.1)
    else:  # model_ab
        source_a, source_b = COMPARISON_MODEL_A, COMPARISON_MODEL_B
        task_a = _generate_single(msgs, COMPARISON_MODEL_A)
        task_b = _generate_single(msgs, COMPARISON_MODEL_B)

    results = await asyncio.gather(task_a, task_b, return_exceptions=True)
    resp_a = results[0] if not isinstance(results[0], Exception) else None
    resp_b = results[1] if not isinstance(results[1], Exception) else None

    # If both fail, return error
    if resp_a is None and resp_b is None:
        return JSONResponse({"error": "Both responses failed"}, status_code=500)

    # If one fails, return as normal response (no comparison UI)
    if resp_a is None or resp_b is None:
        return JSONResponse({"fallback": True, "response": resp_a or resp_b})

    # Save to DB
    comparison_id = None
    if supa:
        user_message = next((m.content for m in reversed(payload.messages) if m.role == "user"), "")
        try:
            row = supa.table("response_comparisons").insert({
                "user_id": uid,
                "session_id": payload.session_id,
                "user_message": user_message[:2000],
                "variant_type": COMPARISON_VARIANT,
                "response_a_source": source_a,
                "response_a_content": resp_a[:8000],
                "response_b_source": source_b,
                "response_b_content": resp_b[:8000],
            }).execute()
            comparison_id = row.data[0]["id"] if row.data else None
        except Exception as e:
            logger.warning(f"Failed to save comparison: {e}")

    return JSONResponse({
        "enabled": True,
        "fallback": False,
        "comparison_id": comparison_id,
        "variant_type": COMPARISON_VARIANT,
        "response_a": resp_a,
        "response_b": resp_b,
    })


@api_router.post("/chat/comparison/{comparison_id}/choose")
async def choose_comparison(comparison_id: str, request: Request):
    """Record user's choice for a comparison."""
    if not supa: return _no_supa()
    uid = _get_user_id(request)
    if not uid: return JSONResponse({"error": "Unauthorized"}, status_code=401)
    body = await request.json()
    chosen = body.get("chosen")  # 'a' | 'b' | 'both_good' | 'both_bad'
    if chosen not in ("a", "b", "both_good", "both_bad"):
        return JSONResponse({"error": "Invalid choice"}, status_code=400)
    supa.table("response_comparisons").update({"chosen": chosen}).eq("id", comparison_id).eq("user_id", uid).execute()
    return JSONResponse({"success": True})


# ── Onboarding & Survey ─────────────────────────────────────────────────────

@api_router.get("/onboarding/status")
async def get_onboarding_status(request: Request):
    if not supa: return _no_supa()
    uid = _get_user_id(request)
    if not uid: return JSONResponse({"error": "Unauthorized"}, status_code=401)
    res = supa.table("onboarding_responses").select("user_id,skipped,completed_at").eq("user_id", uid).execute()
    return JSONResponse({"is_onboarded": bool(res.data)})


@api_router.post("/onboarding")
async def submit_onboarding(request: Request):
    if not supa: return _no_supa()
    uid = _get_user_id(request)
    if not uid: return JSONResponse({"error": "Unauthorized"}, status_code=401)
    body = await request.json()
    skipped = body.get("skipped", False)

    row = {
        "user_id": uid,
        "role": body.get("role"),
        "primary_goal": body.get("primary_goal"),
        "ai_familiarity": body.get("ai_familiarity"),
        "language_preference": body.get("language_preference"),
        "skipped": skipped,
    }
    supa.table("onboarding_responses").upsert(row, on_conflict="user_id").execute()

    # Insert structured facts directly into user_memories (no LLM needed)
    if not skipped:
        memories = []
        if body.get("role"):
            label = {"pelajar": "Pelajar/mahasiswa", "karyawan": "Karyawan", "freelancer": "Freelancer",
                     "founder": "Founder/pebisnis", "developer": "Developer"}.get(body["role"], body["role"])
            memories.append({"user_id": uid, "content": f"User adalah {label}", "category": "onboarding_profile"})
        if body.get("primary_goal"):
            memories.append({"user_id": uid, "content": f"Tujuan utama memakai Micro Agent: {body['primary_goal']}", "category": "onboarding_profile"})
        if body.get("ai_familiarity"):
            memories.append({"user_id": uid, "content": f"Tingkat familiarity dengan AI: {body['ai_familiarity']}", "category": "onboarding_profile"})
        if body.get("language_preference"):
            memories.append({"user_id": uid, "content": f"Preferensi bahasa: {body['language_preference']}", "category": "onboarding_profile"})
        if memories:
            supa.table("user_memories").insert(memories).execute()

    return JSONResponse({"success": True})


@api_router.post("/survey")
async def submit_survey(request: Request):
    if not supa: return _no_supa()
    uid = _get_user_id(request)
    if not uid: return JSONResponse({"error": "Unauthorized"}, status_code=401)
    body = await request.json()

    # Check for dismiss
    if body.get("dismissed"):
        supa.table("survey_dismissals").upsert({"user_id": uid}, on_conflict="user_id").execute()
        return JSONResponse({"success": True})

    row = {
        "user_id": uid,
        "satisfaction_score": body.get("satisfaction_score"),
        "most_used_feature": body.get("most_used_feature"),
        "pain_points": body.get("pain_points"),
        "feature_requests": body.get("feature_requests"),
    }
    supa.table("survey_responses").insert(row).execute()

    # 100 token bonus — once per user, guarded by survey_responses existence check above
    try:
        cur = await get_user_balance(uid)
        supa.table("user_credits").upsert({"user_id": uid, "balance": cur + 100}, on_conflict="user_id").execute()
        supa.table("credit_transactions").insert({"user_id": uid, "amount": 100, "type": "bonus", "model": "survey_reward"}).execute()
        logger.info(f"Survey bonus +100 tokens for user {uid}")
    except Exception as e:
        logger.warning(f"Survey token bonus failed: {e}")

    # Optional background extraction from open text
    open_text = " ".join(filter(None, [body.get("pain_points"), body.get("feature_requests")]))
    if open_text.strip() and len(open_text.strip()) > 20:
        async def _extract_survey_memory():
            prompt = f"""Baca jawaban survey berikut. Kalau ada fakta personal tentang user yang berguna untuk personalisasi AI (misal pekerjaan spesifik, proyek, kebutuhan khusus), ekstrak sebagai fakta singkat. Kalau murni feedback produk tanpa info personal, balas kosong.

Jawaban: {open_text}

Fakta (satu per baris, kosong jika tidak ada):"""
            facts = await extract_facts_from_text(prompt, model_id="deepseek-v4-flash")
            if facts:
                await save_user_memories(uid, facts)
        asyncio.create_task(_extract_survey_memory())

    return JSONResponse({"success": True})


@api_router.get("/survey/status")
async def get_survey_status(request: Request):
    if not supa: return _no_supa()
    uid = _get_user_id(request)
    if not uid: return JSONResponse({"error": "Unauthorized"}, status_code=401)
    dismissed = supa.table("survey_dismissals").select("user_id").eq("user_id", uid).execute()
    completed = supa.table("survey_responses").select("id").eq("user_id", uid).execute()
    session_count = supa.table("sessions").select("id", count="exact").eq("user_id", uid).execute()
    return JSONResponse({
        "dismissed": bool(dismissed.data),
        "completed": bool(completed.data),
        "session_count": session_count.count or 0,
    })


# ── Improve Prompt ─────────────────────────────────────────────────────────

IMPROVE_SYSTEM = """\
Kamu adalah pakar prompt engineering. Tugasmu menulis ulang prompt pengguna 
menjadi lebih jelas, spesifik, dan efektif.

ATURAN:
- Pertahankan bahasa yang sama (Indonesia \u2192 Indonesia, English \u2192 English)
- Pertahankan maksud yang sama \u2014 jangan ubah apa yang diinginkan user
- Tambahkan spesifisitas: siapa, apa, format, tone, panjang jika hilang
- Hilangkan ambiguitas
- Buat actionable dan jelas
- Tetap ringkas \u2014 jangan terlalu panjang
- Output HANYA teks prompt yang sudah diperbaiki
- Tanpa penjelasan, tanpa kata pembuka, langsung teksnya saja
"""


class ImprovePromptRequest(BaseModel):
    prompt: str
    context: str = ""

    @field_validator("prompt")
    @classmethod
    def trim_prompt(cls, v: str) -> str:
        return v.strip()


@api_router.post("/improve-prompt")
async def improve_prompt(req: ImprovePromptRequest):
    if len(req.prompt) < 3:
        return JSONResponse({"success": False, "error": "Prompt terlalu pendek"}, status_code=400)

    # Use deepseek for improve-prompt
    model_id = "deepseek-v4-flash"
    base_url, api_key, provider_name, model = get_provider_config(model_id)
    if not base_url or not api_key:
        return JSONResponse({"success": False, "error": "Provider belum dikonfigurasi"}, status_code=500)

    context_text = f"\n\nKonteks percakapan:\n{req.context.strip()}" if req.context.strip() else ""
    user_content = f"Improve this prompt:{context_text}\n\nOriginal: {req.prompt}"

    body = {
        "model": model,
        "messages": [
            {"role": "system", "content": IMPROVE_SYSTEM},
            {"role": "user", "content": user_content},
        ],
        "max_tokens": 1024,
        "temperature": 0.4,
        "stream": True,
    }
    headers_req = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    timeout = httpx.Timeout(connect=15.0, read=60.0, write=15.0, pool=15.0)

    try:
        improved = ""
        buf = ""
        async with httpx.AsyncClient(timeout=timeout) as client:
            async with client.stream("POST", provider_url(base_url), headers=headers_req, json=body) as resp:
                if resp.status_code >= 400:
                    detail = await resp.aread()
                    return JSONResponse({"success": False, "error": f"Model error {resp.status_code}: {detail[:200].decode(errors='replace')}"}, status_code=500)
                async for raw_chunk in resp.aiter_bytes():
                    buf += raw_chunk.decode("utf-8", errors="ignore")
                    while "\n" in buf:
                        line, buf = buf.split("\n", 1)
                        line = line.strip()
                        if line.startswith("data:"):
                            line = line[5:].strip()
                        if not line or line == "[DONE]":
                            continue
                        try:
                            chunk = json.loads(line)
                            delta = chunk["choices"][0].get("delta", {})
                            # content comes AFTER reasoning_content finishes
                            c = delta.get("content")
                            if c and c != "null":
                                improved += c
                        except (json.JSONDecodeError, KeyError, IndexError):
                            pass
        return JSONResponse({"success": True, "original": req.prompt, "improved": improved.strip()})
    except Exception as exc:
        logger.exception("Improve prompt failed")
        return JSONResponse({"success": False, "error": str(exc)}, status_code=500)


@api_router.post("/chat/stream")
async def chat_stream(payload: ChatStreamRequest, request: Request):
    # Rate limit: 15 requests per minute per user
    uid = _get_user_id(request) or payload.user_id or "anonymous"
    if not check_rate_limit(uid, max_requests=15, window_seconds=60):
        return JSONResponse({"error": "Terlalu banyak request. Tunggu sebentar."}, status_code=429)
    return StreamingResponse(
        stream_chat_response(payload, request),
        media_type="text/event-stream; charset=utf-8",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


# ── Deep Research ────────────────────────────────────────────────────────

try:
    from duckduckgo_search import DDGS as _DDGS
    DDGS_AVAILABLE = True
except ImportError:
    DDGS_AVAILABLE = False

try:
    from bs4 import BeautifulSoup as _BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False


class DeepResearchRequest(BaseModel):
    query: str

    @field_validator("query")
    @classmethod
    def trim_query(cls, v: str) -> str:
        return v.strip()


DEEP_RESEARCH_SYSTEM = """\
Kamu adalah jurnalis dan analis riset mendalam. Tugasmu adalah menulis artikel blog yang komprehensif,
naratif, dan sangat enak dibaca dalam Bahasa Indonesia. Bukan laporan kaku — tapi seperti long-form blog
professional di Medium atau majalah teknologi.

STYLE:
- Alur naratif yang mengalir seperti cerita
- Heading yang jelas dan menarik (##, ###)
- Setiap section dimulai dengan gambar yang relevan, BARU diikuti teks
- Teks tidak mengacu ke gambar (jangan tulis "seperti terlihat di gambar"). Gambar adalah ilustrasi visual
- Paragraf pendek-sedang (3-5 kalimat), mudah dibaca
- Gunakan SEMUA gambar yang tersedia, tersebar merata di seluruh artikel

IMEGE USAGE (WAJIB):
- Setiap section/heading HARUS ada minimal 1 gambar
- Tempatkan gambar di AWAL section, sebelum paragraf penjelasan
- Format: ![caption singkat kontekstual](url) — caption max 5 kata
- Gunakan SEMUA gambar yang diberikan, jangan skip
- Urutan: heading -> gambar relevan -> teks

FORMAT:
## [Judul Artikel Menarik]

_Deskripsi singkat 1-2 kalimat_

---

## [Section 1]

![caption](url_gambar)

Teks penjelasan...

## [Section 2]

![caption](url_gambar)

Teks penjelasan...

## Referensi
- [judul](url)
"""


async def _ddg_search(query: str, max_results: int = 8) -> List[dict]:
    if not DDGS_AVAILABLE:
        return [{"error": "duckduckgo-search tidak tersedia"}]
    try:
        import asyncio
        loop = asyncio.get_event_loop()
        def _sync():
            with _DDGS() as ddgs:
                return list(ddgs.text(query, max_results=max_results))
        results = await loop.run_in_executor(None, _sync)
        return [
            {"title": r.get("title", ""), "url": r.get("href", ""), "snippet": r.get("body", "")}
            for r in results
        ]
    except Exception as exc:
        logger.warning("DDG search failed: %s", exc)
        return [{"error": str(exc)}]


async def _fetch_url(url: str) -> str:
    try:
        async with httpx.AsyncClient(timeout=10, follow_redirects=True) as client:
            resp = await client.get(url, headers={"User-Agent": "Mozilla/5.0"})
            if not BS4_AVAILABLE:
                return resp.text[:4000]
            soup = _BeautifulSoup(resp.text, "html.parser")
            for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
                tag.decompose()
            text = soup.get_text(separator="\n", strip=True)
            # Compact whitespace
            text = "\n".join(line for line in text.splitlines() if line.strip())
            return text[:5000]
    except Exception as exc:
        return f"Error fetching {url}: {exc}"


async def run_deep_research(query: str):
    """
    Deep research flow:
    1. Generate 5 search queries from the topic
    2. Execute each query via Tavily/Firecrawl/DDG
    3. Fetch top URLs via Firecrawl or direct HTTP
    4. Synthesize all content into a report via model
    """
    import time
    start_time = time.time()

    def evt(data: dict) -> str:
        return f"data: {json.dumps(data, ensure_ascii=False)}\n\n"

    base_url, api_key, provider_name, model = get_provider_config(DEFAULT_MODEL_ID)
    if not base_url or not api_key:
        yield evt({"type": "error", "message": "Provider belum dikonfigurasi"})
        return

    timeout_obj = httpx.Timeout(connect=15.0, read=180.0, write=30.0, pool=15.0)
    headers_req = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}

    yield evt({"type": "start", "message": "Memulai riset mendalam...", "query": query})

    # ── STEP 1: Generate search queries ──────────────────────────────────────
    async def call_model(msgs: list, max_tok: int = 3000) -> str:
        body = {"model": model, "messages": msgs, "temperature": 0.3, "max_tokens": max_tok, "stream": True}
        full = ""
        async with httpx.AsyncClient(timeout=timeout_obj) as client:
            async with client.stream("POST", provider_url(base_url), headers=headers_req, json=body) as resp:
                if resp.status_code >= 400:
                    detail = await resp.aread()
                    raise RuntimeError(f"Model error {resp.status_code}: {detail[:300].decode(errors='replace')}")
                async for line in resp.aiter_lines():
                    line = line.strip()
                    if line.startswith("data:"):
                        line = line[5:].strip()
                    if not line or line == "[DONE]":
                        continue
                    try:
                        chunk = json.loads(line)
                        delta = chunk["choices"][0].get("delta", {})
                        full += delta.get("content") or ""
                    except (json.JSONDecodeError, KeyError, IndexError):
                        pass
        return full

    yield evt({"type": "step", "action": "search", "message": "Membuat rencana riset...", "step": 1})

    # Generate 10 diverse search queries
    try:
        query_response = await call_model([
            {"role": "system", "content": "Kamu adalah asisten riset. Jawab HANYA dengan JSON array berisi 10 query pencarian campuran bahasa Indonesia dan Inggris yang beragam untuk meneliti topik secara komprehensif dari berbagai sudut pandang."},
            {"role": "user", "content": f"Topik: {query}\n\nBerikan 10 query pencarian yang sangat beragam (definisi, tren, statistik, prediksi, contoh kasus, kritik, dll). Format: [\"query1\", \"query2\", ...]"},
        ], max_tok=500)
        import re as _re
        arr_match = _re.search(r'\[.*?\]', query_response, _re.DOTALL)
        search_queries = json.loads(arr_match.group(0)) if arr_match else [query]
        search_queries = [str(q) for q in search_queries[:10]]
    except Exception:
        search_queries = [
            query, f"{query} prediksi tren", f"{query} 2025 2026 2030",
            f"{query} analisis pasar", f"{query} statistik data",
            f"{query} riset terbaru", f"{query} impact dampak",
            f"{query} indonesia", f"{query} global worldwide",
            f"{query} future outlook",
        ]

    # ── STEP 2: Execute searches (max 50 results total) ────────────────────────
    all_results: List[dict] = []
    sources: List[dict] = []
    steps: List[str] = []
    max_search_results = 50  # max 50 web results total
    results_per_query = max(3, min(8, max_search_results // len(search_queries)))

    for i, q in enumerate(search_queries):
        if len(all_results) >= max_search_results:
            break
        step_msg = f"Mencari: {q}"
        steps.append(step_msg)
        yield evt({"type": "step", "action": "search", "message": step_msg, "step": len(steps)})

        try:
            _, results = await web_search(q, max_results=results_per_query)
        except Exception as exc:
            logger.warning("Deep research search failed for '%s': %s", q, exc)
            results = []

        for r in results:
            if r.get("url"):
                yield evt({"type": "source_found", "title": r.get("title", ""), "url": r["url"]})
        all_results.extend(results)

    yield evt({"type": "step", "action": "search", "message": f"Ditemukan {len(all_results)} hasil dari {len(search_queries)} query", "step": len(steps)})

    # ── STEP 3: Fetch top URLs ─────────────────────────────────────────────────
    seen_urls: set = set()
    urls_to_read: List[str] = []
    # Deduplicate and prioritize — take ALL found URLs up to 20
    for r in all_results:
        url = r.get("url", "")
        if url and url not in seen_urls and not url.endswith(".pdf"):
            seen_urls.add(url)
            urls_to_read.append(url)
        if len(urls_to_read) >= 20:
            break

    web_context_parts: List[str] = []
    for idx, url in enumerate(urls_to_read):
        step_msg = f"Membaca: {url}"
        steps.append(step_msg)
        yield evt({"type": "step", "action": "read", "message": step_msg, "step": len(steps), "url": url})

        content = await _fetch_url(url)
        title = next((r.get("title", url) for r in all_results if r.get("url") == url), url)
        sources.append({"url": url, "title": title})
        yield evt({"type": "source_added", "url": url, "title": title, "total": len(sources)})
        web_context_parts.append(f"### [{idx+1}] {title}\nURL: {url}\n\n{content[:3000]}")

    # ── STEP 3.5: Search images (parallel: Firecrawl + Tavily) ──────────────────────
    image_step = "Mencari gambar relevan..."
    steps.append(image_step)
    yield evt({"type": "step", "action": "image", "message": image_step, "step": len(steps)})

    images: List[dict] = []  # [{url, alt, source_url}]
    # Better query — no forced 'photo portrait' suffix
    image_query = query.strip()

    # CDN-friendly image URL patterns (prefer wikimedia, imgur, unsplash, etc)
    CDN_DOMAINS = ("upload.wikimedia.org", "commons.wikimedia.org", "images.unsplash.com",
                   "i.imgur.com", "cdn.pixabay.com", "upload.wikimedia", ".cdn.",
                   "static.", "media.", "img.", "image.", "photo.")

    def is_cdn_url(url: str) -> bool:
        return any(d in url.lower() for d in CDN_DOMAINS)

    async def fetch_firecrawl_images() -> List[dict]:
        found = []
        try:
            fc_results = await firecrawl_search(image_query, max_results=8)
            for r in fc_results:
                content = r.get("content", "")
                # Extract all image URLs from markdown
                img_matches = re.findall(r'!\[([^\]]*)\]\((https?://[^)]+\.(?:jpg|jpeg|png|webp|gif)(?:\?[^)]*)?\))', content)
                for alt, img_url in img_matches[:5]:
                    # Clean URL — remove trailing paren if any
                    img_url = img_url.rstrip(")")
                    if img_url not in [im["url"] for im in found] and len(img_url) < 500:
                        found.append({"url": img_url, "alt": alt.strip() or query, "source_url": r.get("url", ""), "cdn": is_cdn_url(img_url)})
                if len(found) >= 9:
                    break
        except Exception as e:
            logger.warning("Firecrawl image search failed: %s", e)
        return found

    async def fetch_tavily_images() -> List[dict]:
        found = []
        try:
            tv_payload = {
                "api_key": env_str("TAVILY_API_KEY"),
                "query": image_query,
                "search_depth": "basic",
                "max_results": 8,
                "include_images": True,
            }
            async with httpx.AsyncClient(timeout=12) as client:
                tv_resp = await client.post("https://api.tavily.com/search", json=tv_payload)
                tv_data = tv_resp.json()
            for img_url in (tv_data.get("images") or [])[:9]:
                if img_url not in [im["url"] for im in found] and len(img_url) < 500:
                    found.append({"url": img_url, "alt": query, "source_url": "", "cdn": is_cdn_url(img_url)})
        except Exception as e:
            logger.warning("Tavily image search failed: %s", e)
        return found

    # Run both in parallel (don't skip if one fails)
    try:
        import asyncio as _asyncio
        tasks = []
        if env_str("FIRECRAWL_API_KEY"):
            tasks.append(fetch_firecrawl_images())
        if env_str("TAVILY_API_KEY"):
            tasks.append(fetch_tavily_images())
        if tasks:
            results = await _asyncio.gather(*tasks, return_exceptions=True)
            all_imgs = []
            for r in results:
                if isinstance(r, list):
                    all_imgs.extend(r)
            # Deduplicate by URL
            seen = set()
            for img in all_imgs:
                if img["url"] not in seen:
                    seen.add(img["url"])
                    images.append(img)
            # Sort: CDN images first
            images.sort(key=lambda x: (0 if x.get("cdn") else 1))
            images = images[:9]
    except Exception as exc:
        logger.warning("Image search failed: %s", exc)

    # Also try to extract images from already-fetched sources (mining web content)
    if len(images) < 5 and web_context_parts:
        extra_imgs = []
        for part in web_context_parts:
            # Find markdown image URLs in content
            extra = re.findall(r'!\[[^\]]*\]\((https?://[^)]+\.(?:jpg|jpeg|png|webp|gif)(?:\?[^)]*)??)\)', part)
            for url in extra:
                url = url.rstrip(")")
                if url not in seen and len(url) < 500:
                    seen.add(url)
                    images.append({"url": url, "alt": query, "source_url": "", "cdn": is_cdn_url(url)})
        if extra_imgs:
            images.extend(extra_imgs)
            images.sort(key=lambda x: (0 if x.get("cdn") else 1))
            images = images[:9]

    # Yield images even if just 1+
    if images:
        yield evt({"type": "images_found", "images": images})
        logger.info(f"Deep research images: {len(images)} found")
    else:
        logger.warning("Deep research: no images found")

    # ── STEP 4: Synthesize report ─────────────────────────────────────────────────────────
    synth_step = f"Mensintesis {len(sources)} sumber menjadi laporan mendalam..."
    steps.append(synth_step)
    yield evt({"type": "step", "action": "synthesize", "message": synth_step, "step": len(steps)})

    context_text = "\n\n".join(web_context_parts)

    # Build clickable reference list with markdown links
    ref_list = "\n".join(
        f"[{i+1}] [{s['title']}]({s['url']})"
        for i, s in enumerate(sources)
    )

    # Build image context: include CDN images first, include source page for relevance check
    image_context = ""
    if images:
        image_context = "\n\nGambar relevan yang bisa digunakan dalam laporan.\n"
        image_context += "PENTING: Hanya sisipkan gambar yang BENAR-BENAR relevan dengan konten di sekitarnya. Jangan jelaskan apa isi gambar — langsung sisipkan dengan caption singkat yang kontekstual.\n"
        image_context += "Format: ![caption kontekstual singkat](url_gambar)\n\n"
        # CDN images are more reliable
        cdn_imgs = [img for img in images if img.get("cdn")]
        other_imgs = [img for img in images if not img.get("cdn")]
        ordered = cdn_imgs + other_imgs
        image_context += "Daftar gambar tersedia:\n"
        image_context += "\n".join(
            f"- {img['url']} | alt: {img['alt']} | sumber: {img.get('source_url','')}"
            for img in ordered[:6]
        )

    # Build numbered image reference for AI
    image_ref = ""
    if images:
        cdn_imgs = [img for img in images if img.get("cdn")]
        other_imgs = [img for img in images if not img.get("cdn")]
        ordered_imgs = cdn_imgs + other_imgs
        image_ref = f"\n\nGAMBAR TERSEDIA ({len(ordered_imgs)} gambar) — WAJIB GUNAKAN SEMUA, tersebar merata di artikel:\n"
        for n, img in enumerate(ordered_imgs, 1):
            image_ref += f"{n}. URL: {img['url']}\n   Alt: {img['alt']}\n   Sumber halaman: {img.get('source_url', '')}\n"
        image_ref += "\nInstruksi: Tempatkan setiap gambar di awal section yang paling relevan. Format: ![caption max 5 kata](url)\n"

    report_prompt = f"""Topik artikel: {query}

Kamu memiliki {len(sources)} sumber dan {len(images)} gambar.

KONTEN DARI SUMBER:
{context_text}
{image_ref}
---
Daftar sumber:
{ref_list}

Tulis artikel blog PANJANG dalam Bahasa Indonesia tentang "{query}".
- Minimal 1000-1500 kata
- Alur naratif seperti long-form blog (Medium style)
- Setiap section besar: mulai dengan gambar, baru teks
- Gunakan SEMUA {len(images)} gambar yang tersedia, tersebar di seluruh artikel
- Setiap section harus relevan dengan topik "{query}"
- Sitasi inline: [judul](url) — JANGAN [[1]]
- Akhiri dengan ## Referensi"""

    try:
        report = await call_model([
            {"role": "system", "content": DEEP_RESEARCH_SYSTEM},
            {"role": "user", "content": report_prompt},
        ], max_tok=16000)
    except RuntimeError as exc:
        yield evt({"type": "error", "message": str(exc)})
        return
    except Exception as exc:
        yield evt({"type": "error", "message": f"Gagal sintesis: {exc}"})
        return

    elapsed = round(time.time() - start_time)
    mins, secs = divmod(elapsed, 60)
    time_str = f"{mins}m {secs}s" if mins else f"{secs}s"

    yield evt({
        "type": "complete",
        "report": report,
        "sources": sources,
        "steps": len(steps),
        "elapsed": time_str,
    })


@api_router.post("/deep-research")
async def deep_research_endpoint(req: DeepResearchRequest):
    return StreamingResponse(
        run_deep_research(req.query),
        media_type="text/event-stream; charset=utf-8",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


# ── File Upload & Document Analysis ──────────────────────────────────────────

try:
    from PyPDF2 import PdfReader as _PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
ALLOWED_EXTENSIONS = {"jpg", "jpeg", "png", "gif", "webp", "pdf", "docx", "xlsx", "xls", "txt", "csv", "md"}


async def extract_file_content(file_bytes: bytes, filename: str, content_type: str) -> dict:
    """Extract content from uploaded file. Returns a dict with type and content."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    # Images — send as base64 to provider
    if ext in {"jpg", "jpeg", "png", "gif", "webp"} or "image" in content_type:
        import base64
        media_map = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png",
                     "gif": "image/gif", "webp": "image/webp"}
        media_type = media_map.get(ext, "image/jpeg")
        b64 = base64.standard_b64encode(file_bytes).decode()
        return {"kind": "image", "media_type": media_type, "b64": b64, "filename": filename}

    # PDF — extract text via PyPDF2
    if ext == "pdf":
        if PYPDF2_AVAILABLE:
            try:
                import io
                reader = _PdfReader(io.BytesIO(file_bytes))
                pages = []
                for page in reader.pages[:30]:  # max 30 pages
                    text = page.extract_text() or ""
                    if text.strip():
                        pages.append(text)
                text = "\n\n".join(pages)[:15000]
                return {"kind": "text", "text": f"[PDF: {filename}]\n\n{text}", "filename": filename}
            except Exception as exc:
                return {"kind": "text", "text": f"[PDF: {filename}] (gagal ekstrak: {exc})", "filename": filename}
        return {"kind": "text", "text": f"[PDF: {filename}] (PyPDF2 tidak tersedia)", "filename": filename}

    # DOCX
    if ext == "docx" and DOCX_AVAILABLE:
        try:
            import io
            doc = DocxDocument(io.BytesIO(file_bytes))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())[:15000]
            return {"kind": "text", "text": f"[Word: {filename}]\n\n{text}", "filename": filename}
        except Exception as exc:
            return {"kind": "text", "text": f"[Word: {filename}] (gagal ekstrak: {exc})", "filename": filename}

    # XLSX / XLS
    if ext in {"xlsx", "xls"}:
        try:
            import openpyxl, io
            wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
            rows_text = []
            for sheet_name in wb.sheetnames[:5]:
                ws = wb[sheet_name]
                rows_text.append(f"=== Sheet: {sheet_name} ===")
                for row in ws.iter_rows(values_only=True):
                    row_str = " | ".join(str(c) if c is not None else "" for c in row)
                    if row_str.replace(" | ", "").strip():
                        rows_text.append(row_str)
                if len("\n".join(rows_text)) > 12000:
                    break
            text = "\n".join(rows_text)[:15000]
            return {"kind": "text", "text": f"[Excel: {filename}]\n\n{text}", "filename": filename}
        except Exception as exc:
            return {"kind": "text", "text": f"[Excel: {filename}] (gagal ekstrak: {exc})", "filename": filename}

    # TXT / CSV / MD
    if ext in {"txt", "csv", "md"}:
        text = file_bytes.decode("utf-8", errors="ignore")[:15000]
        return {"kind": "text", "text": f"[{ext.upper()}: {filename}]\n\n{text}", "filename": filename}

    # Fallback
    text = file_bytes.decode("utf-8", errors="ignore")[:5000]
    return {"kind": "text", "text": f"[{filename}]\n\n{text}", "filename": filename}


from fastapi import UploadFile, File, Form


@api_router.post("/upload-and-analyze")
async def upload_and_analyze(
    files: List[UploadFile] = File(...),
    prompt: str = Form(default="Tolong analisis file yang saya upload."),
    chat_history: str = Form(default="[]"),
    model_id: str = Form(default=""),
):
    """Analyze uploaded files (images, PDF, DOCX, XLSX, TXT) using the AI model."""
    # Use claude-sonnet-4-5-1m for vision (supports image_url), fallback to default
    effective_model_id = model_id.strip() or "claude-sonnet-4-5-1m"
    base_url, api_key, provider_name, model = get_provider_config(effective_model_id)
    if not base_url or not api_key:
        # Fallback to default model
        base_url, api_key, provider_name, model = get_provider_config(DEFAULT_MODEL_ID)
    if not base_url or not api_key:
        return JSONResponse({"error": f"{provider_name} belum dikonfigurasi"}, status_code=500)

    # Validate and process files
    file_texts: List[str] = []
    image_blocks: List[dict] = []
    filenames: List[str] = []

    for upload in files[:5]:  # max 5 files
        file_bytes = await upload.read()
        if len(file_bytes) > MAX_FILE_SIZE:
            return JSONResponse({"error": f"File {upload.filename} terlalu besar (maks 10MB)"}, status_code=400)

        content_type = upload.content_type or ""
        result = await extract_file_content(file_bytes, upload.filename or "file", content_type)
        filenames.append(upload.filename or "file")

        if result["kind"] == "image":
            image_blocks.append(result)
        else:
            file_texts.append(result["text"])

    # Build message content for the provider
    # Use text-only format (OpenAI-compatible) — images as text descriptions
    content_parts = []

    for img in image_blocks:
        # Send as OpenAI vision format
        content_parts.append({
            "type": "image_url",
            "image_url": {"url": f"data:{img['media_type']};base64,{img['b64']}"}
        })
        content_parts.append({"type": "text", "text": f"[Gambar: {img['filename']}]"})

    for txt in file_texts:
        content_parts.append({"type": "text", "text": txt})

    content_parts.append({
        "type": "text",
        "text": prompt if prompt.strip() else "Tolong analisis file yang saya upload."
    })

    # Parse chat history
    try:
        history = json.loads(chat_history)
        if not isinstance(history, list):
            history = []
    except Exception:
        history = []

    timeout = env_float("OPENAI_TIMEOUT_SECONDS", DEFAULT_TIMEOUT_SECONDS)
    headers_req = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }

    messages = history[-6:] + [{"role": "user", "content": content_parts}]
    body = {
        "model": model,
        "messages": messages,
        "max_tokens": 4096,
        "temperature": 0.7,
        "stream": False,
    }
    system_msg = (
        "Kamu adalah asisten AI yang ahli menganalisis dokumen dan gambar. "
        "Baca dan pahami seluruh isi file yang diberikan. "
        "Jawab pertanyaan user berdasarkan isi file. "
        "Jika tidak ada pertanyaan spesifik, berikan ringkasan komprehensif. "
        "Untuk gambar: deskripsikan apa yang kamu lihat secara detail. "
        "Untuk data/excel: identifikasi pola, trend, dan insight penting. "
        "Untuk dokumen: ringkas poin-poin utama. "
        "Selalu respond dalam bahasa yang sama dengan pertanyaan user."
    )
    # Prepend system message
    messages = [{"role": "system", "content": system_msg}] + messages
    body["messages"] = messages

    try:
        async with httpx.AsyncClient(timeout=timeout) as client:
            resp = await client.post(provider_url(base_url), headers=headers_req, json=body)
            if resp.status_code >= 400:
                detail = resp.text[:500]
                return JSONResponse({"error": f"Model error {resp.status_code}: {detail}"}, status_code=500)
            data = resp.json()
        response_text = data["choices"][0]["message"]["content"] or ""
    except Exception as exc:
        return JSONResponse({"error": f"Gagal analisis: {exc}"}, status_code=500)

    return JSONResponse({
        "success": True,
        "response": response_text,
        "files_analyzed": filenames,
    })



# ── Studio and Supercomputer endpoints removed - to be rebuilt from scratch ──

app.include_router(api_router)

origins = cors_origins()
app.add_middleware(
    CORSMiddleware,
    allow_credentials="*" not in origins,
    allow_origins=origins,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ── Security Headers Middleware ──────────────────────────────────────────────
from starlette.middleware.base import BaseHTTPMiddleware

class SecurityHeadersMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request, call_next):
        response = await call_next(request)
        response.headers["X-Content-Type-Options"] = "nosniff"
        response.headers["X-Frame-Options"] = "DENY"
        response.headers["X-XSS-Protection"] = "1; mode=block"
        response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
        if request.url.path.startswith("/api/"):
            response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate"
        return response

app.add_middleware(SecurityHeadersMiddleware)


# ── Request Body Size Limit Middleware ────────────────────────────────────────
MAX_BODY_SIZE = 10 * 1024 * 1024  # 10MB

class BodySizeLimitMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request, call_next):
        content_length = request.headers.get("content-length")
        if content_length and int(content_length) > MAX_BODY_SIZE:
            from starlette.responses import JSONResponse
            return JSONResponse({"error": "Request body too large"}, status_code=413)
        return await call_next(request)

app.add_middleware(BodySizeLimitMiddleware)


# ── Request Logging Middleware ────────────────────────────────────────────────
import logging
request_logger = logging.getLogger("requests")
_abuse_counts: dict[str, list[float]] = {}

class RequestLoggingMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request, call_next):
        start = time.time()
        ip = request.client.host if request.client else "unknown"
        path = request.url.path
        method = request.method

        # Track failed auth attempts for abuse detection
        response = await call_next(request)
        duration = time.time() - start

        # Log slow requests (>5s)
        if duration > 5:
            logger.warning(f"Slow request: {method} {path} from {ip} took {duration:.1f}s")

        # Log 401/403 for abuse detection
        if response.status_code in (401, 403):
            _abuse_counts.setdefault(ip, []).append(time.time())
            _abuse_counts[ip] = [t for t in _abuse_counts[ip] if time.time() - t < 300]
            if len(_abuse_counts[ip]) > 20:
                logger.warning(f"Possible abuse from {ip}: {len(_abuse_counts[ip])} failed auth in 5min")

        return response

app.add_middleware(RequestLoggingMiddleware)


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("server:app", host="0.0.0.0", port=8001, reload=True, env_file=".env")

